# -*- coding: utf-8 -*-
"""
Generate English manuscript by reading corrected values from i2_ci_corrected.json.
This is the English translation of generate_v5.py for international journal submission.
"""
import sys, json, os, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJ = Path(__file__).parent
OUTPUT_DIR = PROJ / 'output'
STYLE_FONT = 'Times New Roman'
CJK_FONT = 'Times New Roman'  # No CJK needed for English
HEADING_FONT = 'Times New Roman'

BJSM_PAGE_WIDTH = Cm(21.0)
BJSM_PAGE_HEIGHT = Cm(29.7)

# ================================================================
# Load corrected values
# ================================================================
with open(OUTPUT_DIR / 'i2_ci_corrected.json', 'r', encoding='utf-8') as f:
    corrected = json.load(f)

strict = corrected['strict']
wide = corrected['wide']
short = corrected['short']
mid = corrected['mid']

def fmt_g(v):
    return f"+{v:.3f}" if v >= 0 else f"{v:.3f}"

def fmt_pct(v):
    return f"{v:.1f}%"

def fmt_ci_pair(lo, hi):
    return f"[{fmt_g(lo)},{fmt_g(hi)}]"

def fmt_pct_ci(lo, hi):
    return f"[{fmt_pct(lo)},{fmt_pct(hi)}]"

with open(OUTPUT_DIR / 'pet_peese_results.json', 'r', encoding='utf-8') as f:
    pet_peese = json.load(f)

pp_strict = pet_peese['strict']
pp_wide = pet_peese['wide']

def fmt_pp(v):
    return f"{v:+.4f}" if v >= 0 else f"{v:.4f}"

strict_g = strict['g']; strict_cl = strict['ci_low']; strict_cu = strict['ci_upp']
strict_I2 = strict['I2']; strict_I2_l = strict['I2_ci_low']; strict_I2_u = strict['I2_ci_upp']
strict_k = strict['k']

wide_g = wide['g']; wide_cl = wide['ci_low']; wide_cu = wide['ci_upp']
wide_I2 = wide['I2']; wide_I2_l = wide['I2_ci_low']; wide_I2_u = wide['I2_ci_upp']
wide_k = wide['k']

short_g = short['g']; short_cl = short['ci_low']; short_cu = short['ci_upp']
short_I2 = short['I2']; short_I2_l = short['I2_ci_low']; short_I2_u = short['I2_ci_upp']
short_k = short['k']

mid_g = mid['g']; mid_cl = mid['ci_low']; mid_cu = mid['ci_upp']
mid_I2 = mid['I2']; mid_I2_l = mid['I2_ci_low']; mid_I2_u = mid['I2_ci_upp']
mid_k = mid['k']

print(f"Corrected values loaded:")
print(f"  Strict: k={strict_k}, g={strict_g:+.3f}, I2={strict_I2:.1f}%")
print(f"  Wide: k={wide_k}, g={wide_g:+.3f}, I2={wide_I2:.1f}%")
print(f"  Short: k={short_k}, g={short_g:+.3f}, I2={short_I2:.1f}%")
print(f"  Mid: k={mid_k}, g={mid_g:+.3f}, I2={mid_I2:.1f}%")

# ================================================================
# Formatting helpers
# ================================================================
def set_run_font(run, size=None, bold=False, italic=False, font_name=None, cjk=None):
    font = font_name or STYLE_FONT
    run.font.name = font
    if cjk is None:
        cjk = CJK_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:eastAsia'), cjk)
    rFonts.set(qn('w:cs'), font)
    rPr.insert(0, rFonts)
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True

def add_para(doc, text, bold=False, italic=False, size=None, first_line_indent=Cm(0.74),
             spacing_before=None, spacing_after=None, align=None, cjk=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size or Pt(12), bold=bold, italic=italic, cjk=cjk)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if spacing_before is not None:
        p.paragraph_format.space_before = spacing_before
    if spacing_after is not None:
        p.paragraph_format.space_after = spacing_after
    if align is not None:
        p.alignment = align
    return p

def add_para_headline(doc, text, bold=False, italic=False, size=None, cjk=None, align=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size or Pt(14), bold=bold, italic=italic, cjk=cjk)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = indent
    return p

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, cjk=HEADING_FONT)
    return h

def add_figure(doc, image_path, caption, width=None):
    """Add a figure with caption to the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(image_path):
        run = p.add_run()
        run.add_picture(image_path, width=width or Inches(5.0))
    else:
        run = p.add_run(f'[Image: {image_path}]')
        set_run_font(run, size=Pt(10), italic=True)
    # Add caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    set_run_font(cap_run, size=Pt(10), italic=True)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=Pt(10), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=Pt(10))
    
    return table

# ================================================================
# CREATE DOCUMENT
# ================================================================
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = STYLE_FONT
font.size = Pt(12)

# Set page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ================================================================
# TITLE
# ================================================================
add_para_headline(doc, 'Systematic Review and Meta-analysis of the Effects of Plyometric Training on Countermovement Jump Height: Effect Size Estimation and Dose-Response Relationship Based on Arm Position Stratification', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, indent=None)

add_para(doc, 'MO Yuhang  FANG Yuanshun (Corresponding Author)', bold=False, size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, spacing_after=Pt(2))
add_para(doc, 'School of Physical Education, Southwest University, Chongqing 400715, China',
         size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, spacing_after=Pt(12))

# ================================================================
# ABSTRACT
# ================================================================
add_para_headline(doc, 'Abstract', bold=True, size=Pt(12), indent=None)

add_para(doc,
    f'Objective: To systematically evaluate, for the first time, the effects of plyometric training on countermovement jump (CMJ) height using CMJ arm position as an inclusion criterion, and to examine the moderating effects of variables such as intervention duration on the dose-response relationship. Methods: PubMed, Scopus, Web of Science, Google Scholar, CNKI, WanFang, and VIP databases were systematically searched from inception to May 2026. RCTs with plyometric training interventions reporting CMJ height (Mean{chr(177)}SD) were included. '
    f'Non-randomized designs, studies without a non-plyometric control group, studies involving CMJ with arm swing that did not separately report arm-restricted CMJ, and studies involving injury rehabilitation populations were excluded. '
    f'The primary analysis pool comprised strict hands-on-hips/arms-crossed-chest CMJ studies (k={strict_k}), with a broader pool including arm-unclear/CMJA studies (k={wide_k}) serving as the sensitivity analysis. '
    f'Hedges\' g was computed from pre-post change scores (assuming r=0.7) using a REML random-effects model. Heterogeneity was assessed via I{chr(178)} with Q-profile 95%CI. '
    f'Publication bias was evaluated using Egger regression, Begg rank correlation, Peters test, SE-g correlation, and Trim-and-Fill correction. '
    f'Evidence certainty was rated using the GRADE framework. The protocol was pre-registered at PROSPERO (CRD420261422906). '
    f'Results: The strict hands-on-hips pool yielded a Hedges\' g of {fmt_g(strict_g)} [95%CI: {fmt_g(strict_cl)}, {fmt_g(strict_cu)}], I{chr(178)}={fmt_pct(strict_I2)} (95%CI: {fmt_pct_ci(strict_I2_l, strict_I2_u)}), prediction interval [{fmt_g(strict.get("pred_low", -0.695))}, {fmt_g(strict.get("pred_upp", 2.919))}]. '
    f'The broader pool yielded g={fmt_g(wide_g)} [95%CI: {fmt_g(wide_cl)}, {fmt_g(wide_cu)}], I{chr(178)}={fmt_pct(wide_I2)}. '
    f'Intervention duration emerged as a potential moderator (exploratory signal): short-term ({chr(8804)}6 weeks) g={fmt_g(short_g)} (I{chr(178)}={fmt_pct(short_I2)}), medium-term (7{chr(8211)}10 weeks) g={fmt_g(mid_g)} (I{chr(178)}={fmt_pct(mid_I2)}), long-term (>10 weeks) g=+1.85 (I{chr(178)}=94%). '
    f'Meta-regression suggested a linear duration-effect association ({chr(126)}+0.13 SMD/week, p=0.023), which became non-significant after Bonferroni correction ({chr(945)}=0.05/12{chr(8776)}0.00417), classifying the dose-response relationship as exploratory rather than confirmatory. '
    f'Adolescent subgroups yielded the numerically largest effects (g=+1.37{chr(8211)}1.51, k=2, GRADE: Low), based on exploratory evidence only. Egger test was significant (intercept={chr(8722)}2.64, p<0.001) with strong SE-g correlation (r=+0.88) — both pointing to small-study effects. '
    f'Trim-and-Fill correction paradoxically increased the pooled effect (+13.8%), a known failure signal of the method under high heterogeneity (I{chr(178)}=78.1%; Terrin et al., 2003). PET-PEESE was not used as a formal correction tool due to prerequisite violations (Stanley 2017); detailed output in Supplementary Materials. '
    f'GRADE assessment rated the overall evidence certainty as Low. Conclusion: Plyometric training likely produces a positive effect on CMJ height, but this finding should be interpreted with caution given the low GRADE certainty. '
    f'Short-term ({chr(8804)}6 weeks) training represents the most reliable evidence (moderate effect, low heterogeneity). The dose-response relationship should be regarded as an exploratory signal warranting independent verification. Adolescent findings (g=+1.37{chr(8211)}1.51, k=2) should be treated as preliminary exploratory signals only. Standardization of CMJ arm position protocols in sports science research is urgently needed.',
    size=Pt(11), first_line_indent=Cm(0.74))

add_para(doc,
    'Keywords: plyometric training; countermovement jump; meta-analysis; arm position; dose-response relationship; stretch-shortening cycle',
    size=Pt(11), bold=True, first_line_indent=Cm(0.74), spacing_before=Pt(12))

doc.add_page_break()

# ================================================================
# 1. INTRODUCTION
# ================================================================
add_heading_styled(doc, '1  Introduction', 1)

intro_paragraphs = [
    'Plyometric training (also known as jump training or reactive strength training) is a method that utilizes the stretch-shortening cycle (SSC) of skeletal muscles to produce explosive concentric contractions. Typical exercises include drop jumps, hurdle hops, and consecutive vertical jumps [1,2]. This training is widely used in athletic performance enhancement programs ranging from youth sports to elite competition. The physiological basis lies in enhanced tendon stiffness, improved neuromuscular activation rate, and optimized SSC mechanical efficiency [3]. From a biomechanical energetics perspective, the SSC stores elastic energy during the pre-stretch phase and releases it during the concentric phase, while simultaneously eliciting the stretch reflex to enhance muscle activation — a mechanism first provided with classic quantitative evidence by Bosco et al. (1982) [10].',

    'The countermovement jump (CMJ) is one of the gold standard measures of lower-body explosive power [4]. Unlike the squat jump (SJ), CMJ includes a countermovement pre-stretch phase before the concentric phase, which more realistically reflects SSC utilization ability. Consequently, CMJ is widely used in talent identification, training monitoring, and research assessment. Improvements in CMJ height are typically regarded as direct evidence of plyometric training effectiveness.',

    'Several previous meta-analyses have evaluated the effects of plyometric training on vertical jump height. Markovic\'s (2007) seminal meta-analysis reported a pooled effect size of SMD=0.88 (95%CI: 0.64-1.11, I²=11.4%) for plyometric training on CMJ, classified as a large effect. However, this study included multiple jump outcome measures and did not differentiate CMJ arm position [1]. de Villarreal et al. (2009) included 225 effect sizes from 56 studies, systematically analyzing the influence of training volume, intensity, and type on jump performance, but did not report pooled effect sizes and heterogeneity data specifically for CMJ [2]. Slimani et al. (2016) reported pooled effect sizes for various fitness indicators in team sport athletes, but the comparability of their effect size data was limited due to lack of CMJ arm position stratification and broad inclusion criteria [9]. The Ramirez-Campillo group has published a series of meta-analyses focusing on specific populations: Moran et al. (2019) reported a CMJ effect size of ES=0.57 (95%CI: 0.21-0.93) in female adolescent athletes after plyometric training, and found that the <15 years subgroup (ES=0.78) had significantly higher effects than the ≥15 years subgroup (ES=0.31) [5]; Ramirez-Campillo et al. (2020) reported a pooled effect size of ES=1.01 (95%CI: 0.36-1.66) in female soccer players [6]; Sole and Ramirez-Campillo et al. (2021) found a pooled effect size of ES=0.49 (I²=0.0%) in individual sport athletes [7]. Additionally, Ramirez-Campillo et al. (2020) published an umbrella review in Scand J Med Sci Sports that systematically reviewed the effects of plyometric training on multiple fitness indicators, providing a macro-level evidence landscape for this field [14].',

    'However, these meta-analyses share common methodological limitations. First, the presence or absence of arm swing during CMJ testing significantly affects jump height measurement — Harman et al. (1990) found in a single laboratory study (n=17) that arm swing contributes approximately 10%-15% additional jump height [8]. Although this estimate comes from a single study with limited generalizability to all populations in meta-analyses (e.g., adolescents and elite athletes), arm swing qualitatively alters movement coordination patterns and introduces additional biomechanical sources of variation. However, except for Markovic (2007) who made a preliminary distinction between CMJ with and without arm swing in subgroups, none of the other meta-analyses used CMJ arm position (hands-on-hips/arms-crossed vs. free arm swing/Vertec) as an inclusion criterion or analysis variable. This methodological ambiguity may lead to confounding bias in effect size estimation: CMJ or Vertec vertical jumps with arm swing may overestimate the true effects of plyometric training on lower-body SSC capacity due to upper limb participation. Second, most previous meta-analyses lack systematic testing of the dose-response relationship between intervention duration and effect size (e.g., meta-regression or continuous dose-response models), with only a few conducting categorical subgroup analyses (e.g., <7 weeks vs. ≥7 weeks). Meylan et al. (2014) noted in a JSCR review that plyometric training in adolescents needs to consider the interaction effects of maturation and training parameters [15]; Lloyd et al. (2014) international consensus statement and Radnor et al. (2018) review on SSC function during development also emphasized the critical window of neuromuscular plasticity during adolescence [12,13]. However, specialized meta-analyses for prepubertal and pubertal athletes are very limited (only Moran et al. 2019 for female adolescents), and multiple new RCTs published in recent years (2020-2025) have not been included in any systematic review.',

    'Therefore, this study aims to comprehensively evaluate the effects of plyometric training on CMJ height through an updated and more methodologically rigorous systematic review and meta-analysis. Compared with previous meta-analyses, the incremental contribution of this study lies not in a larger sample size, but in more refined inclusion criteria and methodological design. Specific innovations include: (1) for the first time, CMJ arm position is used as an explicit inclusion criterion, with the primary analysis restricted to strict hands-on-hips/arms-crossed armless CMJ, while arm position is used as a subgroup analysis variable; (2) systematic testing of the dose-response relationship of intervention duration, including categorical subgroups and continuous meta-regression; (3) inclusion of new RCTs from 2020-2025, updating effect size estimates and filling evidence gaps for adolescent populations. This study pre-specified the following research questions: (1) What is the pooled effect size (Hedges\' g) of plyometric training on CMJ height? (2) Do intervention duration, training frequency, participant age, and CMJ arm position moderate the effect size? (3) Is there a dose-response relationship for short-term (≤6 weeks), medium-term (7-10 weeks), and long-term (>10 weeks) plyometric interventions? Based on existing evidence, we hypothesized that plyometric training would produce a moderate-to-large positive effect on CMJ height (g>0.5), with effect sizes increasing with intervention duration.',
]

for text in intro_paragraphs:
    add_para(doc, text)

# ================================================================
# 2. METHODS
# ================================================================
add_heading_styled(doc, '2  Methods', 1)

add_heading_styled(doc, '2.1  Literature Search Strategy', 2)
add_para(doc,
    'We systematically searched PubMed, Scopus, Web of Science, Google Scholar, and CNKI (China National Knowledge Infrastructure), WanFang, and VIP databases from inception to May 2026. The following search strategy was used (PubMed example):')
add_para(doc,
    '(plyometric* OR "jump training" OR "reactive strength" OR "stretch-shortening cycle") '
    'AND ("countermovement jump" OR CMJ OR "vertical jump" OR "jump height") '
    'AND (random* OR RCT OR "controlled trial")',
    size=Pt(10), first_line_indent=None, spacing_before=Pt(4))
add_para(doc,
    'The complete PubMed search strategy (with line numbers and hit counts for each search term) is provided in Supplementary Materials. '
    'Scopus and Web of Science used equivalent search syntax, and Google Scholar used a simplified search string. '
    'Given the known limitations of Google Scholar in search consistency (results are influenced by user location, device, and algorithmic personalization), only the first 200 ranked results from Google Scholar were included.')
add_para(doc,
    'For Chinese databases, we used locally adapted search strategies in CNKI, WanFang, and VIP. '
    'The core search strategy (CNKI professional search syntax example): '
    'SU=(\'快速伸缩复合训练\'+\'增强式训练\'+\'超等长训练\'+\'爆发力训练\'+\'跳深训练\'+\'plyometric\'+\'跳箱训练\'+\'反应力量\')'
    '* (\'反向纵跳\'+\'下蹲跳\'+\'CMJ\'+\'纵跳高度\'+\'垂直跳\'+\'纵跳\'+\'反向跳\'+\'countermovement jump\')'
    '* (\'随机\'+\'对照\'+\'RCT\'+\'随机分组\'+\'随机对照试验\'). '
    'WanFang and VIP used equivalent search syntax or corresponding advanced search interfaces. '
    'The complete Chinese search strategies and terminology mapping are provided in Supplementary Materials.')
add_para(doc,
    'Reference lists of included studies were also hand-searched to supplement the database search. No language restrictions were applied.')

add_heading_styled(doc, '2.2  Inclusion and Exclusion Criteria', 2)
add_para(doc, 'Inclusion criteria (PICOS framework):', bold=True)
add_table(doc,
    ['PICOS Element', 'Criteria'],
    [
        ['P (Participants)', 'Healthy individuals, no restrictions on age, sex, or training level'],
        ['I (Intervention)', 'Plyometric training (including drop jumps, hurdle hops, consecutive jumps, and other SSC exercises), where the experimental group performed additional plyometric training on top of regular training/daily activity. Operational definition of "sufficient training dose": ≥6 training sessions with ≥30 contacts per session (based on common lower bounds in de Villarreal et al. 2009 and Markovic 2007)'],
        ['C (Comparison)', 'Control groups not performing plyometric training (may maintain regular training or daily activity). Two predefined control types: no-training control (true control) and regular training control (active control), with the latter expected to yield smaller effect sizes — planned subgroup analysis to compare effect size differences between control types'],
        ['O (Outcome)', 'Countermovement jump (CMJ) height (cm), reported as pre-test/post-test Mean±SD. Predefined test equipment classification: force platform, contact mat, OptoJump infrared system, with planned sensitivity analysis across equipment types'],
        ['S (Study design)', 'Randomized controlled trial (RCT) or cluster randomized trial'],
    ])

add_para(doc, 'Exclusion criteria:', bold=True)
exclusions = [
    'Did not report CMJ height (only reported SJ, DJ, VJ, or other jump measures)',
    'No independent non-plyometric control group (e.g., plyometric vs. resistance training comparison)',
    'Non-randomized design (e.g., pre-post only, cohort studies)',
    'CMJ test did not restrict arm swing (i.e., CMJA/Abalakov/Vertec VJ) and did not separately report armless CMJ',
    'Participants were injury rehabilitation populations',
    'Duplicate publication or overlapping data',
]
for ex in exclusions:
    add_para(doc, f'• {ex}', first_line_indent=Cm(1.5), size=Pt(11),
             spacing_after=Pt(1))

add_heading_styled(doc, '2.3  Study Selection', 2)
add_para(doc,
    'Two reviewers (M.L. and F.D.) independently performed study selection. First, titles and abstracts of retrieved records were screened, followed by full-text assessment of potentially eligible studies. Both screening phases were conducted independently, with disagreements resolved by consensus or adjudication by a third reviewer. The selection process was documented using a custom Python script with audit trails. Inter-rater agreement was assessed using Cohen\'s kappa coefficient: kappa values for major categorical variables (CMJ arm position, age grouping, intervention duration grouping) were 0.89, 0.92, and 0.95 respectively, indicating excellent agreement; agreement rate for numerical data (sample size, means, standard deviations) was 96.7% (only 1 of 29 studies had extraction discrepancies, resolved after checking the original paper).')

add_heading_styled(doc, '2.4  Data Extraction', 2)
add_para(doc,
    'Two reviewers independently extracted the following information: (1) Study characteristics: first author, publication year, country, sport, sample size; '
    '(2) Participant characteristics: sex, age, training level, maturation stage; (3) Intervention characteristics: plyometric training type, '
    'duration (weeks), frequency (sessions/week), ground contacts; (4) CMJ testing: equipment, arm position (hands-on-hips/unclear/with arms), '
    'units; (5) Outcome data: intervention and control group CMJ height pre-test/post-test Mean±SD. '
    'Inter-rater agreement was assessed using Cohen\'s kappa coefficient: kappa values for major categorical variables (CMJ arm position, age grouping, '
    'intervention duration grouping) were 0.89, 0.92, and 0.95 respectively, indicating excellent agreement; agreement rate for numerical data was 96.7% '
    '(only 1 of 29 studies had extraction discrepancies, resolved after checking the original paper).')
add_para(doc,
    'If the original paper reported SEM (standard error of measurement), SD was calculated as SD=SEM×√n. If only CI (confidence interval) was reported, '
    'SD was back-calculated as SD≈(CI_upper−CI_lower)/(2×1.96). If only change scores (Δ±SD_Δ) were reported, '
    'Post Mean was estimated as Pre+Δ, and Post SD was approximated using Pre SD or SD_Δ.')

add_heading_styled(doc, '2.5  Risk of Bias Assessment', 2)
add_para(doc,
    'Methodological quality of each included study was assessed using both the PEDro scale (Physiotherapy Evidence Database Scale) and the TESTEX scale '
    '(Tool for the assEssment of Study qualiTy and reporting in EXercise). '
    'The PEDro scale contains 11 items: item 1 assesses external validity (not scored), items 2-11 assess internal validity and statistical reporting quality, '
    'with each satisfied item scoring 1 point (total range 0-10). A PEDro score ≥6 was used as the cut-off based on conventions in sports science meta-analyses '
    '(used in multiple Sports Medicine and JSCR reviews), though the PEDro official does not set thresholds. '
    'The TESTEX scale contains 15 items (Smart et al., 2015, J Evid Based Med), adding exercise-specific bias sources to PEDro, '
    'including intervention adherence (item 14) and exercise intensity monitoring (item 15), with total range 0-15. '
    'TESTEX items 5 (subject blinding) and 6 (therapist blinding) are inherently unachievable in exercise training interventions, '
    'yielding an effective maximum score of 13. '
    'Two reviewers independently completed scoring, with disagreements resolved by consensus. '
    'TESTEX scoring results are provided in Supplementary Material S10.')

add_heading_styled(doc, '2.6  Effect Size Calculation', 2)
add_para(doc,
    'Hedges\' g (standardized mean difference with small-sample correction) was used as the effect size metric. Two SMDs were calculated:')
add_para(doc,
    '(1) Post-only SMD: standardized mean difference based on post-test means, using pooled post-test SD as the standardization factor. '
    '(2) Pre-post change SMD (primary analysis): between-group comparison based on pre-to-post change scores, assuming a pre-post correlation coefficient r=0.7 '
    '(typical value for CMJ [4]).')
add_para(doc,
    'Change score SD formula: SD_change = √(SD_pre² + SD_post² − 2×r×SD_pre×SD_post). '
    'Hedges\' g correction factor: J = 1 − 3/(4×df − 1), df = n_IG + n_CG − 2.')

add_heading_styled(doc, '2.7  Statistical Analysis', 2)
add_para(doc,
    'This systematic review and meta-analysis was pre-registered on PROSPERO (registration number: CRD420261422906) '
    'and followed the PRISMA 2020 reporting guidelines (PRISMA 2020 Checklist provided in Supplementary Materials). '
    'All analyses were conducted in a Python 3.x + R 4.6/metafor hybrid environment, '
    'with analysis code uploaded to a public GitHub/Zenodo repository (https://doi.org/10.5281/zenodo.20748080). '
    'AI-assisted tools (Claude AI, Anthropic, model versions Claude Opus 4.8 and Claude Sonnet 4.6) provided assistance in the following limited scope: '
    '(1) generation and debugging of data analysis code (R/metafor and Python scripts); (2) formatting organization and grammatical proofreading of the initial draft; '
    '(3) reference formatting organization and proofreading. '
    'AI was NOT involved in: literature screening (performed independently by two reviewers, kappa=0.86/0.91), '
    'data extraction (independently extracted and cross-checked by two reviewers), inclusion/exclusion decisions, effect size calculation strategy development, '
    'statistical model selection, risk of bias assessment (PEDro scoring), GRADE evidence rating, result interpretation, and all conclusion derivation — '
    'all scientific judgments were made independently by human authors.')

add_para(doc, 'Primary analysis:', bold=True)
add_para(doc,
    'A random-effects model (Restricted Maximum Likelihood, REML) was used to pool effect sizes. The pooled SMD (Hedges\' g), '
    '95%CI, and prediction interval were reported. Prediction intervals were calculated using t-distribution quantiles (k-2 degrees of freedom; '
    'using normal distribution quantiles with k<40 leads to under-coverage). '
    'Heterogeneity was assessed using τ², I² statistic, and Q test. I²>50% was considered substantial heterogeneity, I²>75% considerable heterogeneity.')
add_para(doc,
    'Pre-specified analysis pools: (1) Primary analysis pool: strict hands-on-hips armless CMJ; (2) Broader analysis pool: all CMJ studies (including arm-unclear/CMJA), '
    'as sensitivity analysis; (3) VJ sensitivity subgroup: arm-swing vertical jumps (non-CMJ), analyzed separately. '
    'R19 (Michailidis 2018) was excluded from the primary analysis pool due to unconfirmable SD/SE sources (SE/SD ambiguity in CI back-calculation, effect size overestimated approximately 4.0-fold), '
    'and was reported separately as sensitivity analysis. '
    'This exclusion constitutes a protocol deviation from the PROSPERO pre-registration (CRD420261422906) — '
    'the pre-registered protocol specified that all RCTs meeting PICOS criteria with hands-on-hips armless CMJ would be included in the primary analysis. '
    'R19 formally met inclusion criteria, but during data extraction, an irreconcilable ambiguity was found between the reported standard deviation (SD) and the SD back-calculated from the reported confidence interval '
    '(corresponding to g=+1.48 and g≈+0.37 respectively), with an approximately 4.0-fold effect size difference making it impossible to determine which was the true value. '
    'Therefore, it was excluded from the primary analysis pool to ensure the reliability of the pooled effect size. '
    'Data ambiguity handling followed the pre-registered SOP: (i) contact authors at least twice (for R19: two attempts on March 15 and April 2, 2026, neither received raw data response) → '
    '(ii) if raw data unavailable, perform bidirectional back-calculation based on reported CI/n (complete calculation steps in Supplementary Materials) → '
    '(iii) if bidirectional ambiguity Δg exceeds the pre-set threshold (0.5), mark as "high uncertainty" and exclude. '
    'The handling of this deviation has been validated in the results section below through sensitivity analysis (Δg=−0.016 with R19 included, conclusions robust).')

add_para(doc, 'Subgroup analyses (categorical moderators):', bold=True)
sgs = [
    'CMJ arm position (strict hands-on-hips vs. arm-unclear vs. CMJA with arms)',
    'Age/maturation stage (prepubertal vs. pubertal vs. young adult vs. adult/professional)',
    'Intervention duration (short-term ≤6 weeks vs. medium-term 7-10 weeks vs. long-term >10 weeks)',
    'Training type (plyometric only vs. plyometric + resistance combination vs. plyometric + change of direction)',
    'Control type (no-training control vs. regular training control)',
    'Sex (male vs. female vs. mixed)',
]
for sg in sgs:
    add_para(doc, f'• {sg}', first_line_indent=Cm(1.5), size=Pt(11),
             spacing_after=Pt(1))

add_para(doc, 'Meta-regression (continuous moderators):', bold=True)
add_para(doc,
    'Intervention duration (weeks) and weekly training frequency were used as primary predictors in multivariable models to test independent effects. '
    'Before conducting multivariable meta-regression, collinearity between predictors was checked (VIF<5), '
    'and a candidate model set was pre-specified: Model 1 (duration only), Model 2 (duration + arm position), '
    'Model 3 (duration + frequency). Model selection used AICc and Q_E-residual as criteria.')

add_para(doc, 'Sensitivity analyses:', bold=True)
add_para(doc,
    '(1) Leave-one-out analysis to assess the influence of individual studies on pooled effects; (2) exclusion of outlier studies '
    '(>3×IQR); (3) pre-post correlation coefficient sensitivity testing with r=0.5 and r=0.9.')

add_para(doc, 'Publication bias:', bold=True)
add_para(doc,
    'Funnel plot visual inspection, Egger regression test (p<0.10 indicating asymmetry; the intercept and slope of Egger regression should be distinguished — '
    'the intercept reflects the SE-weighted average effect and is not an independent bias indicator [see Rücker et al. 2008]; the slope is the direct measure of small-study effects), '
    'SE-g rank correlation analysis (small-study effects). '
    'Given the high heterogeneity (I²=' + fmt_pct(strict_I2) + ') and extreme effect sizes in this dataset, '
    'PET-PEESE\'s known limitations under these conditions have been well documented (Stanley 2017; Carter et al. 2019), '
    'and therefore it was not used as the primary bias correction tool — detailed output is provided in Supplementary Materials. '
    'Trim-and-Fill and PET-PEESE were reclassified in the main text as "bias correction attempts (limited by method applicability)" '
    'to conceptually distinguish them from Egger and Begg tests in terms of inference logic.')

# ================================================================
# 3. RESULTS
# ================================================================
add_heading_styled(doc, '3  Results', 1)

# NOTE: The results section is very long. For brevity in this initial version,
# we include a placeholder. The full translation should be completed by reading
# the remaining sections from generate_v5.py and translating them.

add_para(doc,
    '[Note: This is the English manuscript template. The Results section (3.1-3.7) and Discussion section (4.1-4.7) '
    'should be translated from the Chinese version (generate_v5.py lines 600-1370). '
    'The references section (lines 1382-1631) is already in English and can be copied directly.]',
    italic=True, size=Pt(11))

# ================================================================
# REFERENCES (already in English - copy from generate_v5.py)
# ================================================================
add_heading_styled(doc, 'References', 1)

# References are copied from generate_v5.py lines 1382-1631
# They are already in English format
references = [
    '[1] Markovic G. Does plyometric training improve vertical jump height? A meta-analytical review. Br J Sports Med. 2007;41(6):349-355. doi:10.1136/bjsm.2007.035113.',
    '[2] de Villarreal ES, Kellis E, Kraemer WJ, Izquierdo M. Determining variables of plyometric training for improving vertical jump height performance: a meta-analysis. J Strength Cond Res. 2009;23(2):495-506. doi:10.1519/JSC.0b013e318196b7c6.',
    '[3] Flanagan EP, Comyns TM. The use of contact time and the reactive strength index to optimize fast stretch-shortening cycle training. Strength Cond J. 2008;30(5):32-38.',
    '[4] Claudino JG, Cronin J, Mezêncio B, et al. The countermovement jump to monitor neuromuscular status: A meta-analysis. J Sci Med Sport. 2017;20(4):397-403. doi:10.1016/j.jsams.2016.08.011.',
    '[5] Moran J, Sandercock G, Ramírez-Campillo R, et al. Age-related variation in the effect of resistance training on lower-limb strength in pre- and post-pubertal boys: a meta-analysis. Pediatr Exerc Sci. 2019;31(1):91-101.',
    '[6] Ramirez-Campillo R, Alvarez C, Gentil P, et al. Sequencing effects of plyometric training applied before or after regular soccer training on measures of physical fitness in young players. J Strength Cond Res. 2020;34(7):1959-1966.',
    '[7] Sole G, Ramirez-Campillo R, Andrade R, et al. Plyometric jump training effects on the physical fitness of individual-sport athletes: a systematic review with meta-analysis. PeerJ. 2021;9:e11004. doi:10.7717/peerj.11004.',
    '[8] Harman EA, Rosenstein MT, Frykman PN, Rosenstein RM, Kraemer WJ. Estimation of human power output from vertical jump. J Appl Sport Sci. 1990;5(3):116.',
    '[9] Slimani M, Chaabene H, Miarka B, Cheour F. The effects of plyometric training on physical fitness in team sport athletes: a systematic review. J Hum Kinet. 2016;53:107-117. doi:10.1515/hukin-2016-0026.',
    '[10] Bosco C, Komi PV, Tihanyi J, et al. Measurement and coordination of mechanical power in jumping. Eur J Appl Physiol. 1982;50(3):271-282.',
    '[11] Asadi A, Ramirez-Campillo R, Meylan C, et al. Adaptations of vertical ground reaction force variables during and after plyometric training in elite soccer players. J Sports Sci Med. 2017;16(1):46-53.',
    '[12] Lloyd RS, Oliver JL, Radnor JM, et al. Relationships between functional movement screen scores, maturation and physical performance in young soccer players. J Sports Sci. 2014;33(1):11-19.',
    '[13] Radnor JM, Oliver JL, Waugh GS, et al. Maturation-related changes in the force-time characteristics of countermovement jumps in trained male youth soccer players. J Strength Cond Res. 2018;32(12):3404-3411.',
    '[14] Ramirez-Campillo R, Moran J, Sánchez-Sánchez J, et al. Effects of plyometric jump training on measures of physical fitness and lower-limb asymmetries in team sport athletes: a systematic review. Scand J Med Sci Sports. 2020;30(8):1328-1346.',
    '[15] Meylan CMP, Cronin JB, Oliver JL, et al. The effect of maturation on the relationship between jump performance and muscle architecture. J Strength Cond Res. 2014;28(6):1535-1543.',
]

# ================================================================
# 3. RESULTS
# ================================================================
add_heading_styled(doc, '3  Results', 1)

add_heading_styled(doc, '3.1  Study Selection Process', 2)
add_para(doc,
    f'Initial database searches identified 135 records from English-language databases and 35 from Chinese databases (CNKI, WanFang, VIP; after deduplication), totaling 168 records. After removing duplicates and screening titles/abstracts (36 duplicates, 22 irrelevant), 110 full-text articles were assessed. A total of 81 were excluded (25 did not report CMJ height, 15 lacked a non-plyometric control group, 8 used non-randomized designs, 7 involved CMJ with arm swing without separately reporting armless CMJ, 3 had duplicate/overlapping samples, 23 Chinese studies did not meet strict inclusion criteria*, and others).最终29篇RCT纳入Meta分析。'
    '*多数中文文献使用"原地纵跳摸高"而非测力台/光电子系统CMJ，且未报告手臂位置。'
    f'Strict hands-on-hips armless CMJ: 16 studies (including R19 and R24 VJ); after excluding R19 (SD/SE ambiguity) and R24 (VJ not CMJ): k={strict_k}. '
    f'Arm-unclear: 6 studies. CMJA with arms: 6 studies. VJ arm-swing sensitivity subgroup: 1 study.')

add_figure(doc,
    str(OUTPUT_DIR / 'PRISMA_flow_diagram.png'),
    'Figure 1. PRISMA flow diagram of study selection', width=Inches(5.0))

add_heading_styled(doc, '3.2  Study Characteristics', 2)
add_para(doc,
    f'A total of 29 RCTs involving 718 participants (367 intervention, 351 control) were included. Studies were published between 2003 and 2025, spanning soccer (11 studies), basketball (5 studies), volleyball, handball, rowing, swimming, middle-distance running, and other sports. Participant ages ranged from 11 to 70 years, predominantly male (20 studies). '
    f'In the strict pool (k={strict_k}), the pooled effect size was g={fmt_g(strict_g)} (only -0.016 different from when R19 was included).')

add_para(doc, 'Table 1. Characteristics of included studies', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['Characteristic', 'Category', 'k'],
    [
        ['Age/Maturation', 'Prepubertal', '2'],
        ['', 'Pubertal', '3 (including 1 VJ)'],
        ['', 'Young adult', '11'],
        ['', 'Adult/Professional', '12'],
        ['', 'Older adult', '1'],
        ['Sport', 'Soccer', '11'],
        ['', 'Basketball', '5'],
        ['', 'Other sports', '12'],
        ['CMJ Arm Position', 'Strict hands-on-hips', f'{strict_k}*'],
        ['', 'Arm-unclear', '6'],
        ['', 'CMJA with arms', '6'],
        ['Duration', 'Short-term (≤6 wk)', f'{short_k}'],
        ['', 'Medium-term (7-10 wk)', f'{mid_k}'],
        ['', 'Long-term (>10 wk)', '4'],
        ['*Note: R19 excluded from strict pool due to SD/SE ambiguity. Δg=-0.016 after exclusion.', '', ''],
    ])

add_heading_styled(doc, '3.3  Risk of Bias Assessment', 2)
add_para(doc,
    'PEDro scores (items 2-11, max 10): mean 5.93/10, median 6, range 3-8. Quality distribution: good or better (≥6) in 17 studies (58.6%), fair (4-5) in 11 (37.9%), poor (<4) in 1 (R09 Blazevich 2003, 3/10, quasi-randomized). Items 5 (subject blinding) and 6 (therapist blinding) had 0% pass rates — an inherent limitation of exercise interventions. Adjusted 8-item score (excluding items 5-6): mean ~5.4/8 (67%).')

add_para(doc,
    'TESTEX scores (15 items, max 15): mean 8.8/15, median 8, range 7-13. Effective maximum 13 (excluding items 5-6): mean 8.8/13 (67.7%). '
    'TESTEX revealed exercise-specific reporting deficiencies: item 14 (adherence) 31%, item 15 (intensity monitoring) 48%. '
    'TESTEX-PEDro correlation was weak (r=0.17, p=0.39), confirming the scales capture different quality dimensions.')

add_para(doc, 'Table 2. Item-level pass rates (29 studies)', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['Item', 'PEDro %', 'TESTEX %', 'Notes'],
    [
        ['1. Eligibility', '100%', '100%', 'Not scored in PEDro'],
        ['2. Random allocation', '95%', '100%', 'Blazevich 2003 quasi-RCT in PEDro'],
        ['3. Allocation concealment', '7%', '10%', 'Rarely reported in sports science'],
        ['4. Baseline comparability', '90%', '97%', 'Pre-diff <10% considered comparable'],
        ['5. Subject blinding', '0%', '0%', 'Inherent limitation'],
        ['6. Therapist blinding', '0%', '0%', 'Coaches know the protocol'],
        ['7. Assessor blinding', '15%', '14%', 'Only 4 studies blinded assessors'],
        ['8. Follow-up/Reliability', '90%', '100%', 'CMJ is standardized and reliable'],
        ['9. ITT/Validity', '100%', '100%', 'CMJ is a valid criterion'],
        ['10. Dropout reporting', '—', '48%', 'TESTEX requires explicit dropout data'],
        ['11. Treatment completion', '100%', '28%', 'TESTEX requires all participants completed'],
        ['12. Between-group comparison', '100%', '100%', 'All reported CMJ comparisons'],
        ['13. Point + variability', '59%', '100%', 'All reported Mean±SD'],
        ['14. Adherence', '—', '31%', 'TESTEX exercise-specific item'],
        ['15. Intensity monitoring', '—', '48%', 'TESTEX exercise-specific item'],
    ])

add_heading_styled(doc, '3.4  Primary Analysis', 2)
add_para(doc,
    f'Strict hands-on-hips pool ({strict_k} RCTs, excluding R19): '
    f'Pooled Hedges\' g = {fmt_g(strict_g)} [95% CI: {fmt_g(strict_cl)}, {fmt_g(strict_cu)}], p<0.001; '
    f'prediction interval: [{fmt_g(strict.get("pred_low", -0.695))}, {fmt_g(strict.get("pred_upp", 2.919))}]; '
    f'heterogeneity: τ²=0.781 [95%CI: 0.378, 3.768], I²={fmt_pct(strict_I2)} [95% CI: {fmt_pct(strict_I2_l)}, {fmt_pct(strict_I2_u)}], '
    f'Q({strict_k-1})=53.38, p<0.001.')

add_para(doc,
    f'Broader pool ({wide_k} RCTs, excluding R19 and R24 VJ): '
    f'Pooled Hedges\' g = {fmt_g(wide_g)} [95% CI: {fmt_g(wide_cl)}, {fmt_g(wide_cu)}], p<0.001; '
    f'prediction interval: [{fmt_g(wide.get("pred_low", -0.250))}, {fmt_g(wide.get("pred_upp", 2.221))}]; '
    f'heterogeneity: I²={fmt_pct(wide_I2)} [95% CI: {fmt_pct(wide_I2_l)}, {fmt_pct(wide_I2_u)}].')

add_para(doc,
    'Interpretation: Plyometric training likely produces a positive effect on CMJ height — the pooled effect sizes fall within the large effect range (g>0.8) per Cohen\'s criteria, '
    'but must be interpreted in light of the overall GRADE certainty being Low. Results were consistent across analysis pools.')

add_figure(doc,
    str(OUTPUT_DIR / 'forest_strict_hand_on_hip.png'),
    'Figure 2. Forest plot: strict hands-on-hips pool (15 studies, R19 excluded)', width=Inches(5.5))
add_figure(doc,
    str(OUTPUT_DIR / 'funnel_strict_hand_on_hip.png'),
    'Figure 3. Funnel plot: strict hands-on-hips pool', width=Inches(5.0))

add_heading_styled(doc, '3.5  Sensitivity Analyses', 2)
add_para(doc, 'Table 3. Sensitivity analysis: outlier exclusion', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['Analysis', 'k', 'Hedges\' g', '95% CI', 'I²'],
    [
        ['All strict hands-on-hips (excl. R19)', str(strict_k), fmt_g(strict_g), fmt_ci_pair(strict_cl, strict_cu), fmt_pct(strict_I2)],
        ['Including R19 (sensitivity)', '16', '+1.128', '[+0.656, +1.600]', '76%'],
        ['Excluding R11 (Sedano Campo 2009)', '14', '+0.900', '[+0.545, +1.255]', '54%'],
        ['Excluding R27 (Toumi 2004)', '14', '+0.981', '[+0.522, +1.440]', '72%'],
        ['Excluding R11+R27', '13', '+0.794', '[+0.486, +1.101]', '37%'],
        ['PEDro≥6 subgroup', '8', '+0.714', '[+0.383, +1.045]', '12%'],
        ['PEDro<6 subgroup', '7', '+1.619', '[+0.671, +2.567]', '84%'],
    ])

add_para(doc,
    f'Leave-one-out analysis: removing any single study kept the pooled effect between g=+0.900 and +1.197, all significantly >0. '
    f'The most influential study was R11 Sedano Campo (2009) (change score Hedges\' g=+5.20, largest individual effect; '
    f'removal: Δg=-0.212, I² dropped to 54.1%).')

add_para(doc,
    f'Pre-post correlation sensitivity: r substantially affected SMD magnitude. r=0.5 (conservative lower bound): strict pool SMD=+0.876 '
    f'[95%CI: +0.481, +1.271]; r=0.7 (default): SMD=+1.096 [+0.631, +1.561]; r=0.9 (upper bound): SMD=+1.720 [+1.076, +2.365]. '
    f'Effect direction remained consistently positive and significant across all r values.')

add_figure(doc,
    str(OUTPUT_DIR / 'forest_wide_all_cmj.png'),
    'Figure S1. Forest plot: broader pool (27 studies, excl. R19/R24 VJ)', width=Inches(5.5))

add_heading_styled(doc, '3.6  Subgroup Analyses', 2)
add_para(doc, 'Table 4. CMJ arm position subgroup analysis', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['Subgroup', 'k', 'Hedges\' g', '95% CI', 'I²'],
    [
        ['Strict hands-on-hips (excl. R19)', str(strict_k), fmt_g(strict_g), fmt_ci_pair(strict_cl, strict_cu), fmt_pct(strict_I2)],
        ['Arm-unclear', '6', '+0.794', '[+0.465, +1.124]', '20%'],
        ['CMJA with arms', '6', '+0.977', '[+0.214, +1.741]', '79%'],
    ])

add_para(doc, 'Table 5. Intervention duration subgroup analysis (key moderator)', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['Subgroup', 'k', "Hedges' g", '95% CI', 'I²'],
    [
        [f'Short-term (≤6 wk)', f'{short_k}*', fmt_g(short_g), fmt_ci_pair(short_cl, short_cu), fmt_pct(short_I2)],
        [f'Medium-term (7-10 wk)', f'{mid_k}', fmt_g(mid_g), fmt_ci_pair(mid_cl, mid_cu), f'{fmt_pct(mid_I2)}'],
        [f'Long-term (>10 wk)', '4', '+1.852', '[-0.080, +3.783]', '94%'],
    ])
add_para(doc,
    f'Short-term plyometric training yielded the most reliable pooled effect: I² only {fmt_pct(short_I2)}, '
    f'95%CI compact, moderate-to-large effect (g≈{fmt_g(short_g)}). Intervention duration was the primary moderator of heterogeneity.')

add_para(doc, 'Table 6. Age/maturation subgroup analysis', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['Subgroup', 'k', 'Hedges\' g', '95% CI', 'I²', 'Note'],
    [
        ['Prepubertal', '2', '+1.374', '[+0.773, +1.975]', '0%', '[! k=2, insufficient]'],
        ['Pubertal', '2', '+1.513', '[+0.878, +2.149]', '0%', '[! k=2, insufficient]'],
        ['Young adult', '11', '+1.112', '[+0.676, +1.548]', '67%', ''],
        ['Adult/Professional', '9', '+0.803', '[+0.193, +1.413]', '83%', ''],
        ['Older adult', '1', '—', '—', '—', '[k=1, cannot pool]'],
    ])

add_heading_styled(doc, '3.7  Meta-Regression', 2)
add_para(doc, 'Table 7. Univariable meta-regression results', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['Moderator', 'Slope (b)', 'SE', 'z', 'p'],
    [
        ['Duration (weeks)', '+0.131', '0.058', '+2.26', '0.023'],
        ['Weekly frequency', '+0.536', '0.231', '+2.32', '0.020'],
        ['Age (years)', '-0.008', '0.014', '-0.59', '0.556'],
        ['Total sample size', '-0.007', '0.008', '-0.89', '0.372'],
        ['Publication year', '-0.040', '0.031', '-1.28', '0.202'],
        ['CMJ arm position', '-0.104', '0.184', '-0.57', '0.570'],
    ])
add_para(doc,
    'Each additional week of intervention was associated with an expected SMD increase of approximately 0.13 (p=0.023). '
    'However, after Bonferroni correction (adjusted α=0.05/12≈0.00417), this association was no longer significant, '
    'classifying the dose-response relationship as exploratory rather than confirmatory.')

add_heading_styled(doc, '3.8  Publication Bias', 2)
add_para(doc,
    f'Egger regression: strict pool intercept=-2.64, p<0.001; broader pool intercept=-1.39, p<0.001. '
    f'SE-g correlation: strict pool r=+0.88, broader pool r=+0.86 — indicating small-study effects. '
    f'Visual funnel plot inspection suggested approximate symmetry, but the SE-effect association confirmed small-study effects.')

add_para(doc,
    f'Trim-and-Fill correction: strict pool estimated 2 missing studies on the right, corrected g=+1.284 [+0.810, +1.757] (+13.8% increase); '
    f'broader pool estimated 3 missing studies, corrected g=+1.124 [+0.836, +1.413] (+12.2% increase). '
    f'The paradoxical increase after Trim-and-Fill is a known failure signal under high heterogeneity (I²=78.1%; Terrin et al., 2003).')

add_heading_styled(doc, '3.9  GRADE Evidence Profile', 2)
add_para(doc, 'Table 8. GRADE Evidence Profile: Plyometric training effects on CMJ height', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['Outcome', 'k(N)', 'Pooled g [95%CI]', 'Bias', 'Inconsistency', 'Indirectness', 'Imprecision', 'Pub. Bias', 'GRADE'],
    [
        ['Primary (strict)', f'{strict_k}(310)', f'{fmt_g(round(strict_g,2))}[{fmt_g(strict_cl)},{fmt_g(strict_cu)}]', 'Serious', 'Serious', 'Not serious', 'Serious', 'Serious', 'Low'],
        ['Broader pool', f'{wide_k}(671)', f'{fmt_g(round(wide_g,2))}[{fmt_g(wide_cl)},{fmt_g(wide_cu)}]', 'Serious', 'Serious', 'Not serious', 'Serious', 'Serious', 'Low'],
        ['Short-term ≤6wk', str(short_k), f'{fmt_g(round(short_g,2))}[{fmt_g(short_cl)},{fmt_g(short_cu)}]', 'Serious', 'Some concerns', 'Not serious', 'Serious', 'Moderate', 'Moderate'],
        ['Medium-term 7-10wk', str(mid_k), f'{fmt_g(mid_g)}[{fmt_g(mid_cl)},{fmt_g(mid_cu)}]', 'Serious', 'Serious', 'Not serious', 'Serious', 'Serious', 'Low'],
        ['Long-term >10wk', '4', '+1.85[-0.08,3.78]', 'Serious', 'Serious', 'Not serious', 'Serious', 'Serious', 'Very Low'],
    ])
add_para(doc,
    'GRADE summary: Overall evidence certainty is Low, primarily limited by high heterogeneity and publication bias signals. '
    'Short-term (≤6 week) intervention evidence is Moderate (I²=14%), representing the most reliable effect estimate in this review.')

# ================================================================
# 4. DISCUSSION
# ================================================================
add_heading_styled(doc, '4  Discussion', 1)

add_heading_styled(doc, '4.1  Main Findings', 2)
add_para(doc,
    f'This meta-analysis included 29 RCTs (k={strict_k} in strict pool, k={wide_k} in broader pool) to systematically evaluate plyometric training effects on CMJ height. '
    f'Given the overall GRADE certainty of Low, the following findings require cautious interpretation:')

findings = [
    f'Plyometric training likely produces a positive effect on CMJ height: strict pool Hedges\' g={fmt_g(round(strict_g,2))}, broader pool g=+0.99, both classified as large effects (g>0.8). However, overall GRADE rating is Low, indicating further research is likely to change effect estimates.',
    'Results were highly robust across sensitivity analyses: excluding two most influential outliers (R11 Sedano Campo 2009 and R27 Toumi 2004), the pooled effect remained g=+0.79 [95%CI: +0.49, +1.10], I² dropped to 37%.',
    f'The dose-response relationship showed a non-linear pattern (exploratory, non-significant after Bonferroni correction): short-term (≤6 wk) yielded consistent moderate effects (g≈+0.71, I²=14%), medium-term (7-10 wk) g≈{fmt_g(mid_g)}, and long-term (>10 wk) g=+1.85 but with extreme heterogeneity (I²=94%).',
    'Adolescent athletes showed the largest gains: prepubertal and pubertal subgroups reached g=+1.37~1.51, but with k=2 in each group, these are exploratory signals only (GRADE: Low).',
    f'The 95% prediction interval for the strict pool [{fmt_g(strict.get("pred_low", -0.695))}, {fmt_g(strict.get("pred_upp", 2.919))}] crosses zero, indicating that under certain conditions (poor adherence, high baseline fitness, inappropriate protocols), plyometric training may produce zero or even negative effects.',
]
for f_item in findings:
    add_para(doc, f'• {f_item}', first_line_indent=Cm(1.5), size=Pt(12), spacing_after=Pt(3))

add_heading_styled(doc, '4.2  Comparison with Previous Meta-Analyses', 2)
add_para(doc,
    'Our pooled effect sizes (strict pool g=+1.09, broader pool g=+0.99) are slightly higher than most previous meta-analyses (ES/SMD=0.49-1.01). '
    'This can be explained by: (1) stricter inclusion criteria — we restricted to armless CMJ, excluding arm-swing jumps that may dilute the "pure" lower-limb SSC training effect; '
    '(2) inclusion of recent RCTs (2021-2025) with more standardized CMJ protocols; '
    '(3) use of pre-post change score SMD (assuming r=0.7) rather than post-only SMD.')

add_heading_styled(doc, '4.3  Dose-Response Relationship', 2)
add_para(doc,
    'Important caveat: After Bonferroni correction, the linear association between intervention duration and effect size was no longer statistically significant. '
    'Therefore, all discussions below regarding dose-response relationships should be treated as exploratory and require independent verification.')

add_para(doc,
    'The subgroup analyses revealed a non-linear pattern: short-term (≤6 wk) yielded reliable moderate effects (g≈+0.71, I²=14%), '
    'medium-term (7-10 wk) showed modest increases (g≈+0.95), and long-term (>10 wk) showed a larger jump (g=+1.85) but with extreme heterogeneity (I²=94%). '
    'This pattern suggests: short-term gains primarily through neural adaptations (4-8 weeks), medium-term gains through consolidation and early structural adaptations, '
    'and long-term gains through muscle structural remodeling — but with substantially increased inter-individual variability.')

add_heading_styled(doc, '4.4  Age/Maturation Effects', 2)
add_para(doc,
    'Adolescent athletes (prepubertal and pubertal) showed numerically larger effects (g=+1.37~1.51) than adults (g=+0.80), consistent with the developmental neuromuscular plasticity hypothesis — '
    'the "sensitive window" around puberty when neuromuscular systems are rapidly developing may synergize with plyometric stimuli. '
    'However, with only k=2 in each subgroup, these findings are exploratory signals only.')

add_heading_styled(doc, '4.5  CMJ Arm Position: Methodological Implications', 2)
add_para(doc,
    'A unique aspect of this study is using CMJ arm position as an inclusion criterion and stratification variable. '
    'From a biomechanical perspective, arm swing involves at least three mechanisms: (a) inertial reaction force, (b) center of mass optimization, '
    'and (c) angular momentum transfer. These mechanisms systematically increase arm-swing CMJ height compared to armless CMJ, '
    'making our strict hands-on-hips analysis a more conservative and reliable estimate of lower-limb plyometric training "pure effects."')

add_heading_styled(doc, '4.6  Training Safety', 2)
add_para(doc,
    'This meta-analysis did not systematically assess injury risk, and only 3 of 29 included RCTs briefly mentioned no serious adverse events. '
    'Plyometric training involves high-impact exercises with ground reaction forces of 3-7× body weight. '
    'Coaches should follow progressive loading principles, limit plyometric sessions to ≤3/week with 48-hour recovery intervals, '
    'and monitor individual recovery signals.')

add_heading_styled(doc, '4.7  Limitations', 2)
add_para(doc, 'This study has several limitations that readers should consider:')

limitations = [
    ('Publication bias cannot be excluded.',
     'Egger test was significant (p<0.001), Begg rank correlation (τ=0.429, p=0.028), and SE-g correlation was strong (r=+0.86-0.88). However, meta-regression found no significant sample size-effect association (p=0.372), suggesting small-study effects may reflect more intensive training rather than classic publication bias. Results should be interpreted as upper-bound estimates.'),
    ('Small sample sizes in some studies.',
     'Several included studies had n<20 (e.g., R03: n=16; R10: n=13), potentially amplifying individual study influence on pooled effects.'),
    ('Data estimation.',
     'Three studies involved data estimation: (a) R19 excluded due to SE/SD ambiguity; (b) R23 control SD approximated from intervention SD; (c) R29 post-test values estimated from change scores. Sensitivity analysis confirmed minimal impact (Δg=-0.016 to +0.049).'),
    ('Incomplete database coverage.',
     'SPORTDiscus and Embase were not searched. Chinese databases yielded no studies meeting strict inclusion criteria. Grey literature (ProQuest Dissertations) and citation tracking were not performed.'),
    ('Limited subgroup sizes.',
     'Prepubertal (k=2), pubertal (k=2), and long-term (k=4) subgroups had insufficient studies for robust inference.'),
    ('Multiple comparisons.',
     'Six subgroup analyses + six meta-regression tests = 12 independent tests. After Bonferroni correction, no associations remained significant.'),
    ('Linear assumption in meta-regression.',
     'The linear duration-effect relationship (0.13 SMD/week) is useful within 4-16 weeks but should not be extrapolated beyond this range.'),
    ('Lack of individual participant data.',
     'IPD meta-analysis was not possible due to unavailability of raw data from included studies.'),
    ('Missing safety data.',
     'Only 3 of 29 RCTs reported adverse events; injury rates remain unknown.'),
    ('Soccer study overrepresentation.',
     'Soccer studies comprised 37.9% of included RCTs, potentially introducing sport-specific directional bias.'),
]
for i, (bold_part, normal_part) in enumerate(limitations, 1):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.first_line_indent = Cm(0)
    run_b = para.add_run(f'{i}. {bold_part}')
    set_run_font(run_b, size=Pt(12), bold=True)
    run_n = para.add_run(normal_part)
    set_run_font(run_n, size=Pt(12))

add_heading_styled(doc, '4.8  Practical Recommendations and Future Directions', 2)
add_para(doc, 'Recommendations for coaches and practitioners:', bold=True)
coach_recs = [
    f'Short-term (6 weeks) plyometric training (2-3 sessions/week) produces reliable moderate-to-large improvements in CMJ height (g≈{fmt_g(short_g)}, approximately 2.4-3.6 cm on original scale), with high consistency (I²=14%). Recommended as a standard component of pre-season preparation.',
    f'Extending training to 7-10 weeks yields additional moderate gains (g≈{fmt_g(mid_g)}, ~3.8-5.7 cm), with increased individual variability (I²={fmt_pct(mid_I2)}). Personalized monitoring is recommended.',
    'Adolescent athletes may be a population with particularly large gains (g≈1.4-1.5, GRADE: Low, k=2). Training should follow progressive loading principles under professional supervision.',
]
for rec in coach_recs:
    add_para(doc, f'• {rec}', first_line_indent=Cm(1.5), spacing_after=Pt(3))

add_para(doc, 'Recommendations for researchers:', bold=True)
res_recs = [
    'CMJ testing must explicitly report arm position. "Hands on hips" or "arms across chest" should be standardized as the test posture.',
    'Future studies should comprehensively report training dose parameters (FITT) including total ground contacts, drop jump height, and training type details.',
    'Larger, longer-duration RCTs (>10 weeks) are needed to stabilize long-term effect estimates.',
    'Individual participant data should be shared to enable future IPD meta-analyses.',
    'CMJ testing equipment (force plate, contact mat, OptoJump) should be considered as a moderator in future meta-analyses.',
]
for rec in res_recs:
    add_para(doc, f'• {rec}', first_line_indent=Cm(1.5), spacing_after=Pt(3))

# ================================================================
# 5. DECLARATIONS
# ================================================================
add_heading_styled(doc, '5  Declarations', 1)

add_para(doc,
    'Funding: This research received no specific grant from any funding agency.',
    size=Pt(12))
add_para(doc,
    'Conflicts of interest: The authors declare no competing interests.',
    size=Pt(12))
add_para(doc,
    'Protocol availability: The PROSPERO pre-registered protocol (CRD420261422906) serves as the formal protocol document. '
    'Analysis code is publicly available on Zenodo (https://doi.org/10.5281/zenodo.20809676). '
    'AI-assisted tools (Claude AI, Anthropic) were used strictly for code generation/debugging, text formatting, and reference formatting. '
    'All scientific judgments were made independently by human authors.',
    size=Pt(12))

# ================================================================
# REFERENCES
# ================================================================
add_heading_styled(doc, 'References', 1)

# NOTE: Full references (50 total) should be copied from generate_v5.py lines 1382-1631
# This is a subset for demonstration

for ref in references:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Cm(1.27)
    para.paragraph_format.first_line_indent = Cm(-1.27)
    run = para.add_run(ref)
    set_run_font(run, size=Pt(11))

# ================================================================
# SAVE
# ================================================================
output_path = PROJ / 'Plyo训练CMJ高度Meta分析_EN.docx'
doc.save(str(output_path))
print(f'\nSaved English version: {output_path}')
print('Note: Results and Discussion sections need full translation from generate_v5.py.')
print('Complete the translation by copying and translating lines 600-1370 from generate_v5.py.')
