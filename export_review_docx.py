# -*- coding: utf-8 -*-
"""Generate peer review Word document. Content is in plain strings, no fancy quotes."""

import json, os, sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

STYLE_FONT = 'Microsoft YaHei'

def set_run_font(run):
    run.font.name = STYLE_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), STYLE_FONT)
    rPr.insert(0, rFonts)

def h2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        set_run_font(run)

def h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        set_run_font(run)

def p(text, bold=False, color=None, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    set_run_font(run)
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return para

def tabled(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); set_run_font(r)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            r = table.rows[ri+1].cells[ci].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9); set_run_font(r)
    doc.add_paragraph()

# ======================================================
# CONTENT DATA - stored in a dict to avoid source-code encoding issues
# ======================================================
C = {}

C['title'] = '模拟同行评审报告'
C['subtitle'] = 'Plyometric 训练对反向纵跳 (CMJ) 高度影响的 Meta 分析'
C['meta'] = '审稿人: Claude (AI 初审) | 建议: Major Revision (大修) | 日期: 2026-06-08'
C['overall'] = (
    '这是一项有明确创新价值的 Meta 分析 -- 首次将 CMJ 手臂位置系统纳入纳入标准和亚组分析, '
    '找到了稳健的时长-剂量线性关系。方法框架采用 Python+R 混合架构, 敏感性分析细致'
    '(三篇估算数据的逐一剔除检验)。但当前稿件存在 5 个重大缺陷需要补充, 补齐后可达投稿水平。'
)

C['m1_title'] = '1. 缺失偏倚风险评价 (最致命)'
C['m1_prob'] = (
    '问题: Methods 完全没有提到对纳入研究的方法学质量评价。PRISMA 2020 强制要求报告每篇纳入研究的偏倚风险 '
    '(Risk of Bias), 体育科学领域常用: PEDro 量表 (RCT 专用, 0-10 分)、Cochrane RoB 2 (六大领域)、'
    'TESTEX (运动训练研究专用, 0-15 分)。这本是「两名研究者独立提取」时就应该同时做的。'
)
C['m1_fix'] = (
    '建议: 尽快用 PEDro 量表补评 29 篇研究 (最快, 每篇 5-10 分钟), 然后在 Methods 中增加一节 '
    '"Risk of Bias Assessment", Results 中报告摘要评分 (平均分/范围/低分项), Discussion 中讨论质量对结论的影响。'
)

C['m2_title'] = '2. 缺失 PROSPERO 注册和 PRISMA Checklist'
C['m2_prob'] = '问题: 系统综述的标准实践是在开展前预注册 (PROSPERO)。没有注册号, 投稿时很多期刊直接退回。'
C['m2_fix'] = (
    '建议: 去 PROSPERO (https://www.crd.york.ac.uk/prospero/) 补注册 (记录号格式: CRD420XXXXXX); '
    '填写 PRISMA 2020 Checklist (27 项核查表), 作为 Supplementary 提交。'
)

C['m3_title'] = '3. 预测区间跨零值被淡化处理'
C['m3_prob'] = (
    '问题: 严格池的 95% 预测区间为 [-0.878, +3.134], 下限为负值。这意味着未来的某次 Plyometric 训练干预, '
    '有非零概率完全无效甚至产生负效应。目前这只在 Results 报告了数字, Discussion 完全没有讨论其意义。'
)
C['m3_fix'] = (
    '建议: 在 Discussion 中增加一段, 坦诚解读: '
    '「预测区间跨零提示, 虽然平均效应大且正向, 但在特定条件组合下 (如训练强度不足、受试者基础水平高、'
    '测试方法不一致等), Plyometric 干预可能不产生有意义的 CMJ 增益。这进一步强化了报告训练剂量参数和'
    '采用标准化测试方法的必要性。」'
)

C['m4_title'] = '4. r = 0.7 假设缺乏文献支撑与敏感性检验'
C['m4_prob'] = (
    '问题: Methods 中写「假设前测-后测相关系数 r = 0.7 (CMJ 典型值)」, 但未提供引用。这个假设直接影响 '
    'change-score SMD 的计算 (影响 SD_change 的大小)。'
)
C['m4_fix'] = (
    '建议: 补文献引用 (如 Claudino et al. 2017 报告 CMJ 重测信度 ICC > 0.9; 或参考 Cuijpers et al. 2017 '
    'Res Syn Meth 关于 pre-post correlation 的敏感性分析建议)。同时在敏感性分析中额外检验 r = 0.5 和 '
    'r = 0.9 两种假设, 看合并效应量是否对 r 的选择敏感。'
)

C['m5_title'] = '5. 亚组 k=2 时做了过度结论'
C['m5_prob'] = (
    '问题: 青春期前 (k=2)、青春期 (k=2) 亚组的 I2 = 0%, 文中据此说「完全同质」。但 k=2 时 I2 = 0% '
    '是统计上必然的 (只有两个效应量, I2 永远接近零), 不能作为「一致性高」的证据。另外, Plyo+力量混合组'
    '也只有 k=2, 文中报告 g=+1.97 但没有 CI 且未标注 k=2 的推断局限性。'
)
C['m5_fix'] = (
    '建议: 明确标注 k=2 的亚组为 insufficient for subgroup inference (k < 5 不稳健); '
    '将 Plyo+力量 (k=2) 的 g 值加上 highly tentative 限定; 在亚组分析表中增加一列标记 k < 5 的组。'
)

# Moderate issues
C['mod'] = [
    ('6. 宽版池 k=28 vs. 全文 k=29 的数字链条有裂缝',
     'PRISMA 流程: 135 -> 剔除 44 -> 全文评估 91 -> 排除 60 -> 纳入 29。但「宽版全池」写的是 28。'
     '需要在一个地方明确解释第 29 篇去向 (VJ 敏感性亚组)。建议在 Results 开头加一句: '
     '「29 篇纳入研究中, 28 篇报告了 CMJ 数据 (纳入宽版池), 1 篇仅报告 VJ 数据 (进入 VJ 敏感性亚组)。」'),
    ('7. 「总样本量 ≈ 720+」不精确',
     'Meta 分析的标准实践是报告精确的累计样本量 (total N = sum of all n_int + n_ctrl), 而非近似值。'
     '建议从数据表中精确求和, 替换所有 ≈ 符号。'),
    ('8. 缺失文献检索的精确日期和每条数据库的命中数',
     'Methods 写「检索时限为建库至 2026 年 5 月」但未给出精确检索日期。'
     '正文应呼应 PRISMA 流程图中的数据库命中数。建议补精确日期 + 正文呼应数据库分布。'),
    ('9. 漏斗图对称性论证有逻辑问题',
     'Results 写「合并效应两侧研究分布大致对称 (8 篇高于 / 8 篇低于)」作为反驳 Egger 检验的证据。'
     '漏斗图对称是发表偏倚的必要不充分条件, 且 n<10 时目视检查本身不可靠。不应以此来反驳 Egger 检验。'
     '建议调整措辞, 将两者并列作为不同维度的证据, 承认不确定性。'),
]

C['minor'] = [
    '非盲法筛选的可靠性: 需报告两位研究者姓名首字母、筛选一致性 (Cohen\'s kappa)、分歧解决方式。',
    '缺失补充检索记录: 追溯参考文献列表的额外检出数未报告。',
    'Discussion 引用的 Ramaswamy 2025 和 Salazar-Martinez 2024 不在 References 中, 需补全。',
    '缺失 Abstract、关键词、作者信息、基金来源、利益冲突声明 -- 投稿必需项。',
]

C['score_headers'] = ['评审维度', '评分', '说明']
C['score_rows'] = [
    ['研究问题重要性', '4/5', '体育科学领域, 有一定引用前景'],
    ['方法学严谨性', '3/5', '缺 RoB + 注册 + Checklist'],
    ['创新性', '4/5', 'CMJ 臂位分层是真实创新点'],
    ['分析深度', '4/5', '6 亚组 + Meta 回归 + 多重敏感性'],
    ['讨论合理性', '4/5', '与既往研究比较部分写得很好'],
    ['可复现性', '3/5', '缺 kappa、缺精确检索日期'],
    ['写作质量', '4/5', '结构清晰, 逻辑流畅'],
]

# ======================================================
# BUILD DOCUMENT
# ======================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = STYLE_FONT
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), STYLE_FONT)
rPr.insert(0, rFonts)

# Title page
p(C['title'], bold=True, size=Pt(22))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p(C['subtitle'], bold=True, size=Pt(13))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p(C['meta'], size=Pt(10), color=RGBColor(0x66, 0x66, 0x66))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 1
h2('一、总体评价')
p(C['overall'])

# Section 2
h2('二、重大缺陷 (Must Fix)')

for key in ['m1','m2','m3','m4','m5']:
    h3(C[f'{key}_title'])
    p(C[f'{key}_prob'])
    p(C[f'{key}_fix'], color=RGBColor(0x1F, 0x6F, 0x43))

# Section 3
h2('三、中度问题 (Should Fix)')
for title, content in C['mod']:
    h3(title)
    p(content)

# Section 4
h2('四、次要问题 (Nice to Fix)')
for i, issue in enumerate(C['minor'], start=10):
    p(f'{i}. {issue}')

# Section 5
h2('五、评分总结')
tabled(C['score_headers'], C['score_rows'])

# Section 6
h2('六、优先行动清单')
p('本周可完成 (约 2-3 小时):', bold=True)
for t in [
    '1. PEDro 评分 29 篇 -> 加到 Methods + Results',
    '2. PROSPERO 补注册',
    '3. 精确 N 求和 (不用 ~)',
    '4. r=0.7 补文献引用 + 补充 r=0.5/0.9 敏感性分析',
]:
    p(f'  □ {t}')

p('下周可完成 (约 1-2 小时):', bold=True)
for t in [
    '5. 填写 PRISMA 2020 Checklist (27 项)',
    '6. 撰写 Abstract 初稿 (250-300 字)',
    '7. 补全缺失文献 + 修数字和格式细节',
]:
    p(f'  □ {t}')

p('建议从 PEDro 评分开始 -- 这是审稿人最可能直接拒稿的死穴。',
  bold=True, color=RGBColor(0xC0, 0x39, 0x2B))

# Save
output_path = os.path.join(os.environ.get('USERPROFILE', ''), 'Desktop', '模拟同行评审_Plyo_CMJ_Meta分析.docx')
# Fallback
if not os.path.exists(os.path.dirname(output_path)):
    output_path = 'D:/桌面/模拟同行评审_Plyo_CMJ_Meta分析.docx'

doc.save(output_path)
print(f'Saved: {output_path}')
