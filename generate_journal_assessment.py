# -*- coding: utf-8 -*-
"""
Generate journal acceptance probability assessment Word report.
Based on multi-agent workflow results (10 journal assessors + 1 synthesis agent).
"""
import sys, io, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

STYLE_FONT = 'Times New Roman'
CJK_FONT = '宋体'
HEADING_FONT = '黑体'

def set_run_font(run, size=None, bold=False, italic=False, font_name=None, cjk=None):
    font = font_name or STYLE_FONT
    run.font.name = font
    if cjk is None:
        cjk = CJK_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), cjk)
    rPr.insert(0, rFonts)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, cjk=HEADING_FONT)
    return h

def add_para(doc, text, bold=False, italic=False, size=Pt(11), align=None,
             first_line_indent=Cm(0.74), spacing_after=Pt(6),
             spacing_before=Pt(0), font_name=None):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = spacing_after
    para.paragraph_format.space_before = spacing_before
    para.paragraph_format.first_line_indent = first_line_indent
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True, italic=italic, font_name=font_name)
        else:
            run = para.add_run(part)
            set_run_font(run, size=size, bold=bold, italic=italic, font_name=font_name)
    return para

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=Pt(9), bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=Pt(9))
    doc.add_paragraph()
    return table

def prob_color(p):
    """Return RGB color based on probability: red→orange→yellow→green"""
    try:
        p = float(p)
    except:
        return RGBColor(128, 128, 128)
    if p >= 0.7:
        return RGBColor(0, 128, 0)      # green
    elif p >= 0.5:
        return RGBColor(170, 170, 0)    # yellow-green
    elif p >= 0.35:
        return RGBColor(200, 130, 0)    # orange
    elif p >= 0.2:
        return RGBColor(200, 80, 0)     # deep orange
    else:
        return RGBColor(200, 0, 0)      # red

def add_probability_bar(doc, label, probability, detail=""):
    """Add a visual probability bar"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.first_line_indent = None
    run = para.add_run(f"{label}: ")
    set_run_font(run, size=Pt(10), bold=True)
    pct_str = f"{probability}%" if isinstance(probability, (int, float)) else str(probability)
    run2 = para.add_run(pct_str)
    set_run_font(run2, size=Pt(10), bold=True, cjk=CJK_FONT)
    if detail:
        run3 = para.add_run(f"  — {detail}")
        set_run_font(run3, size=Pt(9))
    return para

# ================================================================
# Build document
# ================================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = STYLE_FONT
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), CJK_FONT)
rPr.insert(0, rFonts)

# ---- TITLE PAGE ----
add_para(doc, '', size=Pt(12))
add_para(doc,
    'Plyometric训练CMJ高度Meta分析\n中英文期刊投稿接受概率评估报告',
    bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=None, spacing_after=Pt(16))
add_para(doc,
    '基于10个多Agent审稿人模拟的综合性投稿策略分析',
    bold=False, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=None, spacing_after=Pt(24))

add_para(doc, '评估方法：多Agent并行评审', bold=True, size=Pt(12),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '10个独立Agent，每个模拟20年经验的资深审稿人/副主编视角，\n分别评估5个英文SCI/SSCI期刊和5个中文CSSCI期刊',
         size=Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '评估日期：2026年6月20日', size=Pt(10),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '论文版本：v3（第7轮修改后，评分9.0/10）', size=Pt(10),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

doc.add_page_break()

# ================================================================
# 1. Executive Summary
# ================================================================
add_heading_styled(doc, '1  执行摘要', 1)

add_para(doc,
    '本报告通过多Agent模拟评审系统，对论文《快速伸缩复合训练对反向纵跳高度影响的系统综述与Meta分析'
    '——基于手臂位置分层的效应量估计与剂量-反应关系》在10个中英文期刊的投稿接受概率进行了独立、'
    '多维度的评估。每个Agent模拟了该期刊方向20年经验的资深审稿人/副主编视角，提供了接受概率、'
    'Desk Reject概率、大修/小修概率、核心优势与劣势、修改建议等详细输出。')

add_para(doc, '**论文核心特征：**', bold=True, size=Pt(11))
add_para(doc,
    '• 研究设计：系统综述与Meta分析，PROSPERO预注册(CRD420261422906)，29篇RCT，718名受试者\n'
    '• 核心创新：首次以CMJ手臂位置作为纳入标准，严格手叉腰池k=15为主分析\n'
    '• 主要结果：Hedges\' g = +1.112 [95%CI: +0.599, +1.626], I² = 78.1% (大效应但高异质性)\n'
    '• 最可靠发现：短期≤6周训练 g = +0.710, I² = 13.8% (中等效应、低异质性)\n'
    '• 证据确定性：GRADE整体为Low（降级理由：偏倚风险+不一致性+不精确性+发表偏倚）\n'
    '• 发表偏倚：Egger p<0.001, SE-g r=+0.88, Trim-and-Fill校正后效应量反增+13.8%\n'
    '• 论文状态：7轮多Agent审查修改完成，累计修复113+项问题，终审评分9.0/10')

add_para(doc, '**综合评估结论：**', bold=True, size=Pt(11))
add_para(doc,
    '该论文方法学质量扎实（开放科学实践、敏感性分析框架、GRADE评估均达到国际SCI期刊上游水平），'
    '但面临两大核心障碍：(1) GRADE Low与"大效应"叙事之间的张力——在所有期刊中审稿人都会关注；'
    '(2) 发表偏倚信号极强(Egger p<0.001, SE-g r=+0.88)——这是运动科学Meta分析中罕见的高水平，'
    '审稿人可能质疑效应量估计的可信度。论文最合适的目标是JSCR（英文）和北京体育大学学报（中文）。'
    '若完成各Agent建议的关键修改，英文期刊接受概率可提升至50-60%，中文期刊可提升至65-75%。')

doc.add_page_break()

# ================================================================
# 2. Overall Probability Comparison
# ================================================================
add_heading_styled(doc, '2  接受概率总览', 1)

add_para(doc, '以下表格汇总了10个期刊的评估结果。接受概率为各Agent基于论文当前状态（含7轮修改后）'
         '的独立评估，并考虑了审核流程各阶段的概率分布。', size=Pt(10))

add_heading_styled(doc, '2.1  英文期刊（SCI/SSCI）', 2)

english_headers = ['期刊名称', 'IF', '层次', '接受概率', 'Desk Reject', '大修', '小修', '直接接受', '推荐度']
english_rows = [
    ['J Strength Cond Res (JSCR)', '3.0-4.0', 'Q1/Q2', '≈47%', '≈15%', '≈50%', '≈28%', '≈7%', '⭐⭐⭐⭐⭐ 首推'],
    ['J Sports Sci (JSS)', '2.5-3.5', 'Q2', '≈22%', '≈30%', '≈45%', '≈18%', '≈2%', '⭐⭐ 不推荐当前形态'],
    ['Biol Sport', '4.0-5.0', 'Q1', '≈35%', '≈20%', '≈50%', '≈25%', '≈5%', '⭐⭐⭐ 可行但有条件'],
    ['Sports Med', '9.0-10.0', 'Q1顶刊', '≈5%', '≈60%', '≈30%', '≈8%', '<1%', '⭐ 极低概率'],
    ['Res Q Exerc Sport', '2.0-3.0', 'Q3', '≈22%', '≈35%', '≈40%', '≈18%', '≈3%', '⭐⭐ 机构壁垒高'],
]
add_table(doc, english_headers, english_rows)

add_heading_styled(doc, '2.2  中文期刊（CSSCI）', 2)

chinese_headers = ['期刊名称', '级别', 'CNKI IF', '接受概率', 'Desk Reject', '大修', '小修', '直接接受', '推荐度']
chinese_rows = [
    ['体育科学', 'CSSCI顶级', '2.0-3.0', '≈30%', '≈35%', '≈45%', '≈15%', '≈3%', '⭐⭐⭐ 有挑战'],
    ['中国体育科技', 'CSSCI', '1.5-2.5', '≈35%', '≈25%', '≈45%', '≈25%', '≈5%', '⭐⭐⭐ 可行'],
    ['北京体育大学学报', 'CSSCI', '1.5-2.5', '≈55%', '≈15%', '≈40%', '≈35%', '≈10%', '⭐⭐⭐⭐⭐ 中文首推'],
    ['体育学刊', 'CSSCI', '1.0-2.0', '≈65%', '≈10%', '≈35%', '≈40%', '≈15%', '⭐⭐⭐⭐ 高概率'],
    ['成都体育学院学报', 'CSSCI扩展', '1.0-1.5', '≈62%', '≈12%', '≈35%', '≈40%', '≈13%', '⭐⭐⭐⭐ 高概率'],
]
add_table(doc, chinese_headers, chinese_rows)

# ================================================================
# 3. Detailed Journal Assessments
# ================================================================
add_heading_styled(doc, '3  英文期刊详细评估', 1)

# ---- JSCR ----
add_heading_styled(doc, '3.1  Journal of Strength and Conditioning Research (JSCR) — 首推', 2)
add_para(doc, '**综合接受概率：≈47%**（最可能结果：大修后接受）', bold=True, size=Pt(11))
add_para(doc, '**核心优势匹配：**', bold=True)
add_para(doc,
    '1. 15/29篇纳入RCT来自JSCR，机构契合度极高——审稿人来自同一学术社群\n'
    '2. 选题（Plyometric训练+CMJ）是该刊最核心的话题领域，编辑和审稿人均为S&C领域专家\n'
    '3. 实践建议（教练可直接使用的短期训练效应）与NSCA读者群需求高度匹配\n'
    '4. PROSPERO+Zenodo开放数据在JSCR发表标准中是加分项')
add_para(doc, '**关键风险：**', bold=True)
add_para(doc,
    '1. 发表偏倚信号(Egger p<0.001, SE-g r=+0.88)在运动科学Meta分析中属极端水平，审稿人可能要求PET-PEESE或选择模型\n'
    '2. 严格池k=15 < Sterne等(2011)推荐的k≥20阈值\n'
    '3. g=+1.11 vs g=+0.79(离群值移除后)的张力：哪个是"真实"效应？审稿人会追问\n'
    '4. 当前稿件为中文，必须翻译为英文——JSCR不接受中文投稿\n'
    '5. GRADE Low与"Large Effect"表述的张力需要全文语言一致性处理')
add_para(doc, '**投稿前必须完成的修改：**', bold=True)
add_para(doc,
    'P0：(1)全文英文化；(2)新增PET-PEESE或选择模型发表偏倚分析；(3)解决g=+1.11 vs +0.79的核心张力——建议将g=+0.79(I²=37%)设为保守主估计\n'
    'P1：(4)应用多重比较校正(Benjamini-Hochberg FDR)；(5)进行Ramirez-Campillo课题组排除后的敏感性分析\n'
    'P2：(6)补充PRISMA Checklist作为补充材料；(7)强化"异质性悖论"（严格池I²>宽版池I²）的讨论')

# ---- JSS ----
add_heading_styled(doc, '3.2  Journal of Sports Sciences (JSS) — 不推荐当前形态', 2)
add_para(doc, '**综合接受概率：≈20-25%**（高风险：大概率收到Major Revision后拒稿）', bold=True, size=Pt(11))
add_para(doc, '**核心问题：** JSS审稿人对GRADE Low + 高I² + 强发表偏倚的"红旗组合"容忍度显著低于JSCR。JSS发表的'
         'Meta分析典型特征为：效应量0.3-0.8, I²<50%, GRADE Moderate-High。本论文g=+1.112, I²=78.1%, '
         'GRADE Low, Egger p<0.001——这四项在JSS审稿人眼中构成"信号级联"，审稿人会质疑：'
         '"如果证据质量如此低，为什么值得发表？"')
add_para(doc, '**Agent评估的关键判断：** "我不建议以当前形态投稿JSS。如果投稿，最可能的结果是收到一封措辞礼貌'
         '但实质否定的拒稿信。更优策略是转投JSCR。"', italic=True)

# ---- Biol Sport ----
add_heading_styled(doc, '3.3  Biology of Sport (Biol Sport) — 有条件可行', 2)
add_para(doc, '**综合接受概率：≈35%**（有条件的大修）', bold=True, size=Pt(11))
add_para(doc, '**关键优势：** 开放科学实践(PROSPERO+Zenodo+GitHub)在Biol Sport是实质性加分项，该刊对方法学创新'
         '（手臂位置分层）的评价高于对证据确定性的苛求。')
add_para(doc, '**关键风险：** 发表偏倚信号强度和Trim-and-Fill反增的解释需要更保守的偏倚校正方法（选择模型/'
         '可信度极限分析），Agent指出当前"向右填充=无偏倚"的解释在统计方法学审稿人面前站不住脚。')

# ---- Sports Med ----
add_heading_styled(doc, '3.4  Sports Medicine (Sports Med) — 极低概率', 2)
add_para(doc, '**综合接受概率：≈3-8%**（大概率Desk Reject）', bold=True, size=Pt(11))
add_para(doc, '**核心障碍：** Sports Med发表权威性综述以指导临床/实践决策。GRADE Low + PI跨零的Meta分析'
         '无法满足"提供权威性指导"的期刊使命。Oliver等(2024)同类Meta分析纳入34篇RCT, n=1396，'
         '远大于本研究的k=15。作者团队为研究生+导师（非知名专家），Sports Med对非邀稿的早期研究者'
         '投稿门槛更高。该Agent建议更实际的目标为Scand J Med Sci Sports (IF~4.5)或Sports Med-Open (IF~5)。')

# ---- RQES ----
add_heading_styled(doc, '3.5  Research Quarterly for Exercise and Sport (RQES) — 机构壁垒', 2)
add_para(doc, '**综合接受概率：≈22%**（大概率拒稿）', bold=True, size=Pt(11))
add_para(doc, '**核心障碍：** RQES明确声明综述类文章"accepted from leading research groups and scientists"。'
         '西南大学体育学院在该刊无发表记录，作者团队无国际知名Meta分析师。方法学质量虽好，'
         '但RQES对作者机构声誉的权重远高于其他期刊。')

# ---- Chinese Journals ----
add_heading_styled(doc, '4  中文期刊详细评估', 1)

# ---- 体育科学 ----
add_heading_styled(doc, '4.1  体育科学 — 可投但有挑战', 2)
add_para(doc, '**综合接受概率：≈30%**', bold=True, size=Pt(11))
add_para(doc, '**优势：** PROSPERO+Zenodo在中文体育学刊极为罕见，方法学创新明确。')
add_para(doc, '**障碍：** (1)作者机构(西南大学)非体育学第一梯队（该刊近年第一作者以北体大/上体大/体育总局科研所为主）；'
         '(2)纯Meta分析在该刊发表难度高（偏向奥运科技攻关/新技术方法/大规模实验）；'
         '(3)缺少基金项目支持；(4)东亚/中国人群RCT极少。')
add_para(doc, '**Agent建议：** 考虑添加北体大/上体大教授级合作者作为共同通讯作者——这是国内体育核心期刊'
         '常见且有效的策略，可大幅改善"新人"感知。')

# ---- 北京体育大学学报 ----
add_heading_styled(doc, '4.2  北京体育大学学报 — 中文期刊首推', 2)
add_para(doc, '**综合接受概率：≈55%（有条件：≈60%→修改后70-75%）**', bold=True, size=Pt(11))
add_para(doc, '**Agent核心判断：** "方法学骨架极好——这是通过外审的最大倚仗。当前版本对北体大学报的适配度约60%，'
         '若能完成CNKI中文文献补充和叙事本土化，接受率可提升至70-75%。"')
add_para(doc, '**必备修改（P0）：** (1)补充CNKI/万方/维普检索；(2)Introduction/Discussion做中文化叙事重构'
         '——对接中国体育政策（如奥运争光计划、体育强国建设纲要）；(3)统一中文术语（"快速伸缩复合训练"）')
add_para(doc, '**关键风险：** 约30-40%为北体大校内稿件，外校投稿需更强学术质量或知名学者推荐。')

# ---- 体育学刊 ----
add_heading_styled(doc, '4.3  体育学刊 — 高概率', 2)
add_para(doc, '**综合接受概率：≈65%**（中高概率，但需体裁适配）', bold=True, size=Pt(11))
add_para(doc, '**Agent核心判断：** "论文方法学质量已远超体育学刊常规稿件平均水平，但该刊极少发表严格意义上'
         '的Meta分析/系统综述。投稿存在over-qualified但在发表偏好上不匹配的张力。"')
add_para(doc, '**关键策略：** 将论文定位从"系统综述与Meta分析"调整为"系统综述（含定量合成）"，'
         '在摘要中强调系统综述为主、Meta分析为补充。补充学校体育维度的讨论（体育课体能训练模块嵌入、'
         '义务教育体育与健康课程标准）。')

# ---- 成都体育学院学报 ----
add_heading_styled(doc, '4.4  成都体育学院学报 — 高概率', 2)
add_para(doc, '**综合接受概率：≈62%**（大概率接受，中等修改）', bold=True, size=Pt(11))
add_para(doc, '**Agent核心判断：** "方法学质量显著高于CSSCI扩展版期刊中位水平。PROSPERO注册、GRADE评估、'
         '多层次敏感性分析在投稿池中极为突出。"')

# ---- 中国体育科技 ----
add_heading_styled(doc, '4.5  中国体育科技 — 可行', 2)
add_para(doc, '**综合接受概率：≈35%**', bold=True, size=Pt(11))
add_para(doc, '**核心障碍：** 该刊偏向发表原创实证研究（实验/调查/测试），纯Meta分析发表占比极低(<5%)。'
         'Agent建议将论文定位为"方法学基准示范"而非"新发现报告"，强调方法论的创新性和可推广性。')

doc.add_page_break()

# ================================================================
# 5. Cross-Journal Analysis
# ================================================================
add_heading_styled(doc, '5  跨期刊综合分析', 1)

add_heading_styled(doc, '5.1  论文在各期刊中的共通优势', 2)
add_para(doc,
    '1. **PROSPERO预注册 + Zenodo公开代码数据**：在几乎所有评估期刊中均被认定为实质性加分项，'
    '在国内中文期刊中更是极为罕见的前沿实践。\n'
    '2. **方法学创新（手臂位置分层）**：首次以CMJ手臂位置作为纳入标准，解决上肢摆动(10-15%跳高)'
    '的测量混杂——这一创新在所有10个Agent评估中均被认定为真实的增量贡献。\n'
    '3. **敏感性分析框架完整**：Leave-one-out + 离群值移除 + r=0.5/0.7/0.9 + R19数据质量敏感性'
    ' + GRADE五域评估 + 三重发表偏倚检验——覆盖了审稿人可能质疑的几乎所有角度。\n'
    '4. **自我批判诚实**：GRADE Low的坦诚报告、Egger/SE-g矛盾的公开讨论、Bonferroni校正后不显著'
    '的明确声明、PI跨零的实践解读——这种诚实态度被多个Agent评为"在投稿中少见且会赢得审稿人尊重"。\n'
    '5. **短期训练的可靠证据**：短期≤6周亚组I²=13.8%是最可靠的发现，所有Agent均认为这对教练/实践者'
    '有直接价值——构成了论文的"安全网"发现。')

add_heading_styled(doc, '5.2  论文在各期刊中的共通劣势', 2)
add_para(doc,
    '1. **发表偏倚信号极端强烈**：Egger p<0.001, SE-g r=+0.88是运动科学Meta分析中罕见的高水平。'
    'Trim-and-Fill校正后效应量反而增加(+13.8%)——这一悖论被多个Agent认定为"红线"而非"解脱"。'
    '5个英文期刊Agent中的4个明确要求补充PET-PEESE或选择模型分析。\n'
    '2. **GRADE Low与论文叙事张力**：多个Agent指出，摘要和封面信中的"大效应"表述与正文GRADE Low'
    '评级在语言上不一致。需要全文使用"may/might/low certainty"措辞替代确定性语言。\n'
    '3. **多重比较问题削弱核心发现**：12个亚组/回归检验中，Bonferroni校正后仅0个显著。'
    '论文的第二大科学卖点（剂量-反应）在统计学上仅属探索性信号。\n'
    '4. **严格池I²(78.1%)反而高于宽版池(67.7%)**：这一"异质性悖论"削弱了"更严格纳入→更干净估计"'
    '的核心叙事。多个Agent要求提供非线性格点回归或更深入的机制讨论。\n'
    '5. **k=15低于漏斗图可靠性阈值(k≥20)**：作者虽已标注这一局限，但发表偏倚评估的可信度因此打折扣。\n'
    '6. **足球/单一课题组主导**：37.9%为足球运动员，Ramirez-Campillo课题组贡献~7/29篇——'
    '生态效度和课题组效应是审稿人会关注的点。')

add_heading_styled(doc, '5.3  投稿排名总榜', 2)

rank_headers = ['排名', '期刊', '语言', '接受概率', 'IF/级别', 'Agent评级']
rank_rows = [
    ['1', '体育学刊', '中文', '≈65%', 'CSSCI', '中高概率'],
    ['2', '成都体育学院学报', '中文', '≈62%', 'CSSCI扩展', '高概率'],
    ['3', '北京体育大学学报', '中文', '≈55%', 'CSSCI', '有条件高概率'],
    ['4', 'J Strength Cond Res', '英文', '≈47%', 'IF 3.0-4.0', '有条件中等概率'],
    ['5', 'Biol Sport', '英文', '≈35%', 'IF 4.0-5.0', '有条件中等概率'],
    ['6', '中国体育科技', '中文', '≈35%', 'CSSCI', '可行但有限制'],
    ['7', '体育科学', '中文', '≈30%', 'CSSCI顶级', '有挑战'],
    ['8', 'J Sports Sci', '英文', '≈22%', 'IF 2.5-3.5', '不推荐当前形态'],
    ['9', 'Res Q Exerc Sport', '英文', '≈22%', 'IF 2.0-3.0', '机构壁垒'],
    ['10', 'Sports Med', '英文', '≈5%', 'IF 9.0-10.0', '极低概率'],
]
add_table(doc, rank_headers, rank_rows)

add_para(doc, '**注意：** 中文期刊和英文期刊的概率不宜直接比较——两者评审体系、文化、标准存在系统性差异。'
         '中文期刊的高概率部分源于CSSCI扩展版/普通CSSCI的对标门槛较低，而英文SCI期刊的审稿标准'
         '在统计方法学和发表偏倚信号解读上更为严苛。'
         '论文的方法学质量更接近SCI期刊标准，选择英文投稿获取的国际认可度和引用潜力更高，'
         '但接受概率较低；中文投稿接受概率较高但学术影响力有限。', italic=True, size=Pt(10))

doc.add_page_break()

# ================================================================
# 6. Recommended Submission Strategy
# ================================================================
add_heading_styled(doc, '6  推荐投稿策略', 1)

add_heading_styled(doc, '6.1  三种投稿路线', 2)

add_para(doc, '**路线A：英文优先（追求国际认可度）**', bold=True, size=Pt(12))
add_para(doc,
    'Step 1 — JSCR首投：完成英文化+PET-PEESE+离群值移除后重新框定主结果(g=+0.79→+1.11范围)。'
    '预估接受概率47%。若被拒：\n'
    'Step 2 — Biol Sport次投：强化发表偏倚多重校正+开放科学叙事。预估35%。\n'
    'Step 3 — J Sports Sci或Scand J Med Sci Sports保底。预估25-30%。\n'
    '时间线：英文投稿周期约6-12个月（含修改）。')

add_para(doc, '**路线B：中文优先（追求高接受率）**', bold=True, size=Pt(12))
add_para(doc,
    'Step 1 — 北京体育大学学报：完成CNKI检索补充+叙事本土化+添加资深合作者。预估60-75%。\n'
    'Step 2 — 体育学刊：体裁适配（调整为系统综述+定量合成框架）。预估65%。\n'
    'Step 3 — 成都体育学院学报保底：预估62%。\n'
    '时间线：中文投稿周期约3-6个月。')

add_para(doc, '**路线C：双线并行（最优但工作量最大）**', bold=True, size=Pt(12))
add_para(doc,
    '同时准备中英文版本，但投稿时需注意一稿不可多投原则。建议策略：\n'
    '先投英文JSCR（最匹配的SCI期刊）→ 若被拒→转投中文北京体育大学学报或体育学刊。\n'
    '该路线最大化论文的"使用价值"：英文版积累国际引用，中文版确保发表。')

add_heading_styled(doc, '6.2  通用投稿前改进（所有期刊适用）', 2)

add_para(doc, '**P0级（不修改将显著降低接受概率）：**', bold=True)
add_para(doc,
    '1. 补充至少一种保守发表偏倚校正方法（PET-PEESE或选择模型），不能仅靠Trim-and-Fill反增来论证无偏倚\n'
    '2. 解决g=+1.11 vs g=+0.79的核心张力——建议将离群值移除后的估计作为保守主估计，或至少对两者之间的'
    '不确定性进行更系统的讨论\n'
    '3. 全文语言一致性：GRADE Low对应"may/might"措辞，而非"produces/shows"\n'
    '4. 将剂量-反应关系明确降级为"探索性信号"——Bonferroni校正后不显著这一事实需要在摘要中体现\n'
    '5. 若投英文期刊：全文翻译为英文（建议使用专业学术翻译服务）\n'
    '6. 若投中文期刊：补充CNKI/万方/维普检索，增加中国体育场景讨论')

add_para(doc, '**P1级（大概率引发审稿人批评，建议修复）：**', bold=True)
add_para(doc,
    '7. 应用Benjamini-Hochberg FDR多重比较校正，报告调整后p值\n'
    '8. 新增Ramirez-Campillo课题组排除后的敏感性分析\n'
    '9. 新增PEDro≥6研究的敏感性分析\n'
    '10. 补充SPORTDiscus/Embase检索或提供不检索的有力理由\n'
    '11. 重新框定k=2亚组（青少年）的GRADE评级为Low而非Moderate\n'
    '12. 补充PRISMA 2020 Checklist作为补充材料')

add_para(doc, '**P2级（提升论文竞争力）：**', bold=True)
add_para(doc,
    '13. 添加非线性格点回归（restricted cubic splines）检验剂量-反应非线性\n'
    '14. 扩充pre-post r=0.7选取依据的文献论证\n'
    '15. 足球亚组敏感性分析\n'
    '16. 考虑邀请知名合作者（特别是中文期刊投稿时）')

doc.add_page_break()

# ================================================================
# 7. Final Recommendation
# ================================================================
add_heading_styled(doc, '7  最终推荐', 1)

add_para(doc, '**综合10个Agent的独立评估，本研究团队给出以下最终推荐：**', bold=True, size=Pt(12))

add_para(doc, '')
add_para(doc, '**首选投稿目标：Journal of Strength and Conditioning Research (JSCR)**',
         bold=True, size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '接受概率 ≈47%（完成P0修改后 → 55-60%）', bold=True, size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '')

add_para(doc, '**次选（中文保底）：北京体育大学学报**', bold=True, size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '接受概率 ≈55%（完成本土化修改后 → 70-75%）', bold=True, size=Pt(10),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

add_para(doc, '')
add_para(doc, '**理由：**', bold=True)
add_para(doc,
    '1. JSCR是Plyometric训练CMJ研究的"事实中心"——29篇纳入RCT中15篇(52%)发表于该刊，'
    '审稿人来自同一学术社群，最可能认可手臂位置分层的创新价值。\n'
    '2. JSCR对高异质性和大效应量的容忍度在SCI体育学期刊中相对最高——I²=78%在该刊发表的'
    'Plyometric训练Meta分析中是常见水平。\n'
    '3. 论文的方法学质量（PROSPERO+GRADE+多重敏感性分析+开放数据）已显著高于JSCR平均投稿水平，'
    '有望获得"方法学严谨"的评价而非"证据不足"的拒稿理由。\n'
    '4. 中文保底路线（北体大学报）接受概率高，可确保论文最终能够发表。\n'
    '5. 若JSCR因发表偏倚信号过强而拒稿，Biol Sport(35%)和体育学刊(65%)均为合理的备选。')

add_para(doc, '')
add_para(doc, '**本评估的关键不确定性提示：**', bold=True, size=Pt(10))
add_para(doc,
    '1. Agent模拟存在内在局限——Agent虽基于期刊公开信息和20年审稿经验的角色扮演进行评估，'
    '但无法完全模拟真实审稿人的判断（特别是对非统计因素的直觉反应和学科内政治因素的考量）。\n'
    '2. 期刊接受率受众多不可预测因素影响：当期刊物排期、当期投稿量、审稿人可用性、编辑个人偏好等。\n'
    '3. 论文完成P0-P1建议修改后的接受概率可能显著高于当前评估——所有Agent均指出关键修改可实质性'
    '提升接受概率。本报告的"当前概率"反映的是论文v3现状，而非修改后的最优状态。\n'
    '4. 中文期刊的概率估计可能偏高——中文体育学核心期刊的隐性标准（基金项目、机构声望、学术关系网络）'
    '在Agent评估中难以完全量化。', italic=True)

doc.add_page_break()

# ================================================================
# 8. Appendix: Agent-by-Agent Summary
# ================================================================
add_heading_styled(doc, '附录：各期刊Agent评估要点汇总', 1)

# English
add_heading_styled(doc, 'A.1  英文期刊Agent核心评语', 2)

english_summary = [
    ('JSCR (47%)',
     '审稿人会发现选题高度相关（15/29来自JSCR）→送审概率高。核心风险：发表偏倚信号+'
     'g=1.11 vs 0.79张力+中文稿件。修改后可达55-60%。',
     'Agent判断：大修后接受。'),
    ('JSS (22%)',
     '致命风险：GRADE Low + I²=78% + Egger p<0.001的"红旗组合"在JSS审稿人容忍度极低。'
     '即使完成所有修改，接受概率仅升至30-35%。',
     'Agent判断：不建议以当前形态投稿。'),
    ('Biol Sport (35%)',
     '开放科学+方法学创新是加分项。关键要求：补充选择模型发表偏倚分析+解决Cover Letter与正文'
     '的R19处理矛盾。有条件的大修后接受。',
     'Agent判断：有条件可投。'),
    ('Sports Med (5%)',
     'GRADE Low+PI跨零+作者团队非知名+证据体偏小(k=15)。该刊发表权威性综述以指导临床决策，'
     '本文无法满足。推荐转投Sports Med-Open或Scand J Med Sci Sports。',
     'Agent判断：Reject with encouragement.'),
    ('RQES (22%)',
     '致命障碍：RQES要求综述来自"leading research groups"，西南大学在该刊无发表记录。'
     '方法学质量虽好但机构壁垒高。推荐JSCR或Biol Sport。',
     'Agent判断：Borderline — more likely reject.'),
]

for title, assessment, verdict in english_summary:
    add_para(doc, f'**{title}**', bold=True, size=Pt(10), first_line_indent=None, spacing_after=Pt(1))
    add_para(doc, f'评语：{assessment}', size=Pt(9), spacing_after=Pt(1))
    add_para(doc, f'{verdict}', size=Pt(9), italic=True, spacing_after=Pt(6))

add_heading_styled(doc, 'A.2  中文期刊Agent核心评语', 2)

chinese_summary = [
    ('体育科学 (30%)',
     '方法学远超中文期刊平均水平，但作者机构非第一梯队+纯Meta分析在该刊发表难+无基金支持。'
     '建议添加北体大/上体大教授级合作者。',
     'Agent判断：Marginally competitive.'),
    ('中国体育科技 (35%)',
     '纯Meta分析在该刊发表占比<5%，体裁错配。建议重新定位为"方法学基准示范"。'
     '需补充CNKI检索+中国体育场景讨论+全文中文化。',
     'Agent判断：有条件可投。'),
    ('北京体育大学学报 (55%)',
     '方法学骨架极好。三大缺口：CNKI检索+叙事本土化+中文术语统一。'
     '修改后可达70-75%。有机会成为高引文章。',
     'Agent判断：中文首推——有条件高概率。'),
    ('体育学刊 (65%)',
     '方法学远超该刊平均水平，但Meta分析体裁在该刊罕见。关键策略：调整为"系统综述（含定量合成）"'
     '定位+补充学校体育维度讨论。',
     'Agent判断：中高概率——over-qualified但有体裁张力。'),
    ('成都体育学院学报 (62%)',
     '方法学质量显著高于CSSCI扩展版中位水平。PROSPERO+GRADE在投稿池中极为突出。'
     '修改要求最低。',
     'Agent判断：大概率接受。'),
]

for title, assessment, verdict in chinese_summary:
    add_para(doc, f'**{title}**', bold=True, size=Pt(10), first_line_indent=None, spacing_after=Pt(1))
    add_para(doc, f'评语：{assessment}', size=Pt(9), spacing_after=Pt(1))
    add_para(doc, f'{verdict}', size=Pt(9), italic=True, spacing_after=Pt(6))

# ================================================================
# Disclaimer & Methodology
# ================================================================
add_heading_styled(doc, 'B.  评估方法论说明', 2)
add_para(doc,
    '本报告采用多Agent并行独立评审方法：\n'
    '• 10个独立Agent同时评估10个目标期刊\n'
    '• 每个Agent被赋予"20年经验资深审稿人/副主编"角色，基于该期刊的已知特征（IF、接受率、'
    '范围、审稿标准、典型发表论文特征）进行评估\n'
    '• 各Agent独立访问论文的结构化数据摘要（包括方法学、主要结果、创新点、局限性）\n'
    '• 第11个Agent进行跨期刊综合分析和排名合成\n'
    '• 总Agent使用量：12个，总Token消耗：约617K，总耗时：约20分钟\n'
    '• 评估基于论文v3版本（第7轮多Agent审查修改后，评分9.0/10）\n\n'
    '**局限性声明：** Agent模拟不能替代真实同行评审。本报告的评估概率为基于公开信息和'
    '审稿经验推断的最佳估计，实际投稿结果可能因众多不可预测因素而不同。建议将本报告视为'
    '投稿策略参考而非确定性预测。', size=Pt(9))

# Save
output_path = 'D:/桌面/期刊投稿接受概率评估报告_多Agent分析.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
