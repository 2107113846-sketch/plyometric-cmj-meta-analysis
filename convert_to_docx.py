from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Read markdown
with open(r'D:\桌面\科研训练\output\Manuscript_Draft_Methods_Results.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Split into lines
lines = content.split('\n')

in_code_block = False

for line in lines:
    line = line.rstrip()
    
    # Handle code blocks
    if line.startswith('```'):
        in_code_block = not in_code_block
        continue
    
    if in_code_block:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        continue
    
    # Skip empty lines
    if not line:
        continue
    
    # Handle headings
    if line.startswith('# ') and not line.startswith('## '):
        p = doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        p = doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        p = doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        p = doc.add_heading(line[5:], level=4)
    # Handle horizontal rules
    elif line == '---':
        doc.add_paragraph()
    # Handle bullet points
    elif line.startswith('- ') or line.startswith('* '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
    # Handle numbered lists
    elif re.match(r'^\d+\.', line):
        p = doc.add_paragraph(line, style='List Number')
    # Handle blockquotes
    elif line.startswith('>'):
        text = line[1:].strip()
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.italic = True
    # Handle table rows (simple)
    elif '|' in line and line.strip().startswith('|'):
        # Skip separator rows
        if re.match(r'^\|[\s\-\|]+\|$', line):
            continue
        # Parse table cells
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            p = doc.add_paragraph(' | '.join(cells))
            p.style = doc.styles['Normal']
    # Handle regular text
    else:
        p = doc.add_paragraph(line)

# Save
doc.save(r'D:\桌面\Manuscript_Draft_Methods_Results.docx')
print('Word document saved successfully!')
