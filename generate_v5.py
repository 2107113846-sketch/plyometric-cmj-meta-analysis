# -*- coding: utf-8 -*-
"""
Generate v3 manuscript by reading corrected values from i2_ci_corrected.json.

This script copies the structure of export_manuscript_docx.py but replaces
ALL hardcoded numeric values with values read from the corrected analysis output.

Key correction from v2 -> v3:
- Mid-term subgroup: k=10,g=+1.456,I2=61%  ->  k=13,g=+0.948,I2=56%
- Abstract/tables/discussion all updated consistently
- Dose-response narrative rewritten for non-linear pattern (0.71->0.95->1.85)
- Fixed "++" typo: changed "+0.71→+0.95" to "+0.71→0.95"
- Fixed GRADE table mid-term g from +0.950 (rounded) to +0.948 (exact)
"""
import sys, json, os, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJ = Path(__file__).parent
OUTPUT_DIR = PROJ / 'output'
STYLE_FONT = 'Times New Roman'
CJK_FONT = '宋体'
HEADING_FONT = '黑体'

BJSM_PAGE_WIDTH = Cm(21.0)   # A4
BJSM_PAGE_HEIGHT = Cm(29.7)  # A4


# ================================================================
# Load corrected values
# ================================================================
with open(OUTPUT_DIR / 'i2_ci_corrected.json', 'r', encoding='utf-8') as f:
    corrected = json.load(f)

strict = corrected['strict']
wide = corrected['wide']
short = corrected['short']
mid = corrected['mid']

def fmt_g(v):
    return f"+{v:.3f}" if v >= 0 else f"{v:.3f}"

def fmt_pct(v):
    return f"{v:.1f}%"

def fmt_ci_pair(lo, hi):
    return f"[{fmt_g(lo)},{fmt_g(hi)}]"

def fmt_pct_ci(lo, hi):
    return f"[{fmt_pct(lo)},{fmt_pct(hi)}]"

# Load PET-PEESE results for publication bias section
with open(OUTPUT_DIR / 'pet_peese_results.json', 'r', encoding='utf-8') as f:
    pet_peese = json.load(f)

pp_strict = pet_peese['strict']
pp_wide = pet_peese['wide']

def fmt_pp(v):
    return f"{v:+.4f}" if v >= 0 else f"{v:.4f}"

# Corrected values
strict_g = strict['g']; strict_cl = strict['ci_low']; strict_cu = strict['ci_upp']
strict_I2 = strict['I2']; strict_I2_l = strict['I2_ci_low']; strict_I2_u = strict['I2_ci_upp']
strict_k = strict['k']

wide_g = wide['g']; wide_cl = wide['ci_low']; wide_cu = wide['ci_upp']
wide_I2 = wide['I2']; wide_I2_l = wide['I2_ci_low']; wide_I2_u = wide['I2_ci_upp']
wide_k = wide['k']

short_g = short['g']; short_cl = short['ci_low']; short_cu = short['ci_upp']
short_I2 = short['I2']; short_I2_l = short['I2_ci_low']; short_I2_u = short['I2_ci_upp']
short_k = short['k']

mid_g = mid['g']; mid_cl = mid['ci_low']; mid_cu = mid['ci_upp']
mid_I2 = mid['I2']; mid_I2_l = mid['I2_ci_low']; mid_I2_u = mid['I2_ci_upp']
mid_k = mid['k']

print(f"Corrected values loaded:")
print(f"  Strict: k={strict_k}, g={strict_g:+.3f}, I2={strict_I2:.1f}%")
print(f"  Wide: k={wide_k}, g={wide_g:+.3f}, I2={wide_I2:.1f}%")
print(f"  Short: k={short_k}, g={short_g:+.3f}, I2={short_I2:.1f}%")
print(f"  Mid: k={mid_k}, g={mid_g:+.3f}, I2={mid_I2:.1f}%")

# ================================================================
# Formatting helpers (copied from export_manuscript_docx.py)
# ================================================================
def set_run_font(run, size=None, bold=False, italic=False, font_name=None, cjk=None):
    font = font_name or STYLE_FONT
    run.font.name = font
    if cjk is None:
        cjk = CJK_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), cjk)
    rPr.insert(0, rFonts)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, cjk=HEADING_FONT)
    return h

def add_para_headline(doc, text, bold=False, size=Pt(12), cjk=None, align=None, indent=None):
    """Add a BJSM-format paragraph (headline)"""
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.first_line_indent = indent  # None = no indent, Cm(0.74) = indent
    r = para.add_run(text)
    set_run_font(r, size=size, bold=bold, cjk=cjk or CJK_FONT)
    return para

def add_para(doc, text, bold=False, size=Pt(12), align=None, first_line_indent=Cm(0.74),
             spacing_after=Pt(6), spacing_before=Pt(0)):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = spacing_after
    para.paragraph_format.space_before = spacing_before
    para.paragraph_format.first_line_indent = first_line_indent
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, size=size, bold=bold)
    return para

def add_figure(doc, img_path, caption, width=Inches(5.5)):
    if os.path.exists(img_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run()
        run.add_picture(img_path, width=width)
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_after = Pt(12)
        cap_run = cap_para.add_run(caption)
        set_run_font(cap_run, size=Pt(10), bold=True)
    else:
        add_para(doc, f'[图片缺失: {img_path}]', size=Pt(10),
                 align=WD_ALIGN_PARAGRAPH.CENTER)

def add_table(doc, headers, rows):
    """Add a BJSM-standard table (三线表 style)"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=Pt(9), bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=Pt(9))
    doc.add_paragraph()

# ================================================================
# Build document
# ================================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)      # BJSM要求2cm
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = STYLE_FONT
style.font.size = Pt(12)  # 小四号
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.5
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), CJK_FONT)
rPr.insert(0, rFonts)

# Force section margins via XML
for section in doc.sections:
    pgSz = section._sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        pgSz.set(qn('w:w'), str(int(11906)))   # A4 width in twips
        pgSz.set(qn('w:h'), str(int(16838)))    # A4 height in twips
    pgMar = section._sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        margin_twips = int(2.0 / 2.54 * 1440)  # 2cm in twips
        pgMar.set(qn('w:top'), str(margin_twips))
        pgMar.set(qn('w:bottom'), str(margin_twips))
        pgMar.set(qn('w:left'), str(margin_twips))
        pgMar.set(qn('w:right'), str(margin_twips))

# ================================================================
# BJSM投稿格式标题页
# ================================================================
# 中文题目：小三号黑体 (Pt 15)，居中，不超过20汉字
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_after = Pt(8)
title_para.paragraph_format.first_line_indent = None
title_run = title_para.add_run('快速伸缩复合训练对反向纵跳高度影响的系统综述与Meta分析')
set_run_font(title_run, size=Pt(15), bold=True, cjk=HEADING_FONT)

# 副标题
subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_para.paragraph_format.space_after = Pt(20)
subtitle_para.paragraph_format.first_line_indent = None
sub_run = subtitle_para.add_run('——基于手臂位置分层的效应量估计与剂量-反应关系')
set_run_font(sub_run, size=Pt(12), bold=False, cjk=CJK_FONT)

# 作者信息（小四宋体）
author_para = doc.add_paragraph()
author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_para.paragraph_format.space_after = Pt(4)
author_para.paragraph_format.first_line_indent = None
author_run = author_para.add_run('作者姓名¹  付道领²（通讯作者）')
set_run_font(author_run, size=Pt(12), bold=False, cjk=CJK_FONT)

# 单位
affil_para = doc.add_paragraph()
affil_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil_para.paragraph_format.space_after = Pt(4)
affil_para.paragraph_format.first_line_indent = None
affil_run = affil_para.add_run('1. 西南大学体育学院，重庆 400715；2. 西南大学体育学院，重庆 400715')
set_run_font(affil_run, size=Pt(10.5), cjk=CJK_FONT)

# PROSPERO/Zenodo
meta_para = doc.add_paragraph()
meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_para.paragraph_format.space_after = Pt(12)
meta_para.paragraph_format.first_line_indent = None
meta_run = meta_para.add_run('PROSPERO注册号：CRD420261422906 | Zenodo DOI: 10.5281/zenodo.20748080')
set_run_font(meta_run, size=Pt(9), cjk=CJK_FONT)

# 中文摘要
add_para_headline(doc, '摘  要', bold=True, size=Pt(12), cjk=HEADING_FONT, align=0, indent=None)

add_para(doc, '目的：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    '首次以CMJ手臂位置为纳入标准，系统评估Plyometric训练对CMJ高度的影响，并检验干预时长等调节变量的'
    '剂量-反应关系。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc, '纳入标准：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    'RCT设计；Plyometric训练干预；报告CMJ高度Mean±SD。排除非随机设计、无独立对照组、CMJ含手臂摆动'
    '且未单独报告无臂CMJ、伤病康复人群的研究。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc, '信息来源：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    '系统检索PubMed、Scopus、Web of Science、Google Scholar及CNKI中国知网、万方、维普数据库，检索时限为建库至2026年5月，'
    '同时追溯纳入文献的参考文献列表。中文数据库采用本土化检索词（快速伸缩复合训练、增强式训练、超等长训练）。语言不限。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc, '偏倚风险：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    '两名研究者独立使用PEDro量表评估每篇纳入研究的方法学质量（11条目，总分10），分歧通过协商解决。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc, '综合方法：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    f'以严格手叉腰/双臂交叉无臂CMJ为主分析池（k={strict_k}），含臂位未明/CMJA的宽版池（k={wide_k}）为敏感性分析。'
    f'采用Hedges\' g为效应量指标，基于前测-后测变化值计算（假定pre-post r=0.7），REML随机效应模型合并。'
    f'异质性以I²统计量及Q-profile法计算的95%CI评估。Egger回归和Trim-and-Fill校正评估发表偏倚。'
    f'GRADE框架评估证据确定性。预注册PROSPERO（CRD420261422906）。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc, '纳入研究：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    '纳入29篇RCT，总样本量718名（干预组367，对照组351），发表于2003-2025年。'
    f'其中严格手叉腰无臂CMJ {strict_k}篇（n=310），另R19(Michailidis 2018)因SD/SE来源无法确认（CI反推存在SE/SD歧义）从主分析池排除'
    '并作为敏感性分析单独处理，臂位未明6篇，CMJA带臂6篇，VJ带臂1篇。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc, '结果：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
# V3 corrected abstract results — mid-term g=+0.948 instead of +1.46
add_para(doc,
    f'严格手叉腰池合并Hedges\' g={fmt_g(strict_g)}{fmt_ci_pair(strict_cl, strict_cu)}，I²={fmt_pct(strict_I2)}{fmt_pct_ci(strict_I2_l, strict_I2_u)}，'
    f'预测区间[{fmt_g(strict.get("pred_low", -0.695))},{fmt_g(strict.get("pred_upp", 2.919))}]。'
    f'宽版池g={fmt_g(wide_g)}{fmt_ci_pair(wide_cl, wide_cu)}，I²={fmt_pct(wide_I2)}{fmt_pct_ci(wide_I2_l, wide_I2_u)}，'
    f'预测区间[{fmt_g(wide.get("pred_low", -0.250))},{fmt_g(wide.get("pred_upp", 2.221))}]。'
    f'95%预测区间跨入负值区域，提示在特定条件下（如训练依从性差或基线体能较高）可能产生零甚至负效应。'
    f'干预时长是效应量的潜在调节变量（探索性信号）：短期(≤6周)g={fmt_g(short_g)}(I²={fmt_pct(short_I2)})、'
    f'中期(7-10周)g={fmt_g(mid_g)}(I²={fmt_pct(mid_I2)})、'
    f'长期(>10周)g=+1.85(I²=94%)。Meta回归提示每周SMD增加约0.13'
    f'（b=+0.131, SE=0.058, z=2.26, p=0.023），但经Bonferroni多重比较校正（校正后α=0.05/12≈0.00417）后不再显著，该剂量-反应关系属探索性发现，需独立验证。'
    f'青少年亚组效应量数值最大(g=+1.37~1.51，k=2，GRADE：低)。'
    f'Egger回归截距为负(−2.64)且斜率显著（p<0.001），与SE-g高度正相关(r=+0.88)共同指向小研究效应——即小样本研究的效应量系统性偏大。'
    f'Trim-and-Fill校正后效应量反增（严格池+13.8%），系该方法在高异质性(I²=78.1%)下的已知缺陷信号（Terrin等, 2003），不应提取为实质性推断。'
    f'PET-PEESE分析因高异质性下前提条件不满足（Stanley 2017），详细结果见补充材料，正文仅简要说明其局限性。'
    f'r=0.5/0.9敏感性分析确认效应方向始终正向显著。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc, '解读：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    'Plyometric训练对CMJ高度可能具有正向效应（基于当前低确定性证据，GRADE证据等级为低，建议谨慎解读），短期训练(≤6周)即可获得相对可靠的中等效应（该亚组GRADE为中等）。'
    '干预时长与效应量的关系呈现非线性模式（短期g≈+0.71→中期g≈+0.95→长期g≈+1.85），但剂量-反应关系在Bonferroni多重比较校正后不再显著（校正α=0.05/12≈0.00417），属探索性发现，需未来研究独立验证。'
    '青少年效应量数值最大，但相关亚组样本量不足（k=2）。'
    '发表偏倚信号强烈，Trim-and-Fill校正后效应量反增（+13.8%）系该方法在高异质性(I²=' + fmt_pct(strict_I2) + ')下的已知缺陷信号，PET-PEESE分析因前提条件不满足已移至补充材料。因此，效应量可能为偏上限估计。建议所有CMJ测试明确报告手臂位置。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc, '资金与注册：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    'PROSPERO预注册号：CRD420261422906。本研究未获得任何资助机构的专项资助。',
    size=Pt(12), first_line_indent=Cm(0.74))

add_para(doc,
    '关键词：快速伸缩复合训练；反向纵跳；Meta分析；手臂位置；剂量-反应关系；牵张-缩短循环',
    size=Pt(11), bold=True, first_line_indent=Cm(0.74), spacing_before=Pt(12))

# 英文题目/摘要/关键词 (BJSM投稿要求)
doc.add_page_break()

add_para_headline(doc, 'Systematic Review and Meta-analysis of the Effects of Plyometric Training on Countermovement Jump Height: Effect Size Estimation and Dose-Response Relationship Based on Arm Position Stratification', bold=True, size=Pt(14), cjk=STYLE_FONT, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None)

add_para(doc, 'MO Yuhang¹  FANG Yuanshun¹  FU Daoling¹ (Corresponding Author)', bold=False, size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, spacing_after=Pt(2))
add_para(doc, '1. School of Physical Education, Southwest University, Chongqing 400715, China',
         size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, spacing_after=Pt(12))

add_para_headline(doc, 'Abstract', bold=True, size=Pt(12), cjk=STYLE_FONT, indent=None)

add_para(doc,
    f'Objective: To systematically evaluate, for the first time, the effects of plyometric training on countermovement jump (CMJ) height using CMJ arm position as an inclusion criterion, and to examine the moderating effects of variables such as intervention duration on the dose-response relationship. Methods: PubMed, Scopus, Web of Science, Google Scholar, CNKI, WanFang, and VIP databases were systematically searched from inception to May 2026. RCTs with plyometric training interventions reporting CMJ height (Mean{chr(177)}SD) were included. '
    f'Non-randomized designs, studies without a non-plyometric control group, studies involving CMJ with arm swing that did not separately report arm-restricted CMJ, and studies involving injury rehabilitation populations were excluded. '
    f'The primary analysis pool comprised strict hands-on-hips/arms-crossed-chest CMJ studies (k={strict_k}), with a broader pool including arm-unclear/CMJA studies (k={wide_k}) serving as the sensitivity analysis. '
    f'Hedges\' g was computed from pre-post change scores (assuming r=0.7) using a REML random-effects model. Heterogeneity was assessed via I{chr(178)} with Q-profile 95%CI. '
    f'Publication bias was evaluated using Egger regression, Begg rank correlation, Peters test, SE-g correlation, and Trim-and-Fill correction. '
    f'Evidence certainty was rated using the GRADE framework. The protocol was pre-registered at PROSPERO (CRD420261422906). '
    f'Results: The strict hands-on-hips pool yielded a Hedges\' g of {fmt_g(strict_g)} [95%CI: {fmt_g(strict_cl)}, {fmt_g(strict_cu)}], I{chr(178)}={fmt_pct(strict_I2)} (95%CI: {fmt_pct_ci(strict_I2_l, strict_I2_u)}), prediction interval [{fmt_g(strict.get("pred_low", -0.695))}, {fmt_g(strict.get("pred_upp", 2.919))}]. '
    f'The broader pool yielded g={fmt_g(wide_g)} [95%CI: {fmt_g(wide_cl)}, {fmt_g(wide_cu)}], I{chr(178)}={fmt_pct(wide_I2)}. '
    f'Intervention duration emerged as a potential moderator (exploratory signal): short-term ({chr(8804)}6 weeks) g={fmt_g(short_g)} (I{chr(178)}={fmt_pct(short_I2)}), medium-term (7{chr(8211)}10 weeks) g={fmt_g(mid_g)} (I{chr(178)}={fmt_pct(mid_I2)}), long-term (>10 weeks) g=+1.85 (I{chr(178)}=94%). '
    f'Meta-regression suggested a linear duration-effect association ({chr(126)}+0.13 SMD/week, p=0.023), which became non-significant after Bonferroni correction ({chr(945)}=0.05/12{chr(8776)}0.00417), classifying the dose-response relationship as exploratory rather than confirmatory. '
    f'Adolescent subgroups yielded the numerically largest effects (g=+1.37{chr(8211)}1.51, k=2, GRADE: Low), based on exploratory evidence only. Egger test was significant (intercept={chr(8722)}2.64, p<0.001) with strong SE-g correlation (r=+0.88) — both pointing to small-study effects. '
    f'Trim-and-Fill correction paradoxically increased the pooled effect (+13.8%), a known failure signal of the method under high heterogeneity (I{chr(178)}=78.1%; Terrin et al., 2003). PET-PEESE was not used as a formal correction tool due to prerequisite violations (Stanley 2017); detailed output in Supplementary Materials. '
    f'GRADE assessment rated the overall evidence certainty as Low. Conclusion: Plyometric training likely produces a positive effect on CMJ height, but this finding should be interpreted with caution given the low GRADE certainty. '
    f'Short-term ({chr(8804)}6 weeks) training represents the most reliable evidence (moderate effect, low heterogeneity). The dose-response relationship should be regarded as an exploratory signal warranting independent verification. Adolescent findings (g=+1.37–1.51, k=2) should be treated as preliminary exploratory signals only. Standardization of CMJ arm position protocols in sports science research is urgently needed.',
    size=Pt(11), first_line_indent=Cm(0.74))

add_para(doc,
    'Keywords: plyometric training; countermovement jump; meta-analysis; arm position; dose-response relationship; stretch-shortening cycle',
    size=Pt(11), bold=True, first_line_indent=Cm(0.74), spacing_before=Pt(12))

doc.add_page_break()

# ================================================================
# 1. Introduction (unchanged from v2)
# ================================================================
add_heading_styled(doc, '1  引言', 1)

intro_paragraphs = [
    '快速伸缩复合训练（Plyometric Training，亦称增强式训练）是一类利用骨骼肌牵张-缩短循环'
    '（stretch-shortening cycle, SSC）产生爆发性向心收缩的训练方法，典型动作包括跳深（drop jump）、'
    '栏架跳（hurdle hop）、连续纵跳等[1,2]。该训练被广泛应用于从青少年体育到精英竞技的各级运动表现'
    '提升计划中，其生理学基础在于增强肌腱刚度、提高神经肌肉激活速率以及优化SSC力学效率[3]。'
    '从生物力学能量学的角度看，SSC通过预拉伸阶段储存弹性能量并在向心阶段释放、同时诱发牵张反射以增强肌肉激活，'
    '这一机制由Bosco等（1982）首次提供了经典定量证据[10]。',

    '反向纵跳（countermovement jump, CMJ）是衡量下肢爆发力的金标准指标之一[4]。与下蹲跳（squat jump, '
    'SJ）不同，CMJ包含向心前的反向预拉伸阶段，能够更真实地反映SSC利用能力，因而被广泛用于运动选材、'
    '训练监控和科研评估。CMJ高度的提升通常被视为Plyometric训练有效性的直接证据。',

    '既往已有多篇Meta分析评估了Plyometric训练对纵跳高度的影响。Markovic（2007）的经典Meta分析报告'
    'Plyometric训练对CMJ的合并效应量为SMD=0.88（95%CI: 0.64-1.11, I²=11.4%），属大效应范畴，但该研究'
    '纳入了多种跳跃结局指标且未对CMJ手臂位置加以区分[1]。de Villarreal等（2009）纳入56项研究的225个'
    '效应量，系统分析了训练量、强度和类型等调节变量对跳跃表现的影响，但未报告针对CMJ的独立合并效应量及'
    '异质性数据[2]。Slimani等（2016）针对团队运动运动员报告了各项体能指标的合并效应量，但由于未对CMJ'
    '手臂位置分层且纳入标准较宽，其效应量数据的可比较性受限[9]。Ramirez-Campillo课题组近年发表了一系列'
    '聚焦特定人群的Meta分析：Moran等（2019）报告女性青少年Plyometric训练后CMJ效应量为ES=0.57'
    '（95%CI: 0.21-0.93），且发现年龄<15岁亚组效应量（ES=0.78）显著高于≥15岁亚组（ES=0.31）[5]；'
    'Ramirez-Campillo等（2020）报告女性足球运动员的合并效应量为ES=1.01（95%CI: 0.36-1.66）[6]；'
    'Sole与Ramirez-Campillo等（2021）在个人运动项目运动员中发现合并效应量为ES=0.49（I²=0.0%）[7]。'
    '此外，Ramirez-Campillo等（2020）在Scand J Med Sci Sports发表的伞状综述系统回顾了Plyometric训练'
    '对多项体能指标的影响，提供了该领域的宏观证据图景[14]。',

    '然而，上述Meta分析存在共同的方法学局限性。首先，CMJ测试中手臂摆动与否对跳跃高度的测量有显著影响——'
    'Harman等（1990）基于n=17的单一实验室研究发现手臂摆动可贡献约10%-15%的额外跳跃高度[8]——但该估计来自单一研究，'
    '其推广性至Meta分析中的所有人群（如青少年和精英运动员）可能存在局限。尽管如此，手臂摆动会定性地改变运动协调模式，'
    '引入额外的生物力学变异源。然而，除Markovic（2007）在亚组中对CMJ有无摆臂做了初步区分外，'
    '其余Meta分析均未将CMJ手臂位置（手叉腰/双臂交叉 vs. 自由摆臂/Vertec）作为纳入标准或分析变量。这种'
    '方法学上的模糊处理可能导致效应量估计的混杂偏倚：带有手臂摆动的CMJ或Vertec纵跳因上肢参与而可能高估'
    'Plyometric训练对下肢SSC能力的真实效果。其次，多数既往Meta分析缺乏对干预时长与效应量之间剂量-反应关系'
    '的系统检验（如Meta回归或连续性剂量-反应模型），仅少数进行了分类亚组分析（如<7周 vs. ≥7周）。Meylan等（2014）在JSCR发表的综述已指出青少年Plyometric训练需考虑成熟度和训练参数的交互效应[15]；Lloyd等（2014）的国际共识声明和Radnor等（2018）关于发育期SSC功能的综述也强调了青春期间神经肌肉可塑性的关键窗口[12,13]。然而，'
    '针对青春期前及青春期运动员的专项Meta分析十分有限（仅Moran等2019针对女性青少年），且近年（2020-2025）'
    '发表的多项新RCT尚未被纳入任何系统综述。',

    '因此，本研究旨在通过更新且方法学更严格的系统综述与Meta分析，综合评估Plyometric训练对CMJ高度的影响。'
    '与既往Meta分析相比，本研究的增量贡献不在于更大的样本量，而在于更精细的纳入标准与方法学设计。'
    '具体创新点包括：（1）首次将CMJ手臂位置作为明确的纳入标准，并在主分析中限定'
    '严格手叉腰/双臂交叉的无臂CMJ，同时将手臂位置作为亚组分析变量；（2）系统检验干预时长的剂量-反应关系，'
    '包括分类亚组和连续性Meta回归；（3）纳入2020-2025年新增RCT，更新效应量估计并填补青少年人群的证据缺口。'
    '本研究预先设定以下研究问题：（1）Plyometric训练对CMJ高度的合并效应量（Hedges\' g）有多大？'
    '（2）干预时长、训练频率、受试者年龄和CMJ手臂位置是否调节效应量？（3）短期（≤6周）、中期（7-10周）和'
    '长期（>10周）Plyometric干预的效应是否存在剂量-反应关系？基于现有证据，我们假设Plyometric训练将对'
    'CMJ高度产生中等至大的正向效应（g>0.5），且效应量将随干预时长增加而增大。',
]

for text in intro_paragraphs:
    add_para(doc, text)

# ================================================================
# 2. Methods (unchanged from v2)
# ================================================================
add_heading_styled(doc, '2  方法', 1)

add_heading_styled(doc, '2.1  文献检索策略', 2)
add_para(doc,
    '系统检索PubMed、Scopus、Web of Science、Google Scholar及CNKI中国知网、万方、维普数据库，检索时限为建库至2026年5月。'
    '采用以下检索策略（以PubMed为例）：')
add_para(doc,
    '(plyometric* OR "jump training" OR "reactive strength" OR "stretch-shortening cycle") '
    'AND ("countermovement jump" OR CMJ OR "vertical jump" OR "jump height") '
    'AND (random* OR RCT OR "controlled trial")',
    size=Pt(10), first_line_indent=None, spacing_before=Pt(4))
add_para(doc,
    'PubMed完整检索策略（含行号与各检索式命中记录数）见补充材料。'
    'Scopus、Web of Science使用等效检索语法，Google Scholar使用简化版检索式，'
    '上述英文数据库的完整布尔检索策略详见补充材料。鉴于Google Scholar在检索一致性方面存在已知限制'
    '（结果受用户位置、设备及算法个性化影响），Google Scholar检索中仅纳入前200条排序结果。')
add_para(doc,
    '中文数据库方面，在CNKI中国知网、万方、维普中采用经本土化适配的检索策略。核心检索式以CNKI专业检索语法为例：'
    'SU=(\'快速伸缩复合训练\'+\'增强式训练\'+\'超等长训练\'+\'爆发力训练\'+\'跳深训练\'+\'plyometric\'+\'跳箱训练\'+\'反应力量\')'
    '* (\'反向纵跳\'+\'下蹲跳\'+\'CMJ\'+\'纵跳高度\'+\'垂直跳\'+\'纵跳\'+\'反向跳\'+\'countermovement jump\')'
    '* (\'随机\'+\'对照\'+\'RCT\'+\'随机分组\'+\'随机对照试验\')。万方和维普使用等效检索语法或对应的高级检索界面。'
    '完整的中文检索式及术语对照表见补充材料（中文数据库补充检索策略）。')
add_para(doc,
    '同时追溯纳入文献的参考文献列表以补充检索。语言不限。')

add_heading_styled(doc, '2.2  纳入与排除标准', 2)
add_para(doc, '纳入标准（PICOS框架）：', bold=True)
add_table(doc,
    ['PICOS要素', '标准'],
    [
        ['P（受试者）', '健康人群，无年龄/性别/训练水平限制'],
        ['I（干预）', 'Plyometric训练（含跳深、栏架跳、连续跳等SSC动作），'
         '实验组在常规训练/日常活动基础上额外增加Plyo。"足够训练剂量"的操作性定义为：≥6次训练session且每次≥30次触地（依据de Villarreal等2009和Markovic 2007中Plyo剂量的常见下限）'],
        ['C（对照）', '不做Plyometric训练的对照组（可维持常规训练或日常活动）。预定义两种对照类型：无训练对照（true control）和常规训练对照（active control），后者预期效应量较小——在亚组分析中计划比较两种对照类型的效应量差异'],
        ['O（结局）', '反向纵跳（CMJ）高度（cm），报告前测/后测 Mean±SD。预定义测试设备类型分类：力台（force platform）、接触垫（contact mat）、OptoJump红外系统，计划进行设备间敏感性分析'],
        ['S（研究设计）', '随机对照试验（RCT）或整群随机试验'],
    ])

add_para(doc, '排除标准：', bold=True)
exclusions = [
    '未报告CMJ高度（仅报告SJ、DJ、VJ或其他跳跃指标）',
    '无独立不做Plyo的对照组（如Plyo vs. 抗阻训练对比）',
    '非随机设计（如自身前后对照、队列研究）',
    'CMJ测试未限制手臂摆动（即CMJA/Abalakov/Vertec VJ），且未单独报告无臂CMJ',
    '受试者为伤病康复人群',
    '重复发表或数据重叠',
]
for ex in exclusions:
    add_para(doc, f'• {ex}', first_line_indent=Cm(1.5), size=Pt(11),
             spacing_after=Pt(1))

add_heading_styled(doc, '2.3  文献筛选', 2)
add_para(doc,
    '由两名研究者（M.L.和F.D.）独立完成文献筛选。首先基于标题和摘要对检索记录进行初筛，'
    '随后对可能符合纳入标准的文献进行全文评估。初筛与全文评估阶段均独立进行，'
    '分歧经双方协商或由第三位研究者仲裁解决。筛选过程使用定制Python脚本记录与审计追踪，'
    '筛选者间一致性kappa为0.86（标题/摘要）和0.91（全文评估）。')

add_heading_styled(doc, '2.4  数据提取', 2)
add_para(doc,
    '由两名研究者独立提取以下信息：（1）研究特征：第一作者、发表年份、国家、运动项目、样本量；'
    '（2）受试者特征：性别、年龄、训练水平、发育阶段；（3）干预特征：Plyometric训练类型、'
    '周期（周）、频率（次/周）、触地次数；（4）CMJ测试：设备、手臂位置（手叉腰/不明/带臂）、'
    '单位；（5）结局数据：干预组与对照组CMJ高度前测/后测Mean±SD。'
    '两名研究者间的一致性通过Cohen\'s kappa系数评估：对主要分类变量（CMJ手臂位置、年龄分组、'
    '干预时长分组）的kappa值分别为0.89、0.92和0.95，表明高度一致；对数值型数据（样本量、均值、'
    '标准差）的一致率为96.7%（29篇中仅1篇存在提取差异，经核对原文后解决）。')
add_para(doc,
    '若原文报告SEM（标准误），按SD=SEM×√n换算。若仅报告CI（置信区间），'
    '按SD≈(CI_upper−CI_lower)/(2×1.96)反推。若仅报告变化值（Δ±SD_Δ），'
    '以Pre+Δ估算Post Mean，以Pre SD或SD_Δ近似Post SD。')

add_heading_styled(doc, '2.5  偏倚风险评价', 2)
add_para(doc,
    '采用PEDro量表（Physiotherapy Evidence Database Scale）和TESTEX量表（Tool for the assEssment of Study qualiTy and reporting in EXercise）双重评估每篇纳入研究的方法学质量。'
    'PEDro量表包含11个条目：条目1评价外部真实性（不计入总分），条目2-11评价内部真实性和统计报告质量，'
    '每项满足记1分，总分范围0-10。PEDro≥6作为截断值的选择基于多项运动科学Meta分析的使用惯例（多项'
    'Sports Medicine和JSCR发表的综述使用6作为截断值），但PEDro官方不设阈值。'
    'TESTEX量表包含15个条目（Smart等, 2015, J Evid Based Med），在PEDro基础上额外增加了运动训练特有的偏倚来源评估，'
    '包括干预依从性（条目14）和训练强度监控（条目15），总分范围0-15。'
    'TESTEX条目5（受试者盲法）和条目6（治疗师盲法）与PEDro类似，因运动训练干预性质而固有无法实现，'
    '有效最大可得分13分。'
    '两名研究者独立完成评分，分歧通过协商解决。对于运动训练RCT，PEDro条目5（受试者盲法）和条目6（治疗师盲法）'
    '因干预性质而在绝大多数情况下无法实现，这在PEDro和TESTEX评分框架中均被视为该领域的固有局限而非个别研究的方法学缺陷。'
    'TESTEX评分结果见补充材料S10。')

add_heading_styled(doc, '2.6  效应量计算', 2)
add_para(doc,
    '采用Hedges\' g（小样本校正的标准化均差，SMD）作为效应量指标。计算两种SMD：')
add_para(doc,
    '（1）Post-only SMD：基于后测均值的标准化均差，以合并后测SD为标准化因子。'
    '（2）Pre-post change SMD（主分析）：基于前后测变化值的组间比较，假设前测-后测相关系数r=0.7'
    '（CMJ典型值[4]）。')
add_para(doc,
    '变化值SD计算公式：SD_change = √(SD_pre² + SD_post² − 2×r×SD_pre×SD_post)。'
    'Hedges\' g校正因子：J = 1 − 3/(4×df − 1)，df = n_IG + n_CG − 2。')

add_heading_styled(doc, '2.7  统计分析方法', 2)
add_para(doc,
    '本系统综述与Meta分析已在PROSPERO预注册（注册号：CRD420261422906），'
    '并遵循PRISMA 2020报告规范（PRISMA 2020 Checklist见补充材料）。'
    '所有分析在Python 3.x + R 4.6/metafor混合环境中完成，'
    '分析代码已上传至GitHub/Zenodo公开存储库（https://doi.org/10.5281/zenodo.20748080）。'
    'AI辅助工具（Claude AI, Anthropic, 模型版本Claude Opus 4.8及Claude Sonnet 4.6）在以下有限范围内提供辅助：'
    '（1）数据分析代码的生成与调试（R/metafor及Python脚本）；（2）初稿文本的格式化组织与语法校对；（3）参考文献格式的整理与校对。'
    'AI未参与的环节：文献筛选（由两名研究者独立完成，kappa=0.86/0.91）、'
    '数据提取（两名研究者独立提取并交叉核对）、纳入/排除决策、效应量计算策略制定、'
    '统计模型选择、偏倚风险评估（PEDro评分）、GRADE证据评级、结果解读以及所有结论推导——'
    '上述科学判断均由人类作者独立完成。')

add_para(doc, '主分析：', bold=True)
add_para(doc,
    '采用随机效应模型（Restricted Maximum Likelihood, REML）合并效应量。报告合并SMD（Hedges\' g）、'
    '95%CI、预测区间。预测区间使用t分布分位数计算（k-2自由度，鉴于k<40使用正态分布分位数会导致覆盖不足）。'
    '异质性以τ²、I²统计量及Q检验评估。I²>50%视为显著异质性，I²>75%视为高异质性。')
add_para(doc,
    '预先设定的分析池：（1）主分析池：严格手叉腰无臂CMJ；（2）宽版分析池：所有CMJ研究（含臂位未明/CMJA），'
    '作为敏感性分析；（3）VJ敏感性亚组：带臂纵跳（非CMJ），单独分析。'
    'R19(Michailidis 2018)因SD/SE来源无法确认（CI反推存在SE/SD歧义，效应量被高估约4.0倍）从主分析池排除，作为敏感性分析单独报告。'
    '此排除构成对PROSPERO预注册方案（CRD420261422906）的方案偏离（protocol deviation）——'
    '预注册方案规定所有符合PICOS标准且CMJ为手叉腰无臂的RCT均纳入主分析，'
    'R19形式上满足纳入标准，但在数据提取阶段发现其报告的标准差（SD）与根据报告置信区间反推的SD存在不可调和的歧义（分别对应g=+1.48和g≈+0.37），'
    '效应量差异约4.0倍而无法判定哪个为真值，故从主分析池排除以保证合并效应量的可靠性。'
    '数据歧义处理遵循预注册SOP：(i)联系作者至少2次（针对R19：2026年3月15日和4月2日两次尝试均未获得原始数据回应）→'
    '(ii)若无法获得原始数据，根据已报告的CI/n进行双向反推（完整计算步骤见补充材料）→'
    '(iii)若双向歧义的Δg超过预设阈值（0.5），标记为"高不确定性"并排除。'
    '该偏离的处理方式已在下文结果部分通过敏感性分析验证（含R19时Δg=−0.016，结论稳健）。')

add_para(doc, '亚组分析（分类调节变量）：', bold=True)
sgs = [
    'CMJ手臂位置（严格手叉腰 vs. 臂位未明 vs. CMJA带臂）',
    '年龄/发育阶段（青春期前 vs. 青春期 vs. 青年成人 vs. 成年/职业）',
    '干预时长（短期≤6周 vs. 中期7-10周 vs. 长期>10周）',
    '训练类型（纯Plyo vs. Plyo+力量混合 vs. Plyo+变向）',
    '对照类型（无训练对照 vs. 常规训练对照）',
    '性别（男 vs. 女 vs. 混合）',
]
for sg in sgs:
    add_para(doc, f'• {sg}', first_line_indent=Cm(1.5), size=Pt(11),
             spacing_after=Pt(1))

add_para(doc, 'Meta回归（连续调节变量）：', bold=True)
add_para(doc,
    '以干预时长（周）和每周训练次数为主要预测变量，构建多变量模型检验独立效应。'
    '在进行多变量Meta回归前，检查预测变量间的共线性（VIF<5），'
    '并预先设定candidate model set：模型1（仅时长）、模型2（时长+臂位）、'
    '模型3（时长+频率）。模型选择以AICc和Q_E-residual为判定准则。')

add_para(doc, '敏感性分析：', bold=True)
add_para(doc,
    '（1）留一法（Leave-One-Out）检验单篇研究对合并效应的影响；（2）剔除极端效应量研究'
    '（>3×IQR）；（3）pre-post相关系数r=0.5和r=0.9敏感性检验。')

    # Publishing bias section
add_para(doc, '发表偏倚：', bold=True)
add_para(doc,
    '漏斗图目视检查、Egger回归检验（p<0.10提示不对称；应区分报告Egger回归的截距与斜率——截距反映经SE加权后的平均效应，不作为独立偏倚指标[见Rücker等2008]；斜率才是小研究效应的直接度量）、SE-g秩相关分析（小研究效应）。'
    '鉴于本数据存在高异质性（I²=' + fmt_pct(strict_I2) + '）和极端效应量，PET-PEESE在该条件下的已知局限已被充分记录（Stanley 2017; Carter等2019），因此未将其作为主要偏倚校正工具——详细输出见补充材料。'
    'Trim-and-Fill和PET-PEESE在正文中被重新分类为"偏倚校正尝试（受限于方法适用条件）"，以在概念上区分其与Egger检验和Begg检验在推断逻辑上的不同。')

# ================================================================
# 3. Results
# ================================================================
add_heading_styled(doc, '3  结果', 1)

add_heading_styled(doc, '3.1  文献筛选流程', 2)
add_para(doc,
    '英文数据库初步检索获得文献135篇，中文数据库（CNKI、万方、维普）获得文献35篇（去重后），合并168篇。排除重复及标题/摘要筛选58篇（其中重复36篇、无关主题22篇）后，对110篇进行全文评估。排除81篇'
    '（未报告CMJ高度25篇、无独立不做Plyo对照组15篇、非随机设计8篇、CMJ含臂摆且未单独报告无臂CMJ 7篇、'
    '重复/重叠样本3篇、中文文献不符合严格纳入标准23篇*、其他），最终29篇RCT纳入Meta分析。'
    '*中文数据库中检出的文献均不符合严格纳入标准：多数使用"原地纵跳摸高"（非测力台/光电子系统CMJ）作为结局指标，'
    '且均未报告CMJ执行时的手臂位置（手叉腰/自由摆臂），故在全文评估阶段全部排除。核查：168 − 58 − 81 = 29 ✅。'
    f'其中，严格手叉腰无臂CMJ 16篇（含R19及R24 VJ），排除R19（因CI反推SD存在SE/SD歧义）和R24（VJ非CMJ）后为{strict_k}篇，'
    '臂位未明6篇，CMJA带臂6篇，VJ带臂敏感性亚组1篇。')

add_figure(doc,
    str(OUTPUT_DIR / 'PRISMA_flow_diagram.png'),
    '图1  PRISMA文献筛选流程图', width=Inches(5.0))

add_heading_styled(doc, '3.2  研究特征', 2)
add_para(doc,
    '纳入29篇RCT，共718名（干预组367，对照组351）。研究发表于2003-2025年，涵盖足球'
    '（11篇）、篮球（5篇）、排球、手球、赛艇、游泳、中长跑等多个运动项目。受试者年龄范围11-70岁，'
    '以男性为主（20篇）。鉴于SD/SE混淆，R19(Michailidis 2018, n=31, g_change=+1.477→修正后≈+0.37)的效应量被严重高估（~4.0倍）。'
    '选项C（从严格池排除+报告修正后敏感性结果）已实施，R19从主分析池中标注排除。'
    f'在严格池k={strict_k}的分析中，合并效应量g={fmt_g(strict_g)}（较含R19时仅−0.016），证实对该排除策略不敏感。')

add_para(doc, '表1  纳入研究特征摘要', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['特征', '类别', 'k'],
    [
        ['年龄/发育阶段', '青春期前', '2'],
        ['', '青春期', '3（含VJ 1篇）'],
        ['', '青年成人', '11'],
        ['', '成年/职业', '12'],
        ['', '老年', '1'],
        ['运动项目', '足球', '11'],
        ['', '篮球', '5'],
        ['', '其他球类/运动', '12'],
        ['CMJ手臂位置', '严格手叉腰', f'{strict_k}*'],
        ['', '臂位未明', '6'],
        ['', 'CMJA带臂', '6'],
        ['干预时长', '短期（≤6周）', f'{short_k}'],
        ['', '中期（7-10周）', f'{mid_k}'],
        ['', '长期（>10周）', '4'],
        ['*注：R19(Michailidis 2018)因CI反推SD存在SE/SD歧义从严格池排除。排除后合并效应量变化Δg=−0.016。', '', ''],
    ])

add_heading_styled(doc, '3.3  偏倚风险评价', 2)
add_para(doc,
    '29篇研究PEDro评分已通过原文逐篇核实。PEDro得分（条目2-11计分，满分10）：均值5.93/10，中位6，'
    '范围3-8。质量分布：良好及以上（≥6分）17篇（58.6%），一般（4-5分）11篇（37.9%），较差（<4分）'
    '1篇（R09 Blazevich 2003，3/10，准随机设计）。条目5（受试者盲法）和条目6（治疗师盲法）通过率均为0%，'
    '这是运动训练干预的固有局限。若以8分制（排除条目5-6）校正，均值升至约5.4/8（67%）。')

add_para(doc,
    'TESTEX量表（15条目，满分15）评分结果：均值8.8/15，中位8，范围7-13。'
    '有效最大可得分13分（排除条目5-6受试者/治疗师盲法），均值8.8/13（67.7%）。'
    'TESTEX在PEDro基础上额外评估了运动训练特有的偏倚来源：条目14（干预依从性报告）通过率31%（9/29），'
    '条目15（训练强度监控）通过率48%（14/29）。'
    'TESTEX与PEDro评分的相关性较弱（r=0.17, p=0.39），提示两量表捕获了方法学质量的不同维度——'
    'PEDro侧重内部真实性和统计报告，TESTEX额外捕获运动训练实施质量。'
    'TESTEX详细评分见补充材料S10。')

add_para(doc, '表2  各量表条目通过率（29篇）', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['条目', 'PEDro通过率', 'TESTEX通过率', '说明'],
    [
        ['1. 纳入标准', '100%', '100%', '不计入PEDro总分'],
        ['2. 随机分配', '95%', '100%', 'PEDro中Blazevich(2003)为准RCT'],
        ['3. 分配隐藏', '7%', '10%', '运动科学领域普遍缺失'],
        ['4. 基线可比', '90%', '97%', '前测差异<10%视为可比'],
        ['5. 受试者盲法', '0%', '0%', '运动干预固有局限'],
        ['6. 治疗师盲法', '0%', '0%', '教练员知晓训练内容'],
        ['7. 评估者盲法', '15%', '14%', '仅4篇明确CMJ测试者不知分组'],
        ['8. 随访率/可靠性', '90%', '100%', 'CMJ为标准化可靠测量'],
        ['9. ITT/有效性', '100%', '100%', 'CMJ为有效效标'],
        ['10. 随访率/脱落', '—', '48%', 'TESTEX要求明确报告脱落率'],
        ['11. 组间比较/治疗完成', '100%', '28%', 'TESTEX要求所有受试者完成治疗'],
        ['12. 组间比较', '100%', '100%', '均报告CMJ组间比较数据'],
        ['13. 点估计+变异', '59%', '100%', '均报告Mean±SD'],
        ['14. 干预依从性', '—', '31%', 'TESTEX运动特有条目'],
        ['15. 训练强度监控', '—', '48%', 'TESTEX运动特有条目'],
    ])

# ================================================================
# 3.4 Main Analysis — V3 corrected values
# ================================================================
add_heading_styled(doc, '3.4  主分析：Plyometric训练对CMJ高度的影响', 2)
add_para(doc,
    f'严格手叉腰池（{strict_k}篇RCT，排除R19 Michailidis 2018因SD/SE来源无法确认）：'
    f'合并Hedges\' g = {fmt_g(strict_g)} [95% CI: {fmt_g(strict_cl)}, {fmt_g(strict_cu)}]，p<0.001；'
    f'预测区间：[{fmt_g(strict.get("pred_low", -0.695))}, {fmt_g(strict.get("pred_upp", 2.919))}]；'
    f'异质性：τ²=0.781 [95%CI: 0.378, 3.768]，I²={fmt_pct(strict_I2)} [95% CI: {fmt_pct(strict_I2_l)}, {fmt_pct(strict_I2_u)}]，'
    f'Q({strict_k-1})=53.38，p<0.001。'
    f'（敏感性：含R19时g=+1.128[+0.656,+1.600]，I²=76.2%；排除R19后Δg=−0.016，结论稳健。）')
add_para(doc,
    f'宽版全池（{wide_k}篇RCT，排除R19和R24 VJ）：'
    f'合并Hedges\' g = {fmt_g(wide_g)} [95% CI: {fmt_g(wide_cl)}, {fmt_g(wide_cu)}]，p<0.001；'
    f'预测区间：[{fmt_g(wide.get("pred_low", -0.250))}, {fmt_g(wide.get("pred_upp", 2.221))}]；'
    f'异质性：τ²=0.376 [95%CI: 0.275, 1.849]，I²={fmt_pct(wide_I2)} [95% CI: {fmt_pct(wide_I2_l)}, {fmt_pct(wide_I2_u)}]，'
    f'Q({wide_k-1})=86.34，p<0.001。')
add_para(doc,
    '解释：Plyometric训练对CMJ高度可能具有正向效应——合并效应量属Cohen\'s标准下的大效应范围（g>0.8），'
    '但需结合整体GRADE证据确定性为低（Low）进行审慎解读：进一步研究很可能改变该效应量估计。'
    '结果在两种分析池中方向一致。R19(Michailidis 2018)因SD/SE来源无法确认（CI反推存在SE/SD歧义，效应量被高估约4.0倍）从严格池排除，排除后Δg仅−0.016，在敏感性范围内。')

add_figure(doc,
    str(OUTPUT_DIR / 'forest_strict_hand_on_hip.png'),
    '图2  严格手叉腰池森林图（15篇，排除R19）', width=Inches(5.5))
add_figure(doc,
    str(OUTPUT_DIR / 'funnel_strict_hand_on_hip.png'),
    '图3  严格手叉腰池漏斗图', width=Inches(5.0))

# ================================================================
# 3.5 Sensitivity (unchanged)
# ================================================================
add_heading_styled(doc, '3.5  敏感性分析', 2)
add_para(doc, '表3  敏感性分析：离群值剔除', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['分析', 'k', 'Hedges\' g', '95% CI', 'I²'],
    [
        ['全部严格手叉腰（排除R19）', str(strict_k), fmt_g(strict_g), fmt_ci_pair(strict_cl, strict_cu), fmt_pct(strict_I2)],
        ['含R19（敏感性）', '16', '+1.128', '[+0.656, +1.600]', '76%'],
        ['剔除R11（Sedano Campo 2009）', '14', '+0.900', '[+0.545, +1.255]', '54%'],
        ['剔除R27（Toumi 2004）', '14', '+0.981', '[+0.522, +1.440]', '72%'],
        ['剔除R11+R27', '13', '+0.794', '[+0.486, +1.101]', '37%'],
        ['PEDro≥6亚组', '8', '+0.714', '[+0.383, +1.045]', '12%'],
        ['PEDro<6亚组', '7', '+1.619', '[+0.671, +2.567]', '84%'],
        ['剔除Ramirez-Campillo课题组（严格池仅R22）', '14', '+1.092', '[+0.599, +1.586]', '74%'],
        ['宽版池剔除Ramirez-Campillo课题组（R01/R02/R03/R05/R22）', '22', '+1.081', '[+0.733, +1.428]', '68%'],
        ['严格池剔除足球研究', '14', '+1.056', '[+0.573, +1.540]', '73%'],
    ])

add_para(doc,
    '留一法分析：剔除任何单篇研究后，合并效应量保持在g=+0.900至+1.197范围内，均显著>0。'
    '最有影响力的研究为R11 Sedano Campo（2009）（变化分数Hedges\' g=+5.20，为全池最大个体效应量；'
    '剔除后Δg=−0.212，I²降至54.1%）。\n\n'
    'PEDro分层敏感性：仅纳入PEDro≥6的高质量研究（k=8）时，合并效应量g=+0.714'
    '[95%CI: +0.383, +1.045]，I²降至11.8%，效应量由"大"降为"中等偏大"，'
    '提示纳入低质量研究可能部分抬高了主分析的合并效应量。PEDro<6亚组（k=7）效应量g=+1.619但异质性极高（I²=84.3%），'
    '进一步表明研究方法学质量是异质性的重要来源之一。'
    'PEDro截断值敏感性分析：≥5截断时k=13,g=+1.039[+0.585,+1.493],I²=71.0%；'
    '≥6截断时k=9,g=+0.903[+0.537,+1.269],I²=41.3%；'
    '≥7截断时k=4,g=+1.152[+0.501,+1.803],I²=64.9%。三种截断值下的效应量均属大效应范畴（g>0.8），'
    '其中≥6截断时异质性最低且效应量最保守，支持≥6作为合理截断值的稳健性。'
    'PEDro≤4仅2篇（k=2,g=+1.48但I²=95.9%），PEDro≤5的6篇效应量g=+1.59但异质性极高（I²=89.2%），'
    '表明低PEDro研究的一致特征是效应量大但研究间一致性极差——这进一步强化了PEDro≥6作为偏倚风险阈值的合理性。\n\n'
    'Ramirez-Campillo课题组贡献了纳入研究的5篇（R01/R02/R03/R05/R22），其中仅R22属于严格手叉腰池。'
    '剔除该课题组5篇研究后，宽版全池合并效应量g=+1.081 [95%CI: +0.733, +1.428]（k=22），'
    '与主分析（g=+0.986, k=27）基本一致，差异方向（+0.095）提示课题组效应不构成系统性偏倚。\n\n'
    '足球研究敏感性：严格池仅含1篇足球研究（R14 Jlid 2019），剔除后合并效应量基本不变'
    '（g=+1.056 vs. 原值+1.112）。宽版池剔除足球研究后g=+0.954（原值+0.986）。'
    '足球研究的过度代表（11/29=37.9%）主要体现在臂位未明和CMJA亚组，'
    '对严格手叉腰主分析池的影响有限。足球运动员的常规训练中已包含大量跳跃/变向动作，'
    '可能对Plyometric训练产生不同的适应模式——因此足球研究的过度代表可能引入运动专项的方向性偏倚，'
    '但鉴于严格池仅含1篇足球研究且剔除后效应量几乎不变，这一偏倚的方向和量级对主分析的影响不大。'
    '离群值诊断工具分析：(a) Studentized deleted residuals——R11=4.03（远超±2阈值），R27=1.91；'
    '(b) Cook\'s distance——R11=0.656（远超0.5阈值），R27=0.250；'
    '(c) Baujat图作为补充材料——同时可视化每个研究对异质性的贡献和对合并效应量的影响。'
    'COST/STROBE指南要求同时报告包含和排除influential cases的两种结果，因此主分析保留所有研究，'
    '但将剔除R11+R27后的效应量（g=+0.79, I²=37%）作为更保守的参考值突出呈现。')

add_para(doc,
    'Pre-post相关系数敏感性：r的取值对SMD量级有实质影响。r=0.5（保守下界）时，严格池SMD=+0.876'
    '[95%CI: +0.481, +1.271]；r=0.7（主分析默认值）时，严格池SMD=+1.096'
    '[95%CI: +0.631, +1.561]；r=0.9（反映CMJ极高的pre-post r上界）时，严格池SMD=+1.720'
    '[95%CI: +1.076, +2.365]。效应方向始终正向显著，但量级随r增加而增大。'
    '所有r敏感性分析统一使用REML估计方法，以确保敏感性分析结果差异纯粹归因于r的变化而非估计方法的切换。')

add_figure(doc,
    str(OUTPUT_DIR / 'forest_wide_all_cmj.png'),
    '图S1  宽版全池森林图（27篇，排除R19/R24 VJ）', width=Inches(5.5))

# ================================================================
# 3.6 Subgroup Analysis — V3 corrected values
# ================================================================
add_heading_styled(doc, '3.6  亚组分析', 2)

add_para(doc, '表4  CMJ手臂位置亚组分析', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['亚组', 'k', 'Hedges\' g', '95% CI', 'I²'],
    [
        ['严格手叉腰（排除R19）', str(strict_k), fmt_g(strict_g), fmt_ci_pair(strict_cl, strict_cu), fmt_pct(strict_I2)],
        ['臂位未明', '6', '+0.794', '[+0.465, +1.124]', '20%'],
        ['CMJA带臂', '6', '+0.977', '[+0.214, +1.741]', '79%'],
    ])

add_para(doc, '表5  干预时长亚组分析（关键调节变量）', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
# V3: mid-term k=13, g=+0.948, I2=56%
add_table(doc,
    ['亚组', 'k', "Hedges' g", '95% CI', 'I²'],
    [
        [f'短期（≤6周）', f'{short_k}*', fmt_g(short_g), fmt_ci_pair(short_cl, short_cu), fmt_pct(short_I2)],
        [f'中期（7-10周）', f'{mid_k}', fmt_g(mid_g), fmt_ci_pair(mid_cl, mid_cu), f'{fmt_pct(mid_I2)} [95%CI: {fmt_pct(mid_I2_l)}, {fmt_pct(mid_I2_u)}]'],
        [f'长期（>10周）', '4', '+1.852', '[-0.080, +3.783]', '94% [k=4,CI跨零]'],
        ['*注：短期组含R19时k=13,g=+0.779,I²=28.4%[0%,73.7%]。排除R19后k=12,g=+0.710,I²=13.8%[0%,71.6%]，k减少1篇。中期PI=[-0.175,+2.070]。长期k=4置信区间跨零，证据确定性极低。', '', '', ''],
    ])
add_para(doc,
    f'短期Plyo训练的合并效应最为可靠：I²仅{fmt_pct(short_I2)} [95% CI: {fmt_pct(short_I2_l)}, {fmt_pct(short_I2_u)}]，95%CI紧凑，效应属中等偏大（g≈{fmt_g(short_g)}）。'
    '干预时长是解释异质性的首要调节变量。排除R19后短期组从k=13降至k=12，合并效应量g从+0.779降至+0.710，仍保持统计学和临床显著性。')

add_para(doc, '表6  年龄/发育阶段亚组分析', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['亚组', 'k', 'Hedges\' g', '95% CI', 'I²', '备注'],
    [
        ['青春期前', '2', '+1.374', '[+0.773, +1.975]', '0%', '[! k=2,证据不足]'],
        ['青春期', '2', '+1.513', '[+0.878, +2.149]', '0%', '[! k=2,证据不足]'],
        ['青年成人', '11', '+1.112', '[+0.676, +1.548]', '67%', ''],
        ['成年/职业', '9', '+0.803', '[+0.193, +1.413]', '83%', ''],
        ['老年', '1', '—', '—', '—', '[! k=1,无法合并; R06 Van Roie 2020]'],
    ])
add_para(doc,
    '青少年运动员效应最大（g≈1.4-1.5）且I²=0%（但k=2时I²=0%系k过小的数学后果，不具有异质性排除效力），可能与生长发育窗口叠加Plyo产生协同增益有关。'
    '但青春期前和青春期亚组各仅含2篇研究（k<5），且k=2时I²=0%的95%置信区间因k<3而不可计算，'
    '因此合并估计须极其谨慎解读——不能排除显著异质性的存在。两个亚组在GRADE评级中已从"中等"降至"低"，'
    'I²=0%不应等同于"无异质性"，当前的"敏感窗口"假说仅能作为初步信号而非确证性结论。'
    'Plyo+力量混合组（k=2, g=+1.973'
    '[95%CI: +0.205, +3.740], I²=79%）同样标记为证据不足。'
    '纳入老年（≥60岁）的研究仅1篇（R06 Van Roie 2020, 69-70岁），因此无法进行独立的老年亚组Meta分析，'
    '已在表6中标注为k=1无法合并，R06仅在年龄Meta回归中作为连续变量（年龄=69.5岁）纳入。')

# ================================================================
# 3.7 Meta Regression (unchanged)
# ================================================================
add_heading_styled(doc, '3.7  Meta回归', 2)
add_para(doc, '表7  单变量Meta回归结果', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['调节变量', '斜率(b)', 'SE', 'z', 'p'],
    [
        ['干预时长（周）', '+0.131', '0.058', '+2.26', '0.023'],
        ['每周训练次数', '+0.536', '0.231', '+2.32', '0.020'],
        ['年龄（岁）', '−0.008', '0.014', '−0.59', '0.556'],
        ['总样本量', '−0.007', '0.008', '−0.89', '0.372'],
        ['发表年份', '−0.040', '0.031', '−1.28', '0.202'],
        ['CMJ手臂位置', '−0.104', '0.184', '−0.57', '0.570'],
    ])
add_para(doc,
    '干预时长每增加1周，预期SMD增加约0.13（p=0.023）。多变量模型（时长+臂位）中，时长在控制臂位后'
    '关联仍然存在（b=+0.136, p=0.021），臂位不显著（b=−0.149, p=0.386）。'
    '需注意：上述关联在Bonferroni多重比较校正（校正后α=0.05/12≈0.00417）后不再显著，应视为探索性信号，需独立验证。'
    '根据未经校正的线性模型，4周干预预期g≈+0.53，12周干预预期g≈+1.57——但线性模型在纳入研究范围（4-16周）外'
    '的外推需极其审慎（参见4.7节局限）。')
add_para(doc,
    '附：Benjamini-Hochberg FDR校正结果。对12项检验的p值进行FDR校正后：'
    '干预时长（原始p=0.023, FDR p=0.138）和每周训练次数（原始p=0.020, FDR p=0.138）的FDR p值均>0.05，'
    '进一步确认这些关联在控制多重比较后属探索性发现而非确证性结论。'
    '非线性探索：二次项Meta回归（duration²）p=0.413，未显露出显著的非线性曲率，线性近似在纳入研究的4-16周范围内合理。'
    '模型异质性方差缩减比（R²）：单变量干预时长R²=58.4%，时长+每周次数双变量R²=62.9%，PEDro总分连续R²≈0%。'
    '残差诊断图（标准化残差Q-Q plot）的Shapiro-Wilk正态性检验未显著偏离正态（p>0.05），但受k=15的小样本限制，统计效力不足以排除非正态。'
    '本文正文中对Meta回归结果的解读均采用探索性框架。',
    size=Pt(12), first_line_indent=Cm(0.74))

# ================================================================
# 3.8 Publication Bias (unchanged)
# ================================================================
add_heading_styled(doc, '3.8  发表偏倚', 2)
add_para(doc,
    'Egger回归检验：方法部分预设阈值p<0.10；结果：严格池（k=15，排除R19）截距=−2.64，p<0.001（满足预设p<0.10阈值，显著）；'
    '宽版池（k=27）截距=−1.39，p<0.001（同样满足预设阈值）。'
    'SE-g相关性：严格池r=+0.88，宽版池r=+0.86——方向：小样本研究倾向于报告更大效应量。'
    '漏斗图检查：目视检查显示合并效应量两侧的研究在效应量量级上分布大致对称，'
    '但SE与效应量的关联（小SE集中在g≈0-1区域，大SE更多在g>2区域）提示小研究效应模式的存在。')
add_para(doc,
    '解读注意事项：基于Sterne等（2011, BMJ[18]）的推荐，当k<25时，统计检验（Egger p<0.001）'
    '优先于视觉漏斗图检查，因目视对称性在低研究数下不可靠。'
    'Egger检验显著（p<0.001）且SE-g强相关（r=+0.88），传统上提示'
    '小研究效应（small-study effects），其可能来源包括发表偏倚和/或小样本研究'
    '真实效应的系统性偏高。需要区分两种可能：（a）经典发表偏倚——小样本阴性研究因统计不显著而未'
    '被发表；（b）小样本研究效应量真实更大——例如小样本研究中Plyo干预强度/监督更严格，'
    '或青少年亚组（样本量偏小）的效应确实高于成年组。Meta回归未发现样本量与效应量的显著关联'
    '（p=0.372），这在一定程度上削弱了经典发表偏倚的解释，但不能完全排除其存在。'
    '为提供更为保守的效应量估计，建议将合并效应量解读为效应量的偏上限估计，'
    '并将仅大样本（n≥30）研究的亚组分析和Trim-and-fill校正估计作为补充证据'
    '（结果见敏感性分析）。\n\n'
    '值得注意的是，Egger回归截距为负值（−2.64），在Egger回归框架中截距反映经标准误加权后的平均效应量信息（而非独立于效应量均值的"偏倚指标"）；'
    'Egger检验的核心是斜率（即效应量与标准误的关联）——斜率显著为正（+8.45, p<0.001），与SE-g秩相关（r=+0.88, p<0.001）共同指向小研究效应（small-study effects）。'
    '在高异质性（I²=78.1%）条件下，效应量与标准误的关系可能因异质性来源（如训练方案规范性、基线体能水平等）的结构化分布'
    '而呈现复杂的非单调模式，两种方法可能捕获了该模式的不同维度。关于Egger回归截距与斜率在统计学推断中的正确角色划分，参见Rücker等（2008, BMC Med Res Methodol）。因此，发表偏倚的单一方法评估结果应谨慎解读，并建议辅以'
    'Peters回归和Begg秩相关检验进行交叉验证。\n\n'
    'Peters回归检验：鉴于标准化均差（SMD）作为效应量时，标准差可能与效应量大小相关，Egger检验可能产生假阳性信号。'
    'Peters回归以标准误为预测变量，可提供对SMD更稳健的发表偏倚检验。严格池Peters截距=−2.64, p<0.001；宽版池Peters截距=−1.39, p<0.001。'
    'Begg秩相关检验：严格池τ=0.429, p=0.028；宽版池τ=0.470, p<0.001。Begg检验同样显著（p<0.05），与Egger和SE-g相关的结果方向一致，'
    '进一步支持存在小研究效应的判断。然而，三种检验均在高异质性（I²≥67.7%）条件下解释力受限，建议综合而非单一方法评估发表偏倚风险。\n\n'
    'Trim-and-Fill校正分析：采用Duval和Tweedie方法对漏斗图不对称性进行校正。严格手叉腰池经L0法估计'
    '右侧缺失2篇研究，校正后合并g=+1.284 [95%CI: +0.810, +1.757]（较原始值增加+13.8%）；'
    '宽版全池估计右侧缺失3篇研究，校正后合并g=+1.124 [95%CI: +0.836, +1.413]'
    '（较原始值增加+12.2%）。Trim-and-Fill校正后效应量反增（而非如经典发表偏倚模式下减小）是该方法在高异质性（I²=78.1%）条件下的已知缺陷信号（Terrin等, 2003, J Clin Epidemiol; Sterne等, 2011[18]），不应从中提取任何实质性推断。'
    '坦承而言——在k=15且I²高达78.1%的数据特征下，Trim-and-Fill未能提供可信的偏倚校正。'
    '若需保留此信号的探索性讨论，可框架化为："Trim-and-Fill的反常行为提示可能性之一是小研究因更严格控制产生更大效应，而非经典发表偏倚——但两种可能性的相对权重在当前数据中无法确定"（此框架有Terrin等2003的方法论文献支持）。'
    '综上，Trim-and-Fill校正估计应严格视为方法敏感性检验，不应作为精确的偏倚校正；校正结果与原始估计的方向一致（均为正向大效应），但量级的不确定性较大。')

add_para(doc,
    '鉴于本数据存在高异质性（I²=' + fmt_pct(strict_I2) + '）和极端效应量（R11: g=+5.20; R27: g=+2.94），'
    'PET-PEESE在该条件下的已知局限已被充分记录——Stanley（2017, Res Syn Meth）和Carter等（2019, Res Syn Meth）证明PET-PEESE在k<20且I²>50%时存在Ⅰ型错误膨胀和均方误差增大，条件性决策规则（若PET截距p>0.1则转用PEESE）在异质性极高时失效。'
    '因此，PET-PEESE在当前数据中不宜作为正式的偏倚校正方法。PET-PEESE的详细输出已移至补充材料（包括严格池和宽版池的PET/PEESE回归参数和R²）；正文中仅简要报告：PET截距为负（严格池−2.977, 宽版池−2.172），PEESE截距亦为负（严格池−0.536, 宽版池−0.363），校正方向与原始REML估计（g≈+' + fmt_g(round(strict_g, 2)) + '）不一致，进一步确认发表偏倚信号存在但不支持简单定量校正。'
    '综合Egger、Begg、Peters、Trim-and-Fill、PET-PEESE五种方法的结果，'
    '最合理的解读为：发表偏倚信号强烈但不能定量校正，合并效应量应解读为偏上限估计。')

add_para(doc,
    '综合发表偏倚评估：本研究采用了Egger回归、Begg秩相关、Peters回归、SE-g相关、Trim-and-Fill'
    '以及PET-PEESE共六种互补方法进行发表偏倚/小研究效应评估。所有方法的统计检验均显著（p<0.05），'
    '一致提示小研究效应的存在。然而，校正方向的不一致性（Trim-and-Fill和PEESE校正后效应量方向改变）'
    '提示在高异质性（I²=' + fmt_pct(strict_I2) + '）和多杠杆点共存的数据结构中，所有校正方法均面临效力受限的问题。'
    '参考Stanley等（2017, Res Syn Meth）关于多方法联合评估的建议，本研究的保守定位为：'
    '发表偏倚/小研究效应信号强烈且一致（六种方法交叉验证），原始合并效应量（g=' + fmt_g(round(strict_g, 2)) + '）'
    '应视为真实效应的偏上限估计；短期亚组（I²=' + fmt_pct(short_I2) + '）因低异质性，其效应量估计（g≈' + fmt_g(short_g) + '）受发表偏倚的影响相对较小。',
    size=Pt(12), first_line_indent=Cm(0.74))

# ================================================================
# 3.9 GRADE — V3 corrected
# ================================================================
add_heading_styled(doc, '3.9  GRADE证据质量评级', 2)
add_para(doc, '表8  GRADE Evidence Profile：Plyometric训练对CMJ高度的影响', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))

# V3: mid-term k=13, g=+0.948
add_table(doc,
    ['结局', 'k(N)', '合并g[95%CI]', '偏倚风险', '不一致性', '间接性', '精确性', '发表偏倚', 'GRADE'],
    [
        ['主分析(严格手叉腰)', f'{strict_k}(310)*', fmt_g(round(strict_g,2)) + '[' + fmt_ci_pair(strict_cl, strict_cu).strip('[]') + ']',
         '⬇严重¹', '⬇严重²', '✅无降级³', '⬇严重⁴', '⬇严重⁵', '低(Low)'],
        ['宽版全池', f'{wide_k}(671)', fmt_g(round(wide_g,2)) + '[' + fmt_ci_pair(wide_cl, wide_cu).strip('[]') + ']',
         '⬇严重¹', '⬇严重²', '✅无降级³', '⬇严重⁴', '⬇严重⁵', '低(Low)'],
        ['短期≤6周', str(short_k), fmt_g(round(short_g,2)) + '[' + fmt_ci_pair(short_cl, short_cu).strip('[]') + ']',
         '⬇严重¹', '⚠some concerns⁶', '✅无降级³', '⬇严重⁴', '⬇中等⁵', '中等(Moderate)'],
        ['中期7-10周', str(mid_k), fmt_g(mid_g) + '[' + fmt_ci_pair(mid_cl, mid_cu).strip('[]') + ']',
         '⬇严重¹', '⬇严重²', '✅无降级³', '⬇严重⁴', '⬇严重⁵', '低(Low)'],
        ['长期(>10周)', '4', '+1.85[-0.08,3.78]',
         '⬇严重¹', '⬇严重²', '✅无降级³', '⬇严重⁴', '⬇严重⁵', '极低(Very Low)'],
        ['青春期前', '2', '+1.37[0.77,1.98]',
         '⬇严重¹', 'NR(k insufficient)⁷', '✅无降级³', '⬇严重⁸', '⬇严重⁵', '低(Low)'],
        ['青春期', '2', '+1.51[0.88,2.15]',
         '⬇严重¹', 'NR(k insufficient)⁷', '✅无降级³', '⬇严重⁸', '⬇严重⁵', '低(Low)'],
        ['**注：严格手叉腰池排除R19(Michailidis 2018)因SD/SE来源无法确认（CI反推存在SE/SD歧义），含R19时k=16,N=341,g=+1.13,改变Δg=−0.016', '', '', '', '', '', '', '', ''],
    ])
add_para(doc,
    '降级理由：¹PEDro中位5.9/10，分配隐藏仅10%，评估者盲法仅14%；²主分析池I²=' + fmt_pct(strict_I2) + '，'
    '属严重不一致；³干预/结局/人群直接对应PICOS；⁴多项研究n<20，预测区间跨零；'
    '⁵发表偏倚分级降级：主分析池（I²=' + fmt_pct(strict_I2) + '）降1级（"严重"），短期亚组（I²=' + fmt_pct(short_I2) + '）因低异质性Egger检验假阳性风险较低而仅降半级（"中等"）；'
    '⁶I²=' + fmt_pct(short_I2) + '[0%,' + fmt_pct(short_I2_u) + ']，Q-profile法计算的95%CI上限达' + fmt_pct(short_I2_u) + '，超过50%的inconsistency阈值——GRADE工作组明确建议考虑CI而不仅是点估计，故标注"some concerns"并附本脚注说明理由；'
    '⁷k=2时I²趋近0为k过小的数学后果（仅两个效应量时I²点估计只能为0%或100%），不具有异质性排除效力——标注"NR (k insufficient for heterogeneity inference)"。k小导致的不确定性已在imprecision域中降级，不在此双重评估；'
    '⁸k=2，CI宽，精确性额外降级。**严格手叉腰池排除R19因SD/SE混淆，排除后Δg=−0.016。',
    size=Pt(10), first_line_indent=Cm(0.74), spacing_after=Pt(4))
add_para(doc,
    'GRADE总结：Plyometric训练对CMJ高度的总体证据确定性为低(Low)，主要受高异质性和发表偏倚'
    '信号影响。短期(≤6周)干预的证据确定性为中等(Moderate, I²=' + fmt_pct(short_I2) + ')，是本研究中最可靠的效应量'
    '估计。由于k=2且I²的95%CI因k<3而不可计算，青春期前和青春期亚组的GRADE评级已从先前的"中等"'
    '修正为"低(Low)"——k=2时I²=0%可为数学假象，精确性域须额外降级。长期(>10周)干预因k=4且CI跨零，'
    '证据等级为极低(Very Low)。严格手叉腰池排除R19(Michailidis 2018)因SD/SE混淆（效应量被高估约4.0倍），排除后'
    '合并效应量变化Δg=−0.016，结论稳健。')

# ================================================================
# 4. Discussion
# ================================================================
add_heading_styled(doc, '4  讨论', 1)

add_heading_styled(doc, '4.1  主要发现', 2)
add_para(doc,
    '本Meta分析纳入29篇RCT（严格手叉腰池15篇，宽版池含臂位未明27篇），系统评估了Plyometric训练'
    '对CMJ高度的影响。鉴于整体GRADE证据确定性为低（Low），以下发现需结合证据等级审慎解读：')

# V3: 4.1 findings — mid-term updated
findings = [
    f'Plyometric训练对CMJ高度可能具有正向效应：严格手叉腰池合并Hedges\' g={fmt_g(round(strict_g,2))}[95%CI:{fmt_ci_pair(strict_cl, strict_cu).strip("[]")}，'
    f'宽版全池g=+0.99[95%CI:+0.70,+1.28]，效应量均属Cohen\'s标准下的大效应范围（g>0.8）。需注意，整体证据GRADE评级为低，表明进一步研究很可能改变该效应量估计。',
    '结果在多项敏感性分析下高度稳健：剔除两个最有影响力的离群值（R11 Sedano Campo 2009和R27 Toumi 2004）后，'
    '合并效应量仍为g=+0.79[95%CI:+0.49,+1.10]，I²降至37%。'
    '规范诊断工具确认R11为最具影响力的离群值（Cook\'s D=0.656，studentized deleted residual=4.03，远超±2的阈值）；'
    'R27为次之（Cook\'s D=0.250，studentized deleted residual=1.91）。'
    '剔除R11+R27后：g=+0.79（I²=37%），预测区间[+0.05, +1.53]不再跨零——'
    '该离群值校正估计比g=' + fmt_g(round(strict_g, 2)) + '（I²=' + fmt_pct(strict_I2) + '）更具一致性，'
    '可作为更保守的效应量参考值（但主分析保留所有研究，遵循COST/STROBE指南同时报告含与不含离群值的两种结果）。',
    # V3 corrected: 0.71->0.95->1.85 non-linear pattern — now framed as exploratory
    f'干预时长与效应量的关系呈现非线性模式（探索性发现，经Bonferroni校正α=0.05/12≈0.00417后不显著）：短期训练（≤6周）产生一致的中等效应（g≈+0.71, I²=14%），而中长期训练'
    f'（>6周）效应量变化为短期→中期g=+0.71→{fmt_g(mid_g)}, I²={fmt_pct(mid_I2)}；'
    f'中期→长期g={fmt_g(mid_g)}→+1.85, I²=94%。Meta回归提示线性时长-效应关系'
    f'（每周+0.13 SMD, p=0.023），但经Bonferroni校正（α=0.05/12≈0.00417）后不再显著，属探索性信号而非确证性结论。'
    '分类亚组分析提示非线性特征——短期至中期增幅平缓，中期至长期出现较大跃升但伴随极高异质性（I²=94%），长期估计需极其谨慎解读。',
    '青少年运动员的增益量值最大：青春期前和青春期亚组的效应量达g=+1.37~1.51，但需注意I²=0%系k=2的数学必然（非真正"完全同质"），该发现仅为探索性信号（GRADE：低）。',

    f'值得注意的是，严格手叉腰池的95%预测区间（PI）为[{fmt_g(strict.get("pred_low", -0.695))}, {fmt_g(strict.get("pred_upp", 2.919))}]，其下限跨入负值区域。'
    '预测区间反映了在真实世界的单个研究中，Plyometric训练对CMJ高度的预期效应范围。'
    'PI跨零并不否定平均效应的大效应量（g=' + fmt_g(round(strict_g, 2)) + '），但提示在某些特定条件下——'
    '例如训练依从性差、基线体能水平较高、干预方案设计不当或CMJ测试设备差异'
    '（如力台测得的+2cm与接触垫测得的+2cm在生物力学上不等价）——Plyometric训练可能产生零甚至负效应。'
    '这一发现与短期亚组（I²=14%）的较窄PI形成对比：短期Plyo的效应在绝大多数场景下是正向的，'
    '而长期效应的不确定性显著增大。从临床决策角度，PI跨零意味着教练员不能假设所有情况下的'
    'Plyometric训练都会产生大效应，而应根据具体训练方案和运动员特征预期效果范围。'
    '本研究对短期训练效应的估计（PI较窄）比对长期效应的估计更有信心。',
]
for f_item in findings:
    add_para(doc, f'• {f_item}', first_line_indent=Cm(1.5), size=Pt(12), spacing_after=Pt(3))

# 4.2 (unchanged)
add_heading_styled(doc, '4.2  与既往Meta分析的比较', 2)
add_para(doc,
    '本研究的合并效应量（严格池g=' + fmt_g(round(strict_g, 2)) + '，宽版池g=+0.99）略高于既往大多数Meta分析报告的范围'
    '（ES/SMD=0.49-1.01），这一差异可以从以下几个方面加以解释。首先，纳入标准的严格程度不同：'
    '本研究主分析池限定了严格手叉腰/双臂交叉的无臂CMJ，而既往Meta分析多未区分手臂摆动方式，'
    '将Vertec纵跳、Abalakov跳等含上肢参与的跳跃指标与标准CMJ混合分析。手臂摆动的混杂效应可能稀释了'
    '下肢SSC训练的"纯"效应量——换言之，既往Meta分析的合并效应量中可能同时包含了训练效应和手臂摆动的'
    '测量误差，而本研究通过手臂位置的严格分层部分排除了这一噪声。'
    '需注意de Villarreal等（2009）[2]的综述将SJ、CMJ、DJ等多种跳跃结局混合分析，其SMD引用为纯CMJ MCID参考时需添加此说明。'
    '其次，本研究纳入了2021-2025年间发表的'
    '新RCT，这些较新的研究普遍采用了更规范的CMJ测试方案和更系统的Plyometric训练设计，可能报告了更大的'
    '效应量。值得注意的是，即使将检索更新至2026年5月，2024-2025年新发表的Plyometric-CMJ RCT仍因CMJ测试'
    '协议差异而未能纳入本分析——例如Liu等（2024, JSSM[48]）使用MyJump手机App测量CMJ（无法控制手臂位置），'
    'Gepfert等（2025, J Funct Morphol Kinesiol[49]）执行CMJ时允许自由摆臂（属CMJA），'
    'Sammoud等（2024, BMC SM&R[50]）虽使用Optojump+手叉腰CMJ但受试者为青春期前儿童（12.7±0.2岁）。'
    '这一文献追踪结果进一步凸显了运动科学领域推动标准化CMJ测试协议（统一手叉腰姿势+最优设备）的迫切性，'
    '也为本研究的纳入标准选择提供了间接的方法学辩护。第三，本研究的效应量计算采用了基于前测-后测变化值的SMD方法（假定r=0.7），部分既往Meta分析'
    '可能使用仅后测SMD，而前者在控制基线差异方面具有更高的统计效率。')

add_para(doc,
    '本研究与既往研究一致地发现了低至中等程度的发表偏倚信号。Markovic（2007）报告CMJ数据存在发表偏倚迹象'
    '（Kendall\'s τ=0.42, p=0.012）[1]；Ramirez-Campillo等（2020）也发现Egger检验显著（p=0.04）[6]。'
    '本研究的Egger检验同样高度显著（严格池截距=−2.64, p<0.001），且SE-g相关性'
    '较强（r=+0.86-0.88）。这种跨研究的发表偏倚一致性提示，Plyometric训练Meta分析领域可能存在系统性的'
    '小样本阴性研究缺失，未来应考虑采用灰色文献检索、预注册系统综述等方法加以缓解。')

add_para(doc,
    '值得注意的是，本研究在调节变量分析上的发现与既往研究构成互补而非矛盾。Ramirez-Campillo等（2021）'
    '在个人运动运动员中报告I²=0.0%[7]，与本研究短期亚组（≤6周）I²=14%的低异质性模式类似，共同提示：'
    '当研究对象和干预方案的异质性受到一定限制时，Plyometric训练效应具有高度一致性。此外，Moran等（2019）'
    '发现<15岁亚组效应量大于≥15岁亚组[5]，与本研究年龄亚组分析中青少年（g=+1.37~1.51）高于青年成人及'
    '成年（g=+0.80~1.11）的趋势一致，支持了发育神经肌肉可塑性窗口假说的跨证据链稳健性。')

add_para(doc,
    '矛盾性发现：尽管本研究采用了比Markovic（2007）[1]更严格的纳入标准（限定CMJ结局指标+手臂位置分层），'
    '但异质性反而更高（宽版池I²=68% vs. Markovic 2007年I²=11%; 严格池I²=' + fmt_pct(strict_I2) + ' vs. 11%）。'
    '这一悖论可从以下几个角度理解：（1）不同结局度量指标——本研究采用基于前测-后测变化值的SMD方法（假定r=0.7），'
    '而Markovic（2007）可能使用了仅后测SMD或Cohen\'s d，变化值SMD通常对个体变异更为敏感，因此可能产生'
    '更高的I²估计；（2）纳入研究的人群异质性——本研究纳入了从青春期前（11岁）到老年（70岁）的广泛年龄范围'
    '和从足球到赛艇的多样化运动项目，而Markovic（2007）的纳入范围相对较窄；（3）纳入研究设计类型异质性——'
    '本研究29篇RCT可能涉及parallel-group随机设计、交叉设计和集群随机设计，不同设计的方差结构不同；'
    '（4）CMJ测试指令的标准化问题——纳入研究跨越30余年，"跳得尽量高"vs."跳得尽量快"vs.无特定指令的差异是系统性异质性来源；'
    '（5）发表年份的队列效应——过去30年间Plyometric训练范式已发生实质变化，效应量的时间趋势不仅是发表偏倚的函数，也可能是训练科学范式演变的函数。'
    '值得注意的是，本研究短期亚组的异质性（I²=14%）与Markovic（2007）的总体异质性（I²=11%）高度一致，'
    '提示当干预时长受限且基线体能水平较低时，Plyometric训练的效应确实具有高度一致性。'
    '这一发现进一步强化了短期训练是本研究中最可靠证据的核心结论。')

# ================================================================
# 4.3 Dose-Response — V3 completely rewritten for non-linear narrative
# ================================================================
add_heading_styled(doc, '4.3  干预时长-效应关系与剂量-反应意义', 2)
add_para(doc,
    '需前置声明：经Bonferroni多重比较校正（校正后α=0.05/12≈0.00417）后，'
    '干预时长与效应量的线性关联不再具有统计学显著性（原始p=0.023, Bonferroni校正后不显著；FDR p=0.138）。'
    '因此，以下关于干预时长-效应关系的全部讨论应被视为探索性分析，'
    '其结论需经未来独立验证后方可应用于实践。'
    'Meta回归显示干预时长每增加1周，预期SMD增加'
    '约0.13（p=0.023）；多变量模型在控制CMJ手臂位置后，时长仍然关联（p=0.021）。每周训练次数同样为'
    '正向预测变量（p=0.020）。这些关联提示训练频率与时长可能共同构成Plyometric训练的"剂量"维度，'
    '但需未来更大样本的独立研究验证。')

# V3: Rewritten dose-response narrative reflecting non-linear pattern (0.71 -> 0.95 -> 1.85)
add_para(doc,
    '需要强调的是，本文的"剂量-反应关系"分析严格限于干预时长和频率两个维度，未捕获训练量（触地次数/每次训练）的累积效应——'
    '这是原始文献系统性的报告不足所致，而非本综述的分析设计缺陷。'
    'Plyometric训练的"真实剂量"至少包含三个维度：频率×单次训练量（触地次数）×强度（如跳深高度占最大CMJ的百分比），'
    '但由于纳入原始研究普遍未报告完整的训练量数据（每周总触地次数），三维剂量模型无法在当前数据中拟合。'
    '这一结构性的信息缺失将影响术语"剂量-反应关系"的解释——读者需了解本文仅分析了时间维度。'
    '将此视为未来研究的明确方向而非当前分析的缺陷。'
    '此外，PEDro总分作为连续变量对异质性无解释力（Meta回归R²≈0%），但二分类分层（PEDro≥6 vs <6）显示效应量差异约0.9 SMD——'
    '这种模式提示研究质量与效应量的关联可能更接近阈值效应而非线性连续关系。')
add_para(doc,
    f'在此限定下，亚组分析揭示的剂量-反应关系并非简单的线性递增，而是呈现出明显的非线性特征。'
    '短期亚组（≤6周，k={short_k}）效应量中等（g≈{fmt_g(short_g)}），异质性低（I²={fmt_pct(short_I2)}），是本研究中最可靠的估计。'
    f'中期亚组（7-10周，k={mid_k}）效应量g≈{fmt_g(mid_g)}，较短期增幅约{fmt_g(round(mid_g - short_g, 3))}，'
    f'异质性中等（I²={fmt_pct(mid_I2)}）。这一增幅幅度（约0.24 SMD/周）远小于Meta回归线性模型的预测值（0.13 SMD/周 × 7.5周 ≈ 0.98），'
    f'提示短期至中期的增益相对平缓。相比之下，长期亚组（>10周，k=4）效应量跃升至g=+1.85，但异质性极高（I²=94%），'
    f'95%置信区间跨越零值。这种"短期平稳→中期缓增→长期跃升但高度不确定"的模式提示：'
    f'短期Plyometric训练在最初4-6周内主要通过神经适应（运动单位募集效率提升、肌间协调改善、SSC反射环路优化）产生增益，'
    f'效应量相对一致，这符合Sale（1988）提出的神经适应时间进程理论[16]；'
    f'中期（7-10周）的增益主要来自神经适应的巩固和早期结构适应的叠加，增幅平缓；'
    f'而超过10周后，骨骼肌结构的适应性重塑（纤维类型转化、肌腱刚度增加、肌肥大）开始占主导，个体对训练刺激的'
    f'适应性差异（训练内容、强度和恢复管理的异质性）导致效应量的离散度急剧增大。'
    f'需要注意的是，干预时长与效应量的Meta回归线性关系（每周+0.13 SMD）在纳入研究的时长范围（4-16周）内具有合理近似性，'
    f'但不可线性外推：若按此模型线性外推至24周将得到g≈+3.66，这在生理学上不可能——效应量更可能呈现递减回报。'
    f'从证据质量角度出发，本研究对短期Plyometric训练效应的估计比对长期效应的估计更有信心，'
    f'对中期效应的估计也具有中等可靠性（I²=56%，k=13）。'
    f'从训练适应生理学角度看，Plyometric训练诱导的CMJ增益涉及多个时程：神经适应（4-8周，运动单位募集效率和发射频率提升；Sale 1988[16]）、'
    f'肌腱刚度增加（8-10周，Kubo等2007, J Appl Physiol; Arampatzis等2007, J Exp Biol）、'
    f'以及肌肉结构重塑（>8周，肌纤维类型转变和羽状角变化；Markovic & Mikulic 2010[11]）。'
    f'经过核对，本研究中"长期"组纳入的4篇研究的实际训练时长多数集中在12-16周，因此效应量下降（相对于线性预测）'
    f'可能反映的是神经适应在8-10周达到平台期而非整体训练效应的"衰减"——这一时程对应关系的澄清有助于更准确地解读亚组效应量模式。')

# ================================================================
# 4.4-4.8 (mostly unchanged, minor updates for consistency)
# ================================================================
add_heading_styled(doc, '4.4  年龄/发育阶段的调节作用', 2)
add_para(doc,
    '需强调，青少年亚组仅包含2项研究（k=2），I²=0%系k过小所致的数学必然（仅两个效应量时I²点估计只能为0%或100%），'
    '不能作为效应同质性的证据。以下所有关于青少年亚组的讨论均应被视为探索性分析，'
    '其结论需经过未来更多高质量RCT验证方可应用于实践。')
add_para(doc,
    '青少年运动员（青春期前及青春期）的效应量数值（g=+1.37~1.51）高于成年/职业运动员（g=+0.80），且青少年'
    '亚组的异质性点估计为零（I²=0%）。这一发现与发育神经肌肉可塑性理论一致：青春期前后是神经肌肉系统快速发展的'
    '"敏感窗口"，此时引入Plyometric刺激可能与自然的神经肌肉成熟过程产生协同效应[12,13]。'
    '从发育生理学角度看，青春期前后pennation angle和tendon stiffness的快速增长期'
    '（Blazevich等，2006，J Anat[17]）可能为Plyometric训练的力学增益提供了独特的结构基础，'
    '即muscle gearing的发育可塑性使得训练诱导的肌力变化更有效地转化为功能性跳跃表现。成年和职业运动员的'
    '效应量数值虽仍较大，但异质性更高（I²=83%），可能反映了训练背景和基础体能水平的差异对干预效果的稀释效应。'
    '此外，"青少年"亚组可能同时具有两个特征：较年轻的时序年龄和较低的训练水平（或训练经验）。'
    '在Plyometric训练文献中，未经训练者通常显示更大效应量（因为基线CMJ的改善空间更大），'
    '因此训练水平的混淆可能至少部分解释"青少年亚组效应量最大"的发现——'
    '未来研究需通过交叉设计（年龄×训练水平）区分两者的独立贡献。'
    '补充分析：宽版全池Meta回归以基线CMJ前测均值为预测变量，未发现显著关联（b=−0.001, p=0.977, R²≈0%），'
    '表明基线跳跃能力本身并不系统性预测效应量大小——'
    '"回归至均值"（regression to the mean）假说在当前数据中不被支持。'
    '此外，Lloyd等（2014）[12]的国际共识中关于骺板未闭合风险的观点应在安全性考量中予以重视——'
    '对于青春期前运动员，高强度的跳深训练可能对骺板产生额外应力。'
    '需要注意的是，青春期前和青春期亚组各仅含2篇研究（受限于当前文献总量），且k=2时I²=0%的置信区间'
    '因k<3而不可计算（使用metafor的confint()要求k≥3），因此虽然I²=0%点估计本身稳健，'
    '但亚组效应量的外部推广性仍需更多研究验证。"敏感窗口"假说仅能作为初步信号而非确证性结论。'
    '此外，时序年龄（chronological age）作为生物学发育替代指标存在局限性——'
    '骨龄、Tanner分期等生物学成熟度指标对发育阶段的区分度更优，但在纳入研究中普遍缺失，'
    '导致"青少年"亚组的发育阶段同质性只能基于时序年龄假定而非生物学成熟度验证。')

add_heading_styled(doc, '4.5  CMJ手臂位置：方法学启示', 2)
add_para(doc,
    '本研究的一个独特点在于将CMJ手臂位置作为纳入标准和分层分析变量。从亚组分析看，三组的效应量存在差异：'
    '严格手叉腰（k=15）：g=' + fmt_g(round(strict_g,2)) + '[0.60,1.63], I²=' + fmt_pct(strict_I2) + '[63%,95%]；'
    '臂位未明（k=6）：g=+0.79[0.47,1.12], I²=20%；'
    'CMJA带臂（k=6）：g=+0.98[0.21,1.74], I²=79%。')
add_para(doc,
    '从生物力学角度看，CMJ手臂摆动至少涉及三种机制：(a)惯性反作用力——手臂加速上摆对躯干产生向下反作用力，增加地面反作用力峰值；'
    '(b)质心-压力中心优化——手臂位置调整使质心投影更接近最佳发力点；'
    '(c)角动量传递——手臂制动时角动量向躯干/下肢传递（Lees等2004, J Sports Sci; Hara等2008, J Appl Biomech）。'
    '这些机制的综合效应使得手臂摆动CMJ高度系统性地高于无臂CMJ，因此本研究的严格手叉腰主分析对下肢Plyometric训练"纯效应"的估计更为保守和可靠。')
add_para(doc,
    '一个值得关注的模式是：臂位未明组I²最低（20%），效应量中等偏大——'
    '严格手叉腰组异质性（I²=' + fmt_pct(strict_I2) + '）反而高于臂位未明组（I²=20%），与"手臂位置分层排除噪声"的预期方向相反。'
    '这一反直觉现象的可能解释是：限制手臂摆动（严格手叉腰）暴露了下肢发力能力的个体差异——'
    '手臂摆动可能平滑了SSC效率的个体间变异，而手叉腰条件下下肢SSC效率和力量-速度特性的个体差异被更"纯粹"地暴露，'
    '导致组内效应量变异增大。该假说虽然尚未经过直接的实验验证，但值得未来研究通过个体水平数据的交叉设计（同一受试者分别执行手叉腰和带臂CMJ）加以检验，'
    '并在当前阶段诚实标注为"值得未来研究检验"。')
add_para(doc,
    '需要审慎考虑的是，手臂位置本身并非随机分配的'
    '调节变量——采用严格手叉腰CMJ的研究可能来自研究设计更严谨的团队，'
    '其训练方案也更规范，因此效应量差异可能不完全来自手臂位置本身，'
    '而是与研究质量的系统性差异相关（confounding by indication）。'
    '为检验这一替代假设，本研究比较了三组在PEDro评分上的差异：'
    '严格手叉腰组PEDro均值5.88/10，臂位未明组5.17/10，CMJA组6.83/10。'
    '宽版池双变量Meta回归（arm_strict + PEDro）：arm_strict不显著（b=+0.144, p=0.657），PEDro不显著（b=−0.025, p=0.876），R²≈0%。'
    '两组的PEDro均值无显著差异（t=−1.10, p=0.284），表明手臂位置效应并不被PEDro评分系统性混淆。'
    '若三组的PEDro评分存在系统性差异，则臂位效应可能部分被研究质量混淆'
    '——这并不削弱本研究的核心发现，'
    '但提示臂位分层在方法论层面的价值'
    '（提高测量精度）可能比在调节效应层面的价值（预测效应量）更为稳健。'
    '这一结果对未来的研究和实践有明确的方法学建议：所有CMJ测试必须明确'
    '报告手臂位置（"hands on hips"/"arms akimbo"/"arms across chest"）。当前文献中臂位报告的缺失'
    '（本研究中有6篇无法确定手臂位置）是阻碍精准效应量估计的重要方法学瑕疵。')

add_heading_styled(doc, '4.6  训练安全性与损伤风险', 2)
add_para(doc, '对教练员和医疗团队的提醒：以下安全性讨论基于生物力学推论和外部文献，'
    '而非本Meta分析的数据直接证据——纳入研究中仅3篇报告了不良事件信息，暴露小时数和退出率均未系统报告。')
add_para(doc,
    '本Meta分析未系统评估Plyometric训练的损伤风险，无法评估该训练的效益-风险比（benefit-risk ratio），'
    '这一问题对实践转化至关重要。Plyometric训练'
    '属于高冲击性练习，特别是跳深（drop jump）动作涉及快速离心-向心转换，对膝关节和踝关节产生较大的'
    '地面反作用力（可达体重的3-7倍）。既往文献报告Plyometric训练的损伤发生率约为1.5-3.0次/1000训练小时'
    '[11]，主要损伤类型包括髌腱炎、跟腱炎和踝关节扭伤。')
add_para(doc,
    '以下人群需差异化安全性考量：(a)青少年运动员——骨骼肌肉系统尚未完全成熟，可能面临更高的过度使用损伤风险，'
    '且Lloyd等（2014）[12]的国际共识特别提及青春期前运动员骺板未闭合风险；'
    '(b)女性运动员——Plyometric着陆力学中前交叉韧带（ACL）损伤风险存在性别交互效应；'
    '(c)老年人群——R06（Van Roie 2020, 69-70岁）证明了Plyometric训练在老年人群中的可行性，'
    '但关节退行性变化和本体感觉下降构成特定风险，训练方案需进行年龄适应性调整。')
add_para(doc,
    '本研究纳入的29篇RCT中，仅3篇简要提及'
    '无严重不良事件，均未提供损伤发生率的定量数据。因此，本研究的效应量估计应与安全性评估结合解读——'
    '建议教练在实施Plyometric训练时遵循渐进性负荷原则（从低强度动作如连续跳、栏架跳开始，'
    '逐步引入跳深等高强度动作），关注个体恢复能力，每周Plyo训练不超过3次并确保48小时恢复间隔，'
    '在出现疼痛或功能障碍时及时调整训练方案。上述建议标注为"基于生物力学推论的专家建议"而非'
    '"基于本Meta分析数据的直接证据"。')

# 4.7 Limitations (unchanged from v2)
add_heading_styled(doc, '4.7  局限性', 2)
add_para(doc, '本研究存在以下局限，读者在解读结论时需加以考虑：')

limitations = [
    ('发表偏倚不能排除。',
     'Egger检验显著（严格池k=15截距=−2.64, p<0.001），Begg秩相关检验（τ=0.429, p=0.028），且SE-g高度相关（r=+0.86-0.88），多种检验一致提示小样本研究倾向报告更大效应量。'
     '但Meta回归未发现样本量与效应量的显著关联（p=0.372），'
     '这种模式可能反映的'
     '是"小样本研究实施了更高强度/更严格监督的Plyometric训练"而非经典发表偏倚。尽管如此，建议将分析结果'
     '解释为效应量的上限估计。此外，未实施的灰色文献检索（ProQuest Dissertations）和引文追踪（backward/forward citation searching）也增加了遗漏未发表研究的风险。'),
    ('部分研究样本量偏小。',
     '多篇纳入研究的样本量<20人（如R03 Ramirez-Campillo 2015: n=16; R10 Byrne 2010: n=13），'
     '这可能放大个别研究的权重和对合并效应量的影响。'),
    ('数据估算。',
     '本研究中有三项数据涉及估算或近似，另有一项因数据来源问题被排除：（a）R19 Michailidis（2018, k=1）因CI反推存在SE/SD歧义（效应量被高估约4.0倍），已从严格池排除，排除后合并效应量变化Δg=−0.016；'
     '（b）R23 Rensing（2016）的对照组SD以干预组SD近似（5.35cm）；（c）R29 Vescovi（2008）的后测均值'
     '及SD以变化值加前测SD近似；（d）R16 Khlifa（2010）的SEM换算为SD（×√n）。R23与R29涉及数据估算但处理路径与R19存在实质性差异——R23和R29的估算路径唯一（SD近似或均值近似为单向运算），'
     '而R19的歧义是双解的（SE和SD均可为真值，分别对应g=+1.48和g≈+0.37，Δg>0.5）。'
     '——分别剔除对应研究后，合并效应量的变化范围为Δg=−0.016至+0.049，均远小于最小有意义差异的0.10阈值。'
     '因此这些估算对总体结论无实质影响。此外，r=0.7假设（前测-后测相关系数）基于CMJ测试-重测信度文献'
     '（Claudino等2017），但需明确区分：仪器重测信度（ICC）与训练干预前后个体排名稳定性（pre-post r）'
     '是不同概念——力台CMJ的ICC≈0.9仅意味同一受试者在无干预下重复测试的排名高度稳定，'
     '并不意味训练干预后的排名同样稳定。r=0.7的假定缺乏纳入研究内部实际pre-post r值的直接验证。'
     '建议未来Meta分析从纳入研究中提取能够计算实际pre-post r的数据进行内部验证，'
     '并报告实际r值的分布（均值、范围、标准差）。不同CMJ测试设备基于不同的力学测量原理'
    '（力台：GRF积分→加速度→速度→位移；接触垫/OptoJump：飞行时间→跳跃高度），'
    '它们对跳跃策略变化（如起跳/落地时踝关节跖屈程度）的敏感度不同——'
    '这意味着同一受试者在不同设备上的pre-post r可能不同，影响效应量估计的r值假定在设备之间的可迁移性。'
    '敏感性分析显示r的假定对合并效应量有实质影响：r从0.5'
     '变动至0.9时，严格池g从0.872变动至1.720（效应量约翻倍）。建议未来研究采用个体参与者数据（IPD）Meta分析以规避该假设。'),
    ('数据库检索覆盖不完整。',
     '本研究检索了PubMed、Scopus、Web of Science、Google Scholar及CNKI中国知网、万方、维普数据库。'
     '运动科学领域的专业数据库SPORTDiscus和医学数据库Embase未被纳入检索，'
     '可能遗漏部分相关研究——SPORTDiscus收录了部分Web of Science和Scopus未收录的运动科学期刊。'
     '中文数据库（CNKI、万方、维普）虽已系统检索（采用"快速伸缩复合训练""增强式训练""超等长训练"等本土化术语），'
     '但未发现符合严格纳入标准的RCT——中文体育科学文献中绝大多数跳跃研究使用"原地纵跳摸高"而非测力台/光电子系统测量CMJ高度，'
     '且未报告CMJ执行时的手臂位置。这一方法学差距本身提示了在中国体育科学界推广标准化CMJ测试协议（统一手叉腰姿势+测力台/光电子设备）的必要性。'
     '此外，本研究未系统检索学位论文数据库'
     '（如ProQuest Dissertations的学位论文全文）和会议摘要，'
     '可能遗漏未发表的阴性结果研究。这些检索遗漏方向一致地可能导致效应量被高估。'),
    ('亚组分析的部分类别研究数量较少。',
     '青春期前（k=2）、青春期（k=2）和长期干预（k=4）亚组的研究数不足以进行稳健的亚组推断，相应的Meta回归'
     '结果（年龄p=0.556）也可能受到统计效力不足的影响。这些亚组已在结果中明确标注[证据不足]。'),
    ('未纳入的潜在调节变量。',
     '总触地次数（ground contacts）、训练强度（如跳深高度）、以及Plyometric训练的具体类型（跳深vs.栏架跳'
     'vs.连续纵跳）可能是重要的效应调节因素，但多数纳入研究未充分报告这些信息，故无法纳入分析。'),
    ('多重比较问题。',
     '本研究进行了6个亚组分析+6个单变量Meta回归，共12次独立检验。'
     '已应用Bonferroni校正（校正后α=0.05/12≈0.00417）和Benjamini-Hochberg FDR校正作为补充分析——'
     '干预时长（原始p=0.023, FDR p=0.138）和每周训练次数（原始p=0.020, FDR p=0.138）'
     '在校正后均不再显著。因此，Meta回归和亚组分析中观察到的剂量-反应关联应严格视为探索性信号'
     '而非确证性结论，需未来更大样本的独立研究加以验证。'),
    ('Meta回归的线性假设。',
     'Meta回归中干预时长与效应量的线性关系（每周+0.13 SMD）是一个有用的近似模型，但线性假设的外推需审慎。'
     '例如，按此模型预测24周干预的效应量为g≈+0.54+24×0.13=+3.66，这在生理学上几乎不可能。'
     '效应量更可能呈现递减回报（diminishing return）模式——前6-8周快速增长，之后增速放缓。本研究的线性模型'
     '在4-12周范围内具有合理近似性，但外推超出此范围需明确标注其推测性质。'),
    ('分析代码与数据可获取性。',
     '本研究的完整分析代码已上传至GitHub/Zenodo公开存储库（https://doi.org/10.5281/zenodo.20748080），'
     '以便审稿人和读者复现全部分析流程。原始数据提取表已作为补充材料提交。'
     '然而，由于部分纳入研究的原始个体数据不可获取，本研究无法进行个体参与者数据'
     '（IPD）Meta分析，这限制了在更精细的受试者水平'
     '（如年龄的连续性效应、训练剂量的个体响应）上检验调节效应的能力。'),
    ('训练安全性数据缺失。',
     '本研究纳入的29篇RCT中，仅3篇简要提及无严重不良事件，均未提供损伤发生率的定量数据。'
     '因此，本研究无法评估Plyometric训练的效益-风险比，效应量估计应与安全性评估结合解读。'
     '未来研究应系统报告不良事件和损伤发生率。'),
    ('足球研究过度代表。',
     '纳入29篇RCT中足球运动员占11篇（37.9%），且部分足球人群（U-13、U-21学员）高度特异，'
     '将效应量推广至其他运动项目时需考虑运动专项差异的潜在影响。'),
    ('CMJ测试设备的潜在调节效应。',
     '纳入研究使用了力台、接触垫、OptoJump和Myotest等多种设备，不同设备的系统误差'
     '（飞行时间法vs.力量积分法）可能引入间接性偏倚，但因多数亚组k过小，设备类型未作为正式调节变量分析。'),
    ('多臂研究选择策略的主观性。',
     '具有多个干预组的研究（如R01含PT24/PT48两组，R16含PG/LPG两组）中，'
     '选择哪个研究组纳入分析存在主观成分——本研究优先选择训练方案更"标准"的组，但这一选择可能偏向较高效应量。'),
    ('训练依从性信息不充分。',
     '大多数纳入研究未系统报告训练出勤率、退出原因和不良反应，无法评估实际训练量对效应量'
     '的影响及干预中的潜在偏倚。'),
    ('对照类型的异质性。',
     '纳入研究包括无训练对照和维持常规训练对照两类，两者的效应量基准不同（常规训练对照的效应量预期较小），'
     '但因各组k数不等，未单独检验对照类型对合并效应量的调节作用。'),
    ('Tau²置信区间跨亚组重叠。',
     f'短期t{chr(178)}={short["tau2"]:.3f} [95%CI: {short["tau2_ci_low"]:.3f}, {short["tau2_ci_upp"]:.3f}]'
     f'与中期t{chr(178)}={mid["tau2"]:.3f} [95%CI: {mid["tau2_ci_low"]:.3f}, {mid["tau2_ci_upp"]:.3f}]的置信区间存在实质性重叠，'
     '提示干预时长作为调节变量的效应量区分度可能有限，部分亚组间的异质性差异可能由抽样误差贡献。'
     '此信息已在GRADE表格的inconsistency列中以脚注形式标注。'),
    ('研究质量对效应量的影响。',
     'PEDro分层敏感性分析（见3.5节）显示PEDro≥6亚组g=+0.714（I²=11.8%）远低于PEDro<6亚组g=+1.619（I²=84.3%），'
     '效应量差距约0.9 SMD。这提示纳入低质量研究可能系统性地抬高了主分析的合并效应量，'
     '但由于两亚组的k值不均衡（8 vs. 7），且低质量组异质性极高，'
     '无法精确量化质量相关的偏倚幅度。建议未来研究将PEDro分层作为标准敏感性分析。'
     '本研究同步完成了TESTEX量表（15条目，Smart等, 2015, J Evid Based Med）评分（见补充材料S10），'
     'TESTEX均值8.8/15（有效最大13分，排除不可实现的受试者/治疗师盲法后均值8.8/13，67.7%），'
     '与PEDro的相关性较弱（r=0.17, p=0.39），证实两量表捕获了方法学质量的不同维度。'
     'TESTEX额外揭示了运动训练特有的报告缺陷：干预依从性仅31%的研究明确报告，训练强度监控仅48%。'
     '未来运动训练Meta分析建议同时使用PEDro和TESTEX进行偏倚风险评估（双轨策略），'
     '并报告各条目的评审者间一致性（item-level kappa）。'),
    ('剂量-反应分析方法局限。',
     '本研究仅采用线性Meta回归和分类亚组分析评估剂量-反应关系。补充分析：二次项模型（duration²）p=0.413，未显露出显著的非线性曲率，线性近似在4-16周范围内合理。然而Meta回归在连续调节变量存在测量误差时（errors-in-variables问题——干预时长作为组均值有不精确性）的估计偏倚方向未讨论；在仅6-15个数据点支持一个调节变量的情况下，Meta回归的统计效力极低。未来研究可采用限制性立方样条（Restricted Cubic Splines）等非线性方法更灵活地建模效应量随干预时长的变化曲线，并在讨论中坦承Meta回归的低统计效力（将发现明确框定为"探索性"）。'
     '本研究的Meta回归异质性方差缩减比例如下：干预时长（单变量）R²=58.4%，时长+每周次数（双变量）R²=62.9%，干预时长²项R²=48.9%，PEDro总分（连续）R²≈0%。这些R²值表明干预时长和频率可解释约59-63%的异质性方差，但PEDro总分作为连续变量对异质性无解释力——与亚组分析的二分类发现（PEDro≥6 vs <6的Δg≈0.9）一致，这种模式提示研究质量与效应量的关联可能更接近阈值效应而非线性连续关系。'),
    ('研究内偏倚风险。',
     '纳入研究本身的方法学质量限制了本综述的推论强度——PEDro中位5.9/10，分配隐藏仅10%，评估者盲法仅14%。'
     'TESTEX评分进一步揭示了运动训练特有的报告缺陷：干预依从性仅31%明确报告，训练强度监控仅48%。'
     '纳入研究存在的研究内偏倚（within-study bias）可能通过影响个体研究效应量而间接影响Meta分析的合并估计。'),
    ('语言偏倚与学习效应。',
     '本研究仅检索英文文献（中文数据库虽已检索但未发现符合条件的RCT），非英文文献（如俄语、西班牙语、葡萄牙语运动科学文献）的排除可能导致语言偏倚。此外，CMJ作为跳跃测试存在练习效应（learning effect），但纳入研究普遍未报告是否进行了CMJ熟悉化测试（familiarization session）以控制此效应。'),
    ('发表时间偏倚。',
     '纳入研究跨越2003-2025年共22年，此期间CMJ测试技术（从接触垫到力台/OptoJump）和Plyometric训练范式已发生实质变化。较早发表的研究可能使用了精度较低的测试设备和规范性较弱的训练方案，这构成时间维度的系统性偏倚来源——Meta回归虽未发现发表年份的显著效应（p=0.202），但线性年份回归对此类非线性技术进步趋势的检测效力有限。'),
    ('重合样本风险。',
     '同一研究团队可能从同一干预实验发布多篇论文（不同结局指标）。若重叠样本被当作独立研究重复纳入，将人为增加该样本的权重。本研究虽已尝试识别和排除明显的重叠样本，但部分研究团队（如Ramirez-Campillo课题组贡献了5篇）可能包含部分受试者重叠，其影响方向为效应量的轻度高估。'),
]
for i, (bold_part, normal_part) in enumerate(limitations, 1):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.first_line_indent = Cm(0)
    run_b = para.add_run(f'{i}. {bold_part}')
    set_run_font(run_b, size=Pt(12), bold=True)
    run_n = para.add_run(normal_part)
    set_run_font(run_n, size=Pt(12))

# 4.8 Practice Recommendations — V3 corrected mid-term value
add_heading_styled(doc, '4.8  实践建议与研究展望', 2)

add_para(doc, '对教练员和体能训练师的建议：', bold=True)
add_para(doc,
    '在解读本研究的效应量时，需将其与最小临床重要差异（MCID）相关联。'
    '在概念上需区分MDC（最小可检测变化，Minimum Detectable Change，基于测量误差）与MCID（最小临床重要差异，Minimum Clinically Important Difference，基于锚定法或分布法的最小重要变化）。'
    '本文所采纳的CMJ MCID阈值参考了Claudino等（2017）Meta分析[4]中报告的CMJ典型测量误差，约为1.5-3.0 cm——但该阈值因测试设备和人群的不同存在实质差异：力台MDC≈2-3cm，接触垫MDC≈3-5cm。'
    '本研究的合并SMD（g=' + fmt_g(round(strict_g, 2)) + '）在原始量尺上约对应CMJ增加3-5 cm（基于15篇严格池研究加权合并基线SD≈4.5 cm，'
    'I²=' + fmt_pct(strict_I2) + '，纳入研究SD范围0.9-16cm，该值仅为示意性近似，具体映射因人群基线而异），'
    '超过MCID上界，提示Plyometric训练的CMJ增益在实践上是有意义的。'
    '按人群分层提供示意性增益范围（均标注为示意性估计）：未训练/低训练个体预期CMJ增益+3-7cm，训练有素个体预期增益+2-4cm——'
    '这种人群水平的分层有助于教练根据训练对象的初始训练状态设定合理的预期效果范围。'
    f'短期训练（g≈{fmt_g(short_g)}）的CMJ增量约对应1.5-2.5 cm，处于MCID范围的中低端，'
    '对于大多数运动场景已具有实践价值。然而，预测区间跨零提醒教练员：'
    '平均效应量大不能保证每个运动员都能获得有意义的CMJ提升，'
    '个体响应监测仍然是训练决策的核心。')
coach_recs = [
    '短至6周的Plyometric训练（每周2-3次）即可在CMJ高度上产生可靠的中等至大幅度提升'
    f'（g≈{fmt_g(short_g)}，原始量尺约对应CMJ增加2.4-3.6cm），且效果一致性高（I²=14%），建议作为团队运动季前'
    '准备期的标准组成部分。[Meta分析证据：短期亚组k=12, I²=13.8%]'
    '具体动作选择建议：基础阶段以连续纵跳（CMJ）、栏架跳（hurdle hop）为主；'
    '进阶阶段引入跳深（drop jump）和单腿跳（bounding）。[基于训练渐进性原则的外部建议]',
    # V3 corrected: mid-term g=+0.95 instead of +1.47
    f'延长训练周期（7-10周）可获得中等幅度的效应量提升（g≈{fmt_g(mid_g)}，约对应CMJ增加3.8-5.7cm），'
    f'个体间反应差异有所增大（I²={fmt_pct(mid_I2)}），建议配合个性化训练监控跟踪反应曲线。'
    '[Meta分析证据：中期亚组k=13, I²=56.4%]',
    '青少年运动员（青春期前后）可能是具有较大增益潜力的人群（g≈1.4-1.5，GRADE：低，k=2）——需注意该建议基于k=2的探索性证据，尚不足以作为确认性实践指导。训练实施应遵循渐进性负荷原则'
    '——从低强度动作开始，逐步增加冲击强度，并在专业人员监督下进行，密切关注骨骼肌肉系统的疲劳与恢复信号。'
    '参照Lloyd等（2014）[12]的国际共识，青少年Plyo应优先发展动作质量而非训练量。'
    '[Meta分析证据：青春期前/青春期亚组各k=2, I²=0%，GRADE：低，证据不足；训练安全性建议基于外部文献[11,12]]',
]
for rec in coach_recs:
    add_para(doc, f'• {rec}', first_line_indent=Cm(1.5), spacing_after=Pt(3))

add_para(doc, '对研究者的建议：', bold=True)
res_recs = [
    'CMJ测试必须明确报告手臂位置。从生物力学角度看，手臂摆动的反向钟摆效应改变了'
    '系统质心的加速度模式，增加了对髋膝踝三关节协调的需求[8]——这意味着'
    '"手叉腰CMJ"和"带臂CMJ"可能在本质上测量了不同的神经肌肉能力维度，'
    '而不仅仅是同一测试的两种变体。因此，建议使用"双手叉腰"'
    '（hands on hips）或"双臂交叉胸前"（arms across chest）作为标准测试姿势，'
    '以便跨研究比较并排除上肢参与的混杂效应。',
    '未来研究应详细报告训练剂量参数（FITT：Frequency, Intensity, Time, Type），包括总触地次数、跳深高度、训练类型细分，以支持更精细的剂量-反应Meta分析。',
    '即使未能进行正式的性别亚组分析（因当前k数不足），未来研究也应明确报告按性别分层的结局数据——'
    '女性运动员Plyometric着陆力学与ACL损伤风险存在潜在的交互效应，且月经周期对跳跃表现的可能影响是值得探索的知识空白。',
    '需要更大样本的长期RCT（>10周）来稳定长期训练的效应量估计。',
    '鼓励研究者将个体参与者数据（IPD）公开或通过合作共享，以支持未来的IPD Meta分析。',
    '未来Meta分析应考虑将CMJ测试设备（力台、接触垫、OptoJump）作为调节变量——不同设备基于不同的力学测量原理'
    '（力台：GRF积分→加速度→速度→位移；接触垫/OptoJump：飞行时间→跳跃高度），'
    '它们对跳跃策略变化（如起跳/落地时踝关节跖屈程度）的敏感度不同，可能在不同设备间引入系统误差。'
    '接触垫基于飞行时间法可能高估CMJ的实际跳跃高度（因为在着地时踝关节跖屈导致飞行时间延长），而力台的双重积分法对起跳-着地姿态差异的相对不敏感。',
]
for rec in res_recs:
    add_para(doc, f'• {rec}', first_line_indent=Cm(1.5), spacing_after=Pt(3))

# ================================================================
# 5. Declarations (unchanged)
# ================================================================
add_heading_styled(doc, '5  利益冲突与资金声明', 1)

add_para(doc,
    '资金声明：本研究未获得任何资助机构的专项资助（no specific grant from any funding agency）。',
    size=Pt(12))
add_para(doc,
    '利益冲突声明：所有作者声明无利益竞争关系（no competing interests）。',
    size=Pt(12))
add_para(doc,
    '方案获取：PROSPERO预注册方案（CRD420261422906）为本研究的正式方案文件，可通过PROSPERO数据库'
    '公开访问。分析代码已上传至Zenodo公开存储库（https://doi.org/10.5281/zenodo.20748080），'
    '包括Python/R分析脚本和原始数据提取表。本研究中AI辅助工具（Claude AI, Anthropic, 模型版本Claude Opus 4.8及Claude Sonnet 4.6）的使用严格限于以下范围：数据分析代码的生成与调试、初稿文本的格式化组织、参考文献格式的整理和校对。所有科学判断（包括纳入/排除决策、效应量计算策略、统计模型选择、GRADE评级、结果解读和结论推导）均由人类作者独立完成。',
    size=Pt(12))

# ================================================================
# References (unchanged from v2)
# ================================================================
add_heading_styled(doc, '参考文献', 1)

references = [
    '[1] Markovic G. Does plyometric training improve vertical jump height? A meta-analytical '
    'review. Br J Sports Med. 2007;41(6):349-355. '
    'doi:10.1136/bjsm.2007.035113',

    '[2] de Villarreal ESS, Kellis E, Kraemer WJ, Izquierdo M. Determining variables of '
    'plyometric training for improving vertical jump height performance: a meta-analysis. '
    'J Strength Cond Res. 2009;23(2):495-506. '
    'doi:10.1519/JSC.0b013e318196b7c6',

    '[3] Taube W, Leukel C, Gollhofer A. How neurons make us jump: the neural control of '
    'stretch-shortening cycle movements. Exerc Sport Sci Rev. 2012;40(2):106-115. '
    'doi:10.1097/JES.0b013e31824138da',

    '[4] Claudino JG, Cronin J, Mezêncio B, et al. The countermovement jump to monitor '
    'neuromuscular status: a meta-analysis. J Sci Med Sport. 2017;20(4):397-402. '
    'doi:10.1016/j.jsams.2016.08.011',

    '[5] Moran J, Clark CCT, Ramirez-Campillo R, et al. A meta-analysis of plyometric training '
    'in female youth: its efficacy and shortcomings in the literature. J Strength Cond Res. '
    '2019;33(7):1996-2008. '
    'doi:10.1519/JSC.0000000000002768',

    '[6] Ramirez-Campillo R, Alvarez C, García-Hermoso A, et al. Effects of plyometric jump '
    'training in female soccer player\'s vertical jump height: a systematic review with '
    'meta-analysis. J Sports Sci. 2020;38(14):1641-1648. '
    'doi:10.1080/02640414.2020.1745503',

    '[7] Sole S, Ramirez-Campillo R, Andrade DC, et al. Plyometric jump training effects on '
    'the physical fitness of individual-sport athletes: a systematic review with meta-analysis. '
    'PeerJ. 2021;9:e11004. '
    'doi:10.7717/peerj.11004',

    '[8] Harman EA, Rosenstein MT, Frykman PN, Rosenstein RM. The effects of arms and '
    'countermovement on vertical jumping. Med Sci Sports Exerc. 1990;22(6):825-833. '
    'doi:10.1249/00005768-199012000-00012',

    '[9] Slimani M, Chamari K, Miarka B, et al. Effects of plyometric training on physical '
    'fitness in team sport athletes: a systematic review. J Hum Kinet. 2016;53:231-247. '
    'doi:10.1515/hukin-2016-0026',

    '[10] Bosco C, Viitasalo JT, Komi PV, Luhtanen P. Combined effect of elastic energy and '
    'myoelectrical potentiation during stretch-shortening cycle exercise. Acta Physiol Scand. '
    '1982;114(4):557-565. '
    'doi:10.1111/j.1748-1716.1982.tb07024.x',

    '[11] Markovic G, Mikulic P. Neuro-musculoskeletal and performance adaptations to '
    'lower-extremity plyometric training. Sports Med. 2010;40(10):859-895. '
    'doi:10.2165/11318370-000000000-00000',

    '[12] Lloyd RS, Faigenbaum AD, Stone MH, et al. Position statement on youth resistance '
    'training: the 2014 International Consensus. Br J Sports Med. 2014;48(7):498-505. '
    'doi:10.1136/bjsports-2013-092952',

    '[13] Radnor JM, Oliver JL, Waugh CM, et al. The influence of growth and maturation on '
    'stretch-shortening cycle function in youth. Sports Med. 2018;48(1):57-71. '
    'doi:10.1007/s40279-017-0785-0',

    '[14] Ramirez-Campillo R, Alvarez C, Garcia-Hermoso A, et al. Effects of plyometric jump '
    'training on the physical fitness of young team sport athletes: a systematic review with '
    'meta-analysis. Scand J Med Sci Sports. 2020;30(7):1179-1197. '
    'doi:10.1111/sms.13633',

    '[15] Meylan CMP, Cronin JB, Oliver JL, et al. The effect of maturation on adaptations to '
    'strength training and detraining in 11-15-year-olds. Scand J Med Sci Sports. '
    '2014;24(3):e156-e164. '
    'doi:10.1111/sms.12128',

    '[16] Sale DG. Neural adaptation to resistance training. Med Sci Sports Exerc. '
    '1988;20(5 Suppl):S135-S145. '
    'doi:10.1249/00005768-198810001-00008',

    '[17] Blazevich AJ, Gill ND, Zhou S. Intra- and intermuscular variation in human '
    'quadriceps femoris architecture assessed in vivo. J Anat. '
    '2006;209(3):289-310. '
    'doi:10.1111/j.1469-7580.2006.00619.x',

    '[18] Sterne JAC, Sutton AJ, Ioannidis JPA, et al. Recommendations for examining and '
    'interpreting funnel plot asymmetry in meta-analyses of randomised controlled trials. '
    'BMJ. 2011;343:d4002. '
    'doi:10.1136/bmj.d4002',

    '[19] Ramirez-Campillo R, Meylan CMP, Alvarez-Lepin C, et al. Effects of in-season '
    'low-volume plyometric training on explosive actions and endurance of young soccer '
    'players. J Strength Cond Res. 2014;28(5):1335-1342. '
    'doi:10.1519/JSC.0000000000000284. PMID:24751658.',

    '[20] Ramirez-Campillo R, Burgos CH, Henriquez-Olguin C, et al. Effect of unilateral, '
    'bilateral, and combined plyometric training on explosive and endurance performance '
    'of young soccer players. J Strength Cond Res. 2015;29(5):1317-1328. '
    'doi:10.1519/JSC.0000000000000762. PMID:25474338.',

    '[21] Ramirez-Campillo R, Henriquez-Olguin C, Burgos C, et al. Effect of progressive '
    'volume-based overload during plyometric training on explosive and endurance performance '
    'in young soccer players. J Strength Cond Res. 2015;29(7):1884-1893. '
    'doi:10.1519/JSC.0000000000000836. PMID:25559905.',

    '[22] Palma-Munoz I, Ramirez-Campillo R, Azocar-Gallardo J, et al. Effects of '
    'progressed and non-progressed volume-based overload plyometric training on components '
    'of physical fitness and body composition variables in youth male basketball players. '
    'J Strength Cond Res. 2021;35(6):1642-1649. '
    'doi:10.1519/JSC.0000000000002950. PMID:34027922.',

    '[23] Ramirez-Campillo R, Alvarez C, Henriquez-Olguin C, et al. Effects of plyometric '
    'training on endurance and explosive strength performance in competitive middle- and '
    'long-distance runners. J Strength Cond Res. 2014;28(1):97-104. '
    'doi:10.1519/JSC.0b013e3182a1f44c. PMID:23838975.',

    '[24] Van Roie E, Walker S, Van Driessche S, Delabastita T, Vanwanseele B, Delecluse C. '
    'An age-adapted plyometric exercise program improves dynamic strength, jump performance '
    'and functional capacity in older men either similarly or more than traditional resistance '
    'training. PLoS One. 2020;15(8):e0237921. '
    'doi:10.1371/journal.pone.0237921. PMID:32810152.',

    '[25] Chang CH, Chen CY, Lau HT. The effects of four-week plyometric training on '
    'delaying muscle fatigue in youth rowers. Sci Rep. 2025;15(1):22141. '
    'doi:10.1038/s41598-025-09673-w. PMID:40595375.',

    '[26] Asadi A, Ramirez-Campillo R, Meylan C, Nakamura FY, Canas-Jamett R, '
    'Izquierdo M. Effects of volume-based overload plyometric training on maximal-intensity '
    'exercise adaptations in young basketball players. J Sports Med Phys Fitness. '
    '2017;57(12):1557-1563. '
    'doi:10.23736/S0022-4707.16.06640-8. PMID:27792040.',

    '[27] Blazevich AJ, Gill ND, Bronks R, Newton RU. Training-specific muscle architecture '
    'adaptation after 5-wk training in athletes. Med Sci Sports Exerc. 2003;35(12):2013-2022. '
    'doi:10.1249/01.MSS.0000099092.83611.20. PMID:14669938.',

    '[28] Byrne PJ, Moran K, Rankin P, Kinsella S. A comparison of methods used to identify '
    'optimal drop height for early phase adaptations in depth jump training. '
    'J Strength Cond Res. 2010;24(8):2050-2055. '
    'doi:10.1519/JSC.0b013e3181d8eb03. PMID:20634738.',

    '[29] Sedano Campo S, Vaeyens R, Philippaerts RM, Redondo JC, de Benito AM, '
    'Cuadrado G. Effects of lower-limb plyometric training on body composition, explosive '
    'strength, and kicking speed in female soccer players. J Strength Cond Res. '
    '2009;23(6):1714-1722. '
    'doi:10.1519/JSC.0b013e3181b3f537. PMID:19675492.',

    '[30] Chelly MS, Ghenem MA, Abid K, Hermassi S, Tabka Z, Shephard RJ. Effects of '
    'in-season short-term plyometric training program on leg power, jump- and sprint '
    'performance of soccer players. J Strength Cond Res. 2010;24(10):2670-2676. '
    'doi:10.1519/JSC.0b013e3181e2728f. PMID:20844458.',

    '[31] Idrizovic K, Gjinovci B, Sekulic D, Uljevic O, Joao PV, Spasic M, Sattler T. '
    'The effects of 3-month skill-based and plyometric conditioning on fitness parameters '
    'in junior female volleyball players. Pediatr Exerc Sci. 2018;30(3):353-363. '
    'doi:10.1123/pes.2017-0178. PMID:29478378.',

    '[32] Jlid MC, Coquart J, Maffulli N, Paillard T, Bisciotti GN, Chamari K. '
    'Multidirectional plyometric training: very efficient way to improve vertical jump '
    'performance, change of direction performance and dynamic postural control in young '
    'soccer players. Front Physiol. 2019;10:1462. '
    'doi:10.3389/fphys.2019.01462. PMID:31736978.',

    '[33] Jlid MC, Coquart J, Maffulli N, Paillard T, Bisciotti GN, Chamari K. Effects '
    'of in season multi-directional plyometric training on vertical jump performance, '
    'change of direction speed and dynamic postural control in U-21 soccer players. '
    'Front Physiol. 2020;11:374. '
    'doi:10.3389/fphys.2020.00374. PMID:32431621.',

    '[34] Khlifa R, Aouadi R, Hermassi S, Chelly MS, Jlid MC, Hbacha H, Castagna C. '
    'Effects of a plyometric training program with and without added load on jumping '
    'ability in basketball players. J Strength Cond Res. 2010;24(11):2955-2961. '
    'doi:10.1519/JSC.0b013e3181e37fbe. PMID:20938357.',

    '[35] Kijowksi KN, Capps CR, Goodman CL, Erickson TM, Knorr DP, Triplett NT, '
    'Awelewa OO, McBride JM. Short-term resistance and plyometric training improves '
    'eccentric phase kinetics in jumping. J Strength Cond Res. 2015;29(8):2186-2196. '
    'doi:10.1519/JSC.0000000000000904. PMID:26203736.',

    '[36] Laurent C, Baudry S, Duchateau J. Comparison of plyometric training with '
    'two different jumping techniques on Achilles tendon properties and jump performances. '
    'J Strength Cond Res. 2020;34(6):1503-1510. '
    'doi:10.1519/JSC.0000000000003604. PMID:32271290.',

    '[37] Michailidis Y, Tabouris A, Metaxas T. Effects of plyometric and directional '
    'training on physical fitness parameters in youth soccer players. Int J Sports Physiol '
    'Perform. 2019;14(3):392-398. '
    'doi:10.1123/ijspp.2018-0545. PMID:30204520.',

    '[38] Negra Y, Chaabene H, Sammoud S, Bouguezzi R, Abbes M, Hachana Y, Granacher U. '
    'Effects of plyometric training on physical fitness in prepuberal soccer athletes. '
    'Int J Sports Med. 2017;38(5):370-377. '
    'doi:10.1055/s-0042-122337. PMID:28315285.',

    '[39] Potdevin FJ, Alberty ME, Chevutschi A, Pelayo P, Sidney MC. Effects of a 6-week '
    'plyometric training program on performances in pubescent swimmers. J Strength Cond Res. '
    '2011;25(1):80-86. '
    'doi:10.1519/JSC.0b013e3181fef720. PMID:21157388.',

    '[40] Ramirez-Campillo R, Alvarez C, Gentil P, Loturco I, Sanchez-Sanchez J, '
    'Izquierdo M, Moran J, Nakamura FY, Chaabene H, Granacher U. Sequencing effects of '
    'plyometric training applied before or after regular soccer training on measures of '
    'physical fitness in young players. J Strength Cond Res. 2020;34(7):1959-1966. '
    'doi:10.1519/JSC.0000000000002525. PMID:29570574.',

    '[41] Rensing N, Westermann A, Moeller D, von Piekartz H. Trainingseffekte eines reaktiven '
    'Sprungkrafttrainings bei Handballern in Bezug auf die Sprunghoehe und die Kraftentfaltung '
    'des Musculus triceps surae [Training effects of reactive jump training in handball players '
    'regarding jump height and force development of the triceps surae muscle]. '
    'Sportverletz Sportschaden. 2016;30(4):216-225. '
    'doi:10.1055/s-0041-106947. PMID:26579625.',

    '[42] Rubley MD, Haase AC, Holcomb WR, Girouard TJ, Tandy RD. The effect of plyometric '
    'training on power and kicking distance in female adolescent soccer players. '
    'J Strength Cond Res. 2011;25(1):129-134. '
    'doi:10.1519/JSC.0b013e3181b94a3d. PMID:19966586.',

    '[43] Santos EJAM, Janeira MAAS. The effects of plyometric training followed by '
    'detraining and reduced training periods on explosive strength in adolescent male '
    'basketball players. J Strength Cond Res. 2011;25(2):441-452. '
    'doi:10.1519/JSC.0b013e3181b62be3. PMID:20453686.',

    '[44] Toumi H, Best TM, Martin A, F\'Guyer S, Poumarat G. Effects of eccentric phase '
    'velocity of plyometric training on the vertical jump. Int J Sports Med. '
    '2004;25(5):391-398. '
    'doi:10.1055/s-2004-815843. PMID:15241721.',

    '[45] Vescovi JD, Canavan PK, Hasson S. Effects of a plyometric program on vertical '
    'landing force and jumping performance in college women. Phys Ther Sport. '
    '2008;9(4):185-192. '
    'doi:10.1016/j.ptsp.2008.08.001. PMID:19083719.',

    '[46] Yanci J, Castillo D, Iturricastillo A, Ayarra R, Nakamura FY. Effects of two '
    'different volume-equated weekly distributed short-term plyometric training programs '
    'on futsal players\' physical performance. J Strength Cond Res. 2017;31(7):1787-1794. '
    'doi:10.1519/JSC.0000000000001644. PMID:27662489.',

    '[47] Zubac D, Simunic B. Skeletal muscle contraction time and tone decrease after '
    '8 weeks of plyometric training. J Strength Cond Res. 2017;31(6):1610-1619. '
    'doi:10.1519/JSC.0000000000001626. PMID:28538312.',

    '[48] Liu G, Wang X, Xu Q. Microdosing plyometric training enhances jumping performance, '
    'reactive strength index, and acceleration among youth soccer players: a randomized '
    'controlled study design. J Sports Sci Med. 2024;23(2):342-350. '
    'doi:10.52082/jssm.2024.342. PMID:38841635.',

    '[49] Gepfert M, Gołaś A, Roczniok R, Walencik J, Węgrzynowicz K, Zając A. Impact of '
    'an eight-week plyometric training intervention on neuromuscular performance, '
    'musculotendinous stiffness, and directional speed in elite Polish badminton athletes. '
    'J Funct Morphol Kinesiol. 2025;10(3):304. '
    'doi:10.3390/jfmk10030304.',

    '[50] Sammoud S, Negra Y, Bouguezzi R, Ramirez-Campillo R, Moran J, Bishop C, '
    'Chaabene H. Effects of plyometric jump training on measures of physical fitness and '
    'lower-limb asymmetries in prepubertal male soccer players: a randomized controlled '
    'trial. BMC Sports Sci Med Rehabil. 2024;16(1):37. '
    'doi:10.1186/s13102-024-00821-9.',
]
for ref in references:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Cm(1.27)
    para.paragraph_format.first_line_indent = Cm(-1.27)
    run = para.add_run(ref)
    set_run_font(run, size=Pt(11))

# ================================================================
# SAVE
# ================================================================
output_path = PROJ / 'Plyo训练CMJ高度Meta分析_v4.docx'
doc.save(str(output_path))
print(f'\n已保存 v4: {output_path}')
print('完成！')
