# -*- coding: utf-8 -*-
"""第四轮多Agent独立评审 + LLM决策框架 — 生成审核报告Word文档"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

STYLE_FONT = 'Times New Roman'
CJK_FONT = '宋体'
HEADING_FONT = '黑体'

def set_run_font(run, size=None, bold=False, font_name=None, cjk=None):
    font = font_name or STYLE_FONT
    run.font.name = font
    if cjk is None:
        cjk = CJK_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), cjk)
    rPr.insert(0, rFonts)
    if size:
        run.font.size = size
    run.bold = bold

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, cjk=HEADING_FONT)
    return h

def add_para(doc, text, bold=False, size=Pt(11), align=None,
             first_line_indent=Cm(0.74), spacing_after=Pt(4)):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = spacing_after
    para.paragraph_format.first_line_indent = first_line_indent
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, size=size, bold=bold)
    return para

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=Pt(8), bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=Pt(8))
    doc.add_paragraph()

# ================================================================
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = STYLE_FONT
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), CJK_FONT)
rPr.insert(0, rFonts)

# ================================================================
# 标题页
# ================================================================
add_para(doc, '', size=Pt(8))
add_para(doc,
    'Plyometric训练对CMJ高度影响的Meta分析',
    bold=True, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=None, spacing_after=Pt(4))
add_para(doc,
    '第四轮多Agent独立评审 + LLM决策框架',
    bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=None, spacing_after=Pt(12))

add_para(doc, '审查日期：2026年6月19日', size=Pt(10),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '审查框架：4维度专业Agent并行审查 + 综合LLM决策框架', size=Pt(10),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '审查维度：统计方法学 (Dr. Sigma) | 数据质量 (Dr. Audit) | PRISMA报告规范 (Dr. PRISMA) | 科学解释与实践转化 (Dr. Practice)', size=Pt(9),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '目标期刊：JSCR / J Sports Sci / Biol Sport / Sports Med', size=Pt(9),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '论文路径：D:\\桌面\\Plyo训练CMJ高度Meta分析_初稿.docx', size=Pt(9),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '审查文件数：30+ 个源文件（手稿 + CSV数据 + JSON输出 + Python/R分析代码 + PDF原文）', size=Pt(9),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, spacing_after=Pt(16))

# ================================================================
# 执行摘要
# ================================================================
add_heading_styled(doc, '执行摘要', 1)

add_para(doc, '**总体评估：需要重大修改 (MAJOR REVISION)** — 存在1个新增致命缺陷，多项重大问题未修正', bold=True, size=Pt(11))

add_para(doc,
    '第四轮审查发现，第三轮（2026年6月18日）识别出的4个致命缺陷（FF1-FF4）中：'
    'FF2（样本量）和FF4（参考文献）已修正；FF1（I² CI来源）在手稿文本中部分修正但recompute_i2_ci.py代码中'
    '的亚组ID集合存在错误；FF3（GRADE样本量）状态待确认。')

add_para(doc,
    '第四轮新增了一项致命缺陷（FF5：R19 Michailidis 2018的SD/SE混淆 — SE值0.80被错误当作SD使用，'
    '导致该研究的效应量被高估约3.9倍，进而影响严格池合并效应量）。所有4个Agent均独立确认了这一问题的存在。')

add_para(doc,
    '此外，第四轮确认了多项残余的重大问题（M1-M12中至少4项仍待修正），以及新发现的问题。'
    '核心结论：论文的基础研究设计和分析方法学框架是扎实的，但多个执行层面的数据错误和报告不一致问题在投稿前必须解决。',
    bold=False, size=Pt(11))

# ================================================================
# 第四轮新增致命缺陷
# ================================================================
add_heading_styled(doc, '第一部分：第四轮新增致命缺陷（FF5：本轮新发现）', 1)

add_heading_styled(doc, 'FF5. R19 (Michailidis 2018) SD/SE混淆 — 效应量被高估约3.9倍', 2)

add_para(doc, '**严重等级：致命（FATAL）**', bold=True)
add_para(doc, '**发现Agent：Dr. Audit + Dr. Sigma（独立确认）**')
add_para(doc, '**位置：analysis_ready_effects.csv R19行；clean_and_compute_effects.py R19数据字典**')

add_para(doc, '**问题描述：**', bold=True)
add_para(doc,
    'R19 (Michailidis 2018, IJSPP, n=31, U-13足球学员) 的4个SD值（IG前测、IG后测、CG前测、CG后测）'
    '在analysis_ready_effects.csv中全部设为0.80。数据标注说明"CI_width/3.92≈0.80"。但CI_width/3.92 = SE（标准误），'
    '而非SD（标准差）。真实SD = SE × √n：IG真实SD = 0.80 × √17 = 3.30cm，CG真实SD = 0.80 × √14 = 2.99cm。')

add_para(doc, '**证据链：**', bold=True)
add_para(doc,
    '1. 若SD=0.80，则SE=SD/√17=0.194，CI_width=3.92×0.194=0.76cm。这不符合CI_width/3.92=0.80的前提。'
    '2. 若SE=0.80，则CI_width=3.92×0.80=3.14cm，CI=均值±1.57，这在U-13青少年CMJ数据中合理。'
    '3. 4个SD值全部相等（0.80）在真实数据中几乎不可能（前后测SD完全相等），这是"从同一个CI反推公式批量复制"的典型痕迹。')

add_para(doc, '**影响评估：**', bold=True)
add_para(doc,
    '使用SE=0.80作为SD=0.80，使效应量分母缩小约3.9倍，g值被相应高估。'
    '当前g_change ≈ 1.48，真实g_change估计约0.30-0.40。'
    'R19在严格池16篇中占比1/16，修正后可能使严格池合并g下降0.05-0.10。'
    '这是一个数据完整性的根本性错误，与第三轮发现的FF1（I² CI拼接错误）性质类似但影响不同。')

add_para(doc, '**修正方案（3个选项）：**', bold=True)
add_para(doc,
    '选项A（强烈推荐）：回到原始论文PDF确认原文报告的是SE还是SD。若为SE，按SD=SE×√n重新计算，更新analysis_ready_effects.csv和所有下游输出。'
    '选项B：若原文确实报告SD=0.80（Myotest设备在高度同质的U-13精英足球学员中可能出现极小SD），需在手稿中明确标注并论证其合理性。'
    '选项C（最保守）：从严格池中排除R19，标注原因"SD来源无法确认（CI反推存在SE/SD歧义）"，并报告排除后的敏感性分析结果。')

# ================================================================
# 第三轮致命缺陷修正状态追踪
# ================================================================
add_heading_styled(doc, '第二部分：第三轮致命缺陷（FF1-FF4）修正状态追踪', 1)

add_table(doc,
    ['缺陷编号', '问题描述', '第三轮状态', '第四轮验证', '修正状态'],
    [
        ['FF1', 'I²置信区间在数学上不可能\n（点估计落在CI外）',
         '已识别\n需要R Q-profile重新计算',
         '手稿文本中I² CI已更新为正确值\n但recompute_i2_ci.py中short_ids\n和mid_ids集合仍有错误',
         '⚠️ 部分修正\n代码需修复后再运行'],
        ['FF2', '总体样本量错误\n（732 vs 实际718）',
         '已识别\n需要逐篇核对',
         'verify2.py确认：IG=367, CG=351,\nTotal=718。手稿当前报告718\n摘要和正文一致',
         '✅ 已修正'],
        ['FF3', 'GRADE表样本量严重错误\n（严格池N=520 vs 实际341）',
         '已识别\n需要从CSV重新计算',
         '待确认：当前手稿GRADE表\n是否已更新为正确N值',
         '❓ 待验证'],
        ['FF4', '参考文献复制粘贴错误\n（Ref 32/33克隆数据）',
         '已识别\n需要逐条PubMed验证',
         '当前参考文献使用[1]-[19]\n连续编号（非传统格式）\n缺少纳入研究的完整文献列表',
         '⚠️ 部分修正'],
    ])

# ================================================================
# 重大问题修正状态
# ================================================================
add_heading_styled(doc, '第三部分：重大问题（M1-M12）修正状态追踪', 1)

add_table(doc,
    ['编号', '问题', '严重度', '第四轮状态', '备注'],
    [
        ['M1', 'r敏感性表使用过时数据\n（r=0.7行不能复现主分析）', '重大',
         '❌ 未修正',
         'sensitivity_r_sensitivity.py使用DL近似\n而非REML；需改用R/metafor管道'],
        ['M3', '缺少结构化摘要', '重大',
         '✅ 已修正',
         '当前手稿包含完整PRISMA 2020结构化摘要'],
        ['M4', '4个数据库中3个检索策略缺失', '重大',
         '❓ 待验证',
         '需确认补充材料是否已包含Scopus/WoS/\nGoogle Scholar完整检索策略'],
        ['M5', '双重筛选过程未描述', '重大',
         '✅ 已修正',
         '手稿2.3节已新增筛选过程描述\n含kappa=0.86/0.91'],
        ['M6', '缺少资金和利益冲突声明', '重大',
         '✅ 已修正',
         '手稿第5节已包含两项声明'],
        ['M7', 'Meta回归p值内部矛盾\n(p=0.023 vs p=0.026)', '重大',
         '❌ 未修正',
         '多处讨论文本仍使用p=0.026\n（应统一为p=0.023）'],
        ['M8', '参考文献格式不标准', '重大',
         '⚠️ 部分修正',
         '已改为连续编号[1]-[19]\n但仅含19条参考文献（正文引用\n远多于此），且缺纳入研究列表'],
        ['M9', 'R11个体效应量使用错误度量', '重大',
         '❓ 待验证',
         '需检查讨论中R11引用是否\n已统一为变化分数g=+5.20'],
        ['M10', '年龄亚组遗漏R06 (Van Roie 2020)', '重大',
         '❓ 待验证',
         'R06为老年人(69-70y)，\n需在年龄亚组中单独成组或\n明确标注排除原因'],
        ['M12', 'PEDro百分比错误 (55% vs 57%)', '次要→重大',
         '❓ 待验证',
         '重新计算：17/29=58.6% ≥6分\n需确认手稿是否已更正'],
    ])

# ================================================================
# 第四轮专项Agent发现
# ================================================================
add_heading_styled(doc, '第四部分：专项Agent深度审查发现', 1)

# --- Dr. Sigma ---
add_heading_styled(doc, 'A. Dr. Sigma — 统计方法学审查 (评分：5.0/10)', 2)

add_para(doc, '**核心发现：**', bold=True)
add_para(doc,
    '1. I² CI仍来自错误模型运行 — recompute_i2_ci.py中short_ids和mid_ids集合包含R24(VJ)和R25/R28(重复)，需彻底修复。'
    '2. 预测区间(PI)来源不一致 — 手稿使用Analysis_Report.txt的PI（来自Python DL方法），而正确的PI应来自R/metafor REML模型。修正后严格PI=[-0.564,+2.819]（手稿当前[-0.878,+3.134]）。'
    '3. τ² 95%CI已计算但未在手稿中报告 — 现代Meta分析标准推荐报告τ² CI。严格池τ²=0.687[95%CI: 0.345, 3.337]。'
    '4. 多重比较问题 — 12次检验后Bonferroni校正使所有"显著"结果不显著。p值引用不一致（p=0.023 vs p=0.026）。'
    '5. r敏感性表与主分析不兼容 — sensitivity_r_sensitivity.py使用内部Python DL近似而非R/metafor REML管道。')

# --- Dr. Audit ---
add_heading_styled(doc, 'B. Dr. Audit — 数据质量审计 (核心评级：含1致命缺陷)', 2)

add_para(doc, '**审计点概览：**', bold=True)

add_table(doc,
    ['审计点', '内容', '严重等级', '发现'],
    [
        ['AP1', '总样本量', '次要', 'IG=367, CG=351, Total=718 ✅（第三轮已修正）'],
        ['AP2', 'R19 SD/SE混淆', '致命(FF5)', 'SE=0.80被当作SD=0.80使用 → g被高估~3.9×'],
        ['AP3', 'R27 Toumi 2004', '重大(已处理)', 'Smith机SSC非经典Plyo，已标记离群值'],
        ['AP4', 'Ramirez-Campillo 5篇', '次要', '5篇独立样本（不同PMID,不同人群）✅'],
        ['AP5', 'R23 SD近似', '次要', 'CG SD用IG SD近似，RCT前提下可接受'],
        ['AP6', 'R29后测推算', '次要', 'Post=Pre+Δ方法正确 ✅'],
        ['AP7', 'SEM→SD换算', '次要', 'R16和R06换算均正确 ✅'],
        ['AP8', 'PEDro评分', '次要', '整体合理，R09(3分)可能需要重新评估→4分'],
        ['AP9', '分析池成员', '重大', 'RevMan宽版池文件含R24(VJ非CMJ)'],
    ])

add_para(doc,
    '**审计点9详细说明：** RevMan_ready_wide_pool.csv包含29行数据（含R24 Rubley 2011, VJ带臂）。'
    '若宽版池应排除R24（VJ仅用于敏感性亚组），则宽版池的森林图/漏斗图额外包含了一个非CMJ研究。'
    '需从RevMan宽版文件中移除R24并重新生成图形。')

# --- Dr. PRISMA ---
add_heading_styled(doc, 'C. Dr. PRISMA — 报告规范审查 (评分：76/100)', 2)

add_para(doc, '**Top 3关键问题：**', bold=True)
add_para(doc,
    '1. **Item 16b (方案偏离) — 不合规（致命级投稿风险）**：手稿未报告r=0.7假设、严格池vs宽版池区分、'
    '亚组/Meta回归分析是否在PROSPERO中预先设定。任何偏离必须明确报告并说明理由。'
    '2. **数据可用性 — 部分合规但存在严重缺陷**：Zenodo有分析代码，但data_extraction_FINAL.csv大部分为空'
    '（数据硬编码在Python字典中），不支持独立复现。'
    '3. **补充材料完整性无法验证**：手稿提到S1-S7补充材料但实际存在性待确认。S2（4数据库完整检索策略）和'
    'S4（数据提取表）是最高风险项。')

add_para(doc, '**其他PRISMA问题：**', bold=True)
add_para(doc,
    '• Item 5: 仅提供PubMed检索策略，其他3个数据库策略"详见补充材料"但补充材料未确认存在。'
    '• Item 13e: 每个综合中纳入的研究数已正确报告（k=16, k=28等）。'
    '• Item 20a: PROSPERO注册号已报告但无法公开验证。'
    '• Item 22: PRISMA Checklist本身未随稿件提供。'
    '• Item 23a: 利益冲突声明已添加。'
    '• Item 26: 代码DOI已提供但数据提取表不可获取。'
    '• Item 27: AI辅助工具使用未声明。')

# --- Dr. Practice ---
add_heading_styled(doc, 'D. Dr. Practice — 科学解释与实践转化审查 (评分：5.2/10)', 2)

add_para(doc, '**紧急行动项（4项）：**', bold=True)

add_table(doc,
    ['维度', '评分', '风险', '核心问题'],
    [
        ['实践建议证据基础', '3.5/10', '🔴紧急',
         '7项具体推荐中5项无法追溯到Meta分析数据。60-120次触地、30-60cm箱高、'
         '每2-4周测试等具体参数来自外部文献综述，而非本Meta分析发现。'
         '必须在Meta分析证据与外部专家建议之间建立明确区分标注。'],
        ['训练安全性讨论', '3.0/10', '🔴紧急',
         '29篇RCT中0篇提供损伤定量数据。所有安全建议均来自外部共识，'
         '必须标注为"基于外部文献的预防性建议"。'],
        ['MCID转换', '4.0/10', '🟡高',
         '"加权合并基线SD≈4.5cm"来源不明，无法验证。纳入研究SD范围0.8-16cm，'
         '跨20倍差异。Franchi 2022可能不是正确的MCID来源（该文关于分子调控而非CMJ）。'],
        ['剂量-反应线性假设', '5.5/10', '🟡高',
         '亚组分类均值（0.60→1.46→1.85）与每周+0.13线性斜率不一致。'
         '6→7周间~2.5倍效应量跃升需要更充分的生理学论证。'],
    ])

add_para(doc, '**高优先级行动项（5项）：**', bold=True)
add_para(doc,
    '5. 臂位反常数据：严格手叉腰池I²=76% > 臂位未明池I²=20%，与"臂位分层排除噪声"的前提矛盾，需在讨论中面对。'
    '6. 足球研究过度代表（11/29=38%），推广性限制需在局限性中明确讨论。'
    '7. CMJ设备差异（力台vs接触垫vs OptoJump vs Myotest）未作为调节变量，可能引入间接性偏倚。'
    '8. 效应量解释："手稿多处重复强调大效应(g>1.0)"的修辞模式在PI跨零的背景下可能误导读者，建议在讨论开篇即纳入PI跨零和I²的不确定性。'
    '9. 局限性应补充5-7条遗漏项（足球偏倚、设备差异、多臂选择主观性、训练依从性未报告、未区分对照类型含义）。')

# ================================================================
# LLM决策框架
# ================================================================
add_heading_styled(doc, '第五部分：LLM决策框架 — 综合裁决与优先级排序', 1)

add_heading_styled(doc, '5.1 整体科学质量评估', 2)

add_table(doc,
    ['维度', '评分', '评审人', '关键评语'],
    [
        ['统计方法学', '5.0/10', 'Dr. Sigma',
         '方法框架严格(REML+Q-profile+PI+GRADE)\n但执行存在3项致命级错误'],
        ['数据质量', '4.0/10', 'Dr. Audit',
         'FF5(R19 SD/SE混淆)是致命缺陷\n大多数其他数据点经过验证正确'],
        ['PRISMA合规', '76/100', 'Dr. PRISMA',
         '结构化摘要已添加，筛选过程已描述\n但方案偏离未报告 + 数据可用性不足'],
        ['科学解释', '5.2/10', 'Dr. Practice',
         '讨论有深度和诚实性\n但实践建议与Meta证据脱节严重'],
    ])

add_heading_styled(doc, '5.2 致命缺陷裁定', 2)

add_table(doc,
    ['缺陷编号', '描述', '是否独立导致拒稿', '第四轮裁定'],
    [
        ['FF1', 'I² CI来自错误模型运行\n（点估计落在CI外）', '是',
         '维持致命 → 降级为重大\n（手稿文本已修正CI值\n但代码需修复后重新验证）'],
        ['FF2', '总样本量错误\n（732→已修正为718）', '是',
         '已修正 ✅'],
        ['FF3', 'GRADE表样本量错误\n（520→待确认是否>341）', '是',
         '待验证（需检查当前手稿GRADE表）'],
        ['FF4', '参考文献克隆数据', '是',
         '降至重大（编号已修正\n但纳入研究文献仍缺失）'],
        ['FF5', 'R19 SD/SE混淆\n（效应量被高估~3.9倍）', '是',
         '**新增致命缺陷**\n需回到原始论文验证\n修正前分析结果不可靠'],
    ])

add_heading_styled(doc, '5.3 裁决逻辑', 2)

add_para(doc,
    '**裁决：需要重大修改 (MAJOR REVISION)**', bold=True, size=Pt(11))

add_para(doc,
    '虽然FF1-FF4已部分修正，但第四轮新发现的FF5（R19 SD/SE混淆）是一个独立的致命缺陷。'
    '该问题直接影响主分析（严格池）的合并效应量估计。在FF5解决之前，本手稿无法通过任何严格同行评审。')

add_para(doc,
    '然而，本研究在以下几个维度上具有真正的方法学优势，支持"修改后可投稿"的判断：')

add_para(doc,
    '**(a) 研究问题的创新性不可否认** — 首次将CMJ手臂位置作为纳入标准和分层变量，填补了方法学空白。'
    '**(b) 分析方法学框架高于领域平均水平** — REML随机效应+Q-profile I² CI+预测区间+GRADE的组合在运动科学Meta分析中不常见。'
    '**(c) 讨论的诚实性和自我批评性值得肯定** — PI跨零、发表偏倚信号、多重比较未校正、线性外推局限等均已在手稿中讨论。'
    '**(d) 第三轮已修正了多项关键问题** — 结构化摘要、筛选过程描述、资金/利益冲突声明、样本量等已在手稿中体现。')

add_para(doc,
    '**核心差距不在于研究设计或分析框架，而在于数据管理和报告层面的执行细节。** '
    '这一定位意味着：不需要重新设计研究或重新收集数据；需要的是系统性的数据验证（特别是R19）和报告一致性审核。')

# ================================================================
# 完整修正清单
# ================================================================
add_heading_styled(doc, '第六部分：投稿前完整修正清单（按优先级排序）', 1)

add_heading_styled(doc, 'P0 — 致命（投稿前必须完成，否则将被拒稿）', 2)

add_table(doc,
    ['优先级', '编号', '修正内容', '预计工作量'],
    [
        ['P0-1', 'FF5', '解决R19 SD/SE混淆：\n'
         '(a) 回到原始论文PDF确认SD vs SE\n'
         '(b) 按正确值重新计算R19效应量\n'
         '(c) 重新运行完整分析管道\n'
         '(d) 更新所有手稿数值\n'
         '(e) 如无法确认，排除R19并标注', '2-4小时'],
        ['P0-2', 'FF3', '验证GRADE表N值已修正为341(严格池)/624(宽版池)', '30分钟'],
        ['P0-3', 'AP9', '从RevMan宽版池文件移除R24(VJ)，重新生成森林图/漏斗图', '1小时'],
        ['P0-4', 'PRISMA-16b', '添加方案偏离报告：列出PROSPERO预注册与实际方法的任何差异', '1小时'],
    ])

add_heading_styled(doc, 'P1 — 重大（投稿前强烈建议完成）', 2)

add_table(doc,
    ['优先级', '编号', '修正内容', '预计工作量'],
    [
        ['P1-1', 'FF1-残留', '修复recompute_i2_ci.py中short_ids/mid_ids集合\n'
         '使用正确的周数分配重新运行\n'
         '确保所有I² CI与主分析来自同一模型', '2小时'],
        ['P1-2', 'M1', '修复r敏感性表：使用R/metafor管道重新计算\n'
         '确保r=0.7行完全复现主分析结果', '1.5小时'],
        ['P1-3', 'M7', '统一所有p值引用：Meta回归Duration使用p=0.023\n'
         '检查讨论中所有引用位置', '30分钟'],
        ['P1-4', 'Sigma-D5', '将PI替换为来自R/metafor REML的正确值\n'
         '严格[-0.564,+2.819] 宽[-0.205,+2.209]', '30分钟'],
        ['P1-5', 'M8', '补充纳入研究的完整参考文献列表(29篇)\n'
         '每条包含完整DOI/PMID', '2小时'],
        ['P1-6', 'Practice-D5', '重组实践建议：明确区分Meta证据 vs 外部建议\n'
         '使用证据等级标注（★★★Meta发现/★☆专家建议）', '2小时'],
    ])

add_heading_styled(doc, 'P2 — 次要（可在修改过程中处理）', 2)

add_table(doc,
    ['优先级', '编号', '修正内容'],
    [
        ['P2-1', 'M10', '将R06(Van Roie 2020, 老年)纳入年龄亚组分析或明确标注排除原因'],
        ['P2-2', 'M12', '修正PEDro百分比：17/29=58.6%（≥6分）'],
        ['P2-3', 'Sigma-D4', '在主结果中补充τ² 95%CI'],
        ['P2-4', 'Sigma-D7', '在方法中明确说明"每个亚组独立估计τ²"'],
        ['P2-5', 'Practice-D4', '面对臂位反常数据(I²_严格76% > I²_臂位未明20%)'],
        ['P2-6', 'Practice-D6', '文档化"加权合并基线SD≈4.5cm"的计算过程'],
        ['P2-7', 'Practice-D8', '补充遗漏的局限性：足球偏倚(38%)、设备差异、多臂选择策略、训练依从性、对照类型含义'],
        ['P2-8', 'PRISMA-Item5', '确保补充材料包含Scopus/WoS/Google Scholar完整检索策略'],
        ['P2-9', 'PRISMA-Item22', '将PRISMA 2020 Checklist作为补充材料提交'],
        ['P2-10', 'PRISMA-Item27', '声明AI辅助工具在本研究中的使用情况'],
        ['P2-11', 'Audit-P8', '重新评估R09 PEDro评分（3→可能4分），逐项核实'],
    ])

# ================================================================
# 正面评价
# ================================================================
add_heading_styled(doc, '第七部分：值得保留的核心优势', 1)

add_para(doc, '以下优势在第四轮审查中得到确认，应在修改中保留并强化：')
add_para(doc,
    '1. **CMJ手臂位置分层** — 首次将其作为纳入标准和分层变量，是真正的方法学创新。'
    '2. **预测区间的报告与讨论** — 在运动科学Meta分析中不常见，对PI跨零的解读诚实且方法论上前沿。'
    '3. **GRADE评估框架完整** — 尽管个别评级需修正，但框架本身覆盖了所有必要领域。'
    '4. **r敏感性分析范围全面(0.5-0.9)** — 证实了效应方向的方向稳健性，尽管内部一致性需修复。'
    '5. **发表偏倚的多维度分析** — 正确区分了经典发表偏倚和小研究真实效应，Trim-and-Fill分析深入。'
    '6. **训练安全性讨论(4.6节)** — 讨论本身的存在值得肯定（大多数Meta分析忽略），尽管证据基础薄弱。'
    '7. **代码公开(Zenodo DOI)和PROSPERO预注册** — 体现了开放科学态度。'
    '8. **Inter-rater reliability报告(kappa=0.86/0.89/0.91/0.92/0.95)** — 良好的方法学实践。'
    '9. **多臂研究取组逻辑透明且有原则** — 优先取标准方案、保守选择。'
    '10. **局限性讨论的诚实性** — 10条局限性的自我批评展示了学术成熟度。')

# ================================================================
# 各Agent独立评分汇总
# ================================================================
add_heading_styled(doc, '附录：四Agent独立评分与建议汇总', 1)

add_table(doc,
    ['Agent', '评分', '推荐决定', '一句话总结'],
    [
        ['Dr. Sigma\n(统计方法学)', '5.0/10',
         'Major Revision',
         '方法框架先进但执行存在3项致命错误'],
        ['Dr. Audit\n(数据质量)', '4.0/10',
         'Major Revision\n(含1致命缺陷)',
         '核心数据错误(R19 SD/SE)使主分析不可靠'],
        ['Dr. PRISMA\n(报告规范)', '76/100',
         'Minor Revision\n(方案偏离→Major)',
         '结构化摘要已添加，方案偏离未报告是最大缺口'],
        ['Dr. Practice\n(科学解释)', '5.2/10',
         'Major Revision',
         '讨论深度好但实践建议与证据脱节'],
    ])

add_para(doc, '', size=Pt(4))
add_para(doc,
    '**最终综合裁决：MAJOR REVISION（需要重大修改）**', bold=True, size=Pt(12),
    align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc,
    '**核心判断：** 本研究的基础研究设计和方法学框架是扎实的。新增致命缺陷(FF5: R19 SD/SE混淆)和多项未修正的重大问题'
    '需要作者在投稿前系统解决。一旦致命缺陷修正完毕，本手稿有潜力对Plyometric训练文献做出有意义的贡献。'
    '预计修正后重新审查可在1-2轮内达到Accept条件。',
    bold=False, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
    spacing_after=Pt(12))

add_para(doc,
    '审查负责人：Claude Opus 4.8 (1M context) + 4 Agent并行框架',
    size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc,
    f'审查日期：2026年6月19日',
    size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

# ================================================================
# 保存
# ================================================================
output_path = 'D:/桌面/Plyo_CMJ_Meta分析_第四轮综合审查报告.docx'
doc.save(output_path)
print(f'报告已保存至: {output_path}')
print('Done.')
