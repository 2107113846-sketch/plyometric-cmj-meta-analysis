# -*- coding: utf-8 -*-
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

STYLE_FONT = 'Times New Roman'
CJK_FONT = '宋体'
HEADING_FONT = '黑体'

def set_font(run, size=None, bold=False, font_name=None, cjk=None):
    font = font_name or STYLE_FONT
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), cjk or CJK_FONT)
    rPr.insert(0, rFonts)
    if size: run.font.size = size
    if bold: run.font.bold = True

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_font(run, size=Pt(16 if level==1 else 14 if level==2 else 12), bold=True, cjk=HEADING_FONT)
    return h

def add_para(text, bold=False, size=Pt(12), indent_first_line=Cm(0.74)):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = indent_first_line
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p

def add_code_block(text, size=Pt(10)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    set_font(r, size=size, cjk='Courier New')
    return p

# ==== TITLE ====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('中文数据库补充检索策略')
set_font(r, size=Pt(18), bold=True, cjk=HEADING_FONT)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('—— Plyometric训练对CMJ高度影响的Meta分析')
set_font(r, size=Pt(14), cjk=CJK_FONT)

doc.add_paragraph()

# ==== 1 ====
add_heading_styled('1  检索数据库', 2)

add_para(
    '在PubMed、Scopus、Web of Science及Google Scholar四个英文数据库的基础上，'
    '补充检索以下三个中文数据库，'
    '以覆盖中国体育科学领域已发表和未正式发表的相关研究：',
    size=Pt(12))

table = doc.add_table(rows=4, cols=4, style='Table Grid')
headers = ['数据库', '网址', '收录范围', '检索方式']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    r = cell.paragraphs[0].add_run(h)
    set_font(r, size=Pt(11), bold=True)

db_rows = [
    ['中国知网\n(CNKI)', 'www.cnki.net',
     '期刊论文 + 硕博学位论文\n+ 会议论文 + 报纸',
     '专业检索'],
    ['万方\n(WanFang)', 'www.wanfangdata.com.cn',
     '期刊论文 + 学位论文\n+ 会议论文',
     '专业检索'],
    ['维普\n(VIP)', 'www.cqvip.com',
     '期刊论文为主',
     '高级检索'],
]
for i, row_data in enumerate(db_rows):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = ''
        r = cell.paragraphs[0].add_run(cell_text)
        set_font(r, size=Pt(10))

doc.add_paragraph()

# ==== 2 ====
add_heading_styled('2  检索策略', 2)
add_para('检索时限：建库至2026年5月。语言：中文及英文（中文数据库中刊载的英文论文一并纳入）。以下为各数据库的完整检索式。', size=Pt(12))

# CNKI
add_heading_styled('2.1  中国知网 (CNKI) — 专业检索式', 3)
add_para('在CNKI“高级检索”页面切换至“专业检索”选项卡，输入以下检索式：', size=Pt(12))
add_code_block(
    "SU=('快速伸缩复合训练' + '增强式训练' + '超等长训练' + '爆发力训练'\n"
    "    + '跳深训练' + 'plyometric' + '跳箱训练' + '反应力量')\n"
    "* ('反向纵跳' + '下蹲跳' + 'CMJ' + '纵跳高度' + '垂直跳'\n"
    "    + '纵跳' + '反向跳' + 'countermovement jump')\n"
    "* ('随机' + '对照' + 'RCT' + '随机分组' + '随机对照试验')",
    size=Pt(9.5))
add_para('SU=主题（含标题+摘要+关键词）。"+" 表示OR，"*" 表示AND。检索范围为全部期刊 + 硕博论文 + 会议论文。', size=Pt(10), indent_first_line=None)

doc.add_paragraph()

# WanFang
add_heading_styled('2.2  万方 (WanFang) — 专业检索式', 3)
add_para('在万方“高级检索”页面切换至“专业检索”模式，输入以下检索式：', size=Pt(12))
add_code_block(
    '主题:("快速伸缩复合训练" OR "增强式训练" OR "超等长训练"\n'
    '     OR "plyometric" OR "跳深训练" OR "爆发力训练")\n'
    'AND 主题:("反向纵跳" OR "CMJ" OR "纵跳高度" OR "下蹲跳"\n'
    '     OR "垂直跳" OR "纵跳" OR "countermovement jump")\n'
    'AND 主题:("随机" OR "对照" OR "RCT" OR "随机分组"\n'
    '     OR "随机对照试验")',
    size=Pt(9.5))
add_para('万方专业检索使用冒号语法。检索范围选择“学术论文”+“学位论文”+“会议论文”。“为中国科技核心期刊”复选框不勾选（以覆盖更多来源）。', size=Pt(10), indent_first_line=None)

doc.add_paragraph()

# VIP
add_heading_styled('2.3  维普 (VIP) — 高级检索', 3)
add_para('维普不支持专业检索式语法，需在“高级检索”界面逐行填入（逻辑关系全部设为AND）：', size=Pt(12))

table2 = doc.add_table(rows=4, cols=4, style='Table Grid')
h2 = ['行号', '检索字段', '检索词', '逻辑']
for i, h in enumerate(h2):
    cell = table2.rows[0].cells[i]
    cell.text = ''
    r = cell.paragraphs[0].add_run(h)
    set_font(r, size=Pt(11), bold=True)

vip_rows = [
    ['1', 'M=题名或关键词',
     '快速伸缩复合训练 OR 增强式训练 OR 超等长训练 OR plyometric OR 跳深训练 OR 爆发力训练 OR 跳箱训练',
     'AND'],
    ['2', 'M=题名或关键词',
     '反向纵跳 OR 下蹲跳 OR CMJ OR 纵跳高度 OR 垂直跳 OR 纵跳 OR countermovement jump OR 反向跳',
     'AND'],
    ['3', 'M=题名或关键词',
     '随机 OR 对照 OR RCT OR 随机分组 OR 随机对照试验',
     '—'],
]
for i, row_data in enumerate(vip_rows):
    for j, cell_text in enumerate(row_data):
        cell = table2.rows[i+1].cells[j]
        cell.text = ''
        r = cell.paragraphs[0].add_run(cell_text)
        set_font(r, size=Pt(9))

doc.add_paragraph()

# ==== 3 ====
add_heading_styled('3  中英文关键词对照表', 2)
add_para('由于中文体育学术文献中术语使用不统一，以下对照表列出了所有需检索的同义词/近义表达，以确保检索的全面性。', size=Pt(12))

table3 = doc.add_table(rows=7, cols=3, style='Table Grid')
h3 = ['英文标准术语', '中文规范译名', '中文同义词/异名（均需检索）']
for i, h in enumerate(h3):
    cell = table3.rows[0].cells[i]
    cell.text = ''
    r = cell.paragraphs[0].add_run(h)
    set_font(r, size=Pt(10.5), bold=True)

kw_data = [
    ['Plyometric Training',
     '快速伸缩复合训练\n（田麦久《运动训练学》标准译名）',
     '增强式训练、超等长训练、爆发力训练\n跳深训练、plyometric训练、跳箱训练\n反应力量训练、弹跳训练'],
    ['Countermovement Jump\n(CMJ)',
     '反向纵跳',
     '下蹲跳、CMJ、反向跳、反向运动跳\n带反向的垂直跳、有预蹲的纵跳\ncountermovement jump'],
    ['Vertical Jump (VJ)',
     '纵跳 / 垂直跳',
     '垂直起跳、纵跳高度、原地纵跳\n摸高、弹跳高度、跳跃高度'],
    ['Randomized Controlled Trial (RCT)',
     '随机对照试验',
     '随机分组、随机对照、随机\nRCT、随机化实验、随机设计'],
    ['Stretch-Shortening Cycle (SSC)',
     '牵张-缩短循环',
     '拉长-缩短周期、SSC\n超等长收缩、弹性能循环'],
    ['Jump Height',
     '跳跃高度 / 纵跳高度',
     '跳高高度、腾空高度、CMJ高度\n起跳高度、飞行高度'],
]
for i, row_data in enumerate(kw_data):
    for j, cell_text in enumerate(row_data):
        cell = table3.rows[i+1].cells[j]
        cell.text = ''
        r = cell.paragraphs[0].add_run(cell_text)
        set_font(r, size=Pt(9))

doc.add_paragraph()

# ==== 4 ====
add_heading_styled('4  检索执行记录', 2)
add_para('下表为检索执行时建议填写的记录模板，以确保检索过程可复现：', size=Pt(12))

table4 = doc.add_table(rows=4, cols=5, style='Table Grid')
h4 = ['数据库', '检索日期', '检索结果数', '标题/摘要初筛后', '备注']
for i, h in enumerate(h4):
    cell = table4.rows[0].cells[i]
    cell.text = ''
    r = cell.paragraphs[0].add_run(h)
    set_font(r, size=Pt(10.5), bold=True)

rec_rows = [
    ['CNKI中国知网', '______', '______', '______', '含期刊+硕博+会议'],
    ['万方', '______', '______', '______', '含期刊+学位+会议'],
    ['维普', '______', '______', '______', '期刊论文为主'],
]
for i, row_data in enumerate(rec_rows):
    for j, cell_text in enumerate(row_data):
        cell = table4.rows[i+1].cells[j]
        cell.text = ''
        r = cell.paragraphs[0].add_run(cell_text)
        set_font(r, size=Pt(10))

doc.add_paragraph()

# ==== 5 ====
add_heading_styled('5  纳入筛选说明', 2)

notes = [
    '1. 对于中文数据库中检索到的文献，应用与英文检索相同的纳入/排除标准（PICOS框架）。CMJ手臂位置识别标准：中文文献中“双手叉腰”“双臂交叉胸前”“无摆臂”“无手臂摆动”视为严格手叉腰无臂CMJ；“自由摆臂”“自然摆臂”“带臂”“Abalakov”“摸高”等视为CMJA/带臂或臂位未明。',

    '2. 中文文献中“增强式训练”“超等长训练”等术语可能与Plyometric Training并非完全对应——部分中文研究将一般性跳跃训练（非SSC机制）也标记为“增强式训练”。筛选时需仔细阅读方法部分的训练动作描述，确认是否包含跳深、栏架跳、连续纵跳等典型SSC训练动作，将不符合者排除。',

    '3. 中文数据库中重复收录（如同一篇论文在CNKI和万方均有收录）的情况，以首次检索到的记录为准，后续重复记录在筛选流程中标记去重。',

    '4. CNKI学位论文（硕博论文）需注意：若同一研究团队基于同一数据集的学位论文和期刊论文均被检索到，优先纳入同行评审的期刊论文；若学位论文提供了更完整的数据（如SD值），则纳入学位论文版本，并将期刊论文标记为重复。',
]
for note in notes:
    add_para(note, size=Pt(12))

doc.add_paragraph()

# ==== Footer ====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— 本文件为论文《快速伸缩复合训练对反向纵跳高度影响的系统综述与Meta分析》的补充检索方案 —')
set_font(r, size=Pt(9), cjk=CJK_FONT)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('PROSPERO注册号：CRD420261422906 | 目标期刊：北京体育大学学报')
set_font(r2, size=Pt(9), cjk=CJK_FONT)

# Save
output_path = r'D:\桌面\中文数据库补充检索策略.docx'
doc.save(output_path)
print(f'Done: {output_path}')
