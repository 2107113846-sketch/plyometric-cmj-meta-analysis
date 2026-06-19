"""Generate the third-round review report as a Word document on Desktop."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(4)

def h(text, level):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return hd

def p(text, bold=False, size=Pt(11), indent=None, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return para

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(hd)
        run.font.size = Pt(9)
        run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()
    return table

RED = RGBColor(0xCC, 0x00, 0x00)
ORANGE = RGBColor(0xCC, 0x55, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)

# ===== TITLE PAGE =====
p('', size=Pt(8))
p('Plyometric 训练对 CMJ 高度影响的 Meta 分析', bold=True, size=Pt(16))
p('第三轮多 Agent 独立评审 + LLM 决策框架', bold=True, size=Pt(12))
p('')
p('审查日期：2026年6月18日', size=Pt(10))
p('审查框架：6维度专业Agent并行审查 + 综合编辑LLM决策框架', size=Pt(10))
p('审查维度：统计方法学 | 数据质量 | 科学解释 | PRISMA报告规范 | 语言呈现 | 致命缺陷', size=Pt(10))
p('目标期刊：JSCR / J Sports Sci / Biol Sport / Sports Med', size=Pt(10))
p('论文路径：D:\\桌面\\Plyo训练CMJ高度Meta分析_初稿.docx', size=Pt(10))
p('审查文件数：20+ 个源文件（手稿 + CSV数据 + JSON输出 + Python/R分析代码）', size=Pt(10))
doc.add_page_break()

# ===== EXECUTIVE SUMMARY =====
h('执行摘要', 1)

p('总体评估：需要重大修改 (MAJOR REVISION)', bold=True, size=Pt(12))

tbl(
    ['类别', '数量', '说明'],
    [
        ['🔴 致命缺陷（会导致拒稿）', '4', '每个都独立足以拒稿'],
        ['🟠 重大问题（投稿前必须修复）', '12', '系统性的PRISMA合规差距'],
        ['🟡 次要问题（可自行决定修复）', '10', '影响质量但不影响发表'],
        ['🔵 改进建议', '6', '提升竞争力'],
    ]
)

p('核心判断：', bold=True)
p('本研究针对一个重要临床问题，具有真正创新的方法论贡献（手臂位置分层）。分析方法——含r敏感性的变化分数SMD、REML合并、Q-profile I² CI、预测区间、全面的亚组和Meta回归分析以及GRADE评估——比大多数运动科学Meta分析更为严格。然而，当前版本的手稿存在4个致命缺陷，每个都独立足以在编辑筛选阶段被拒稿。基础研究是可靠的，分析方法学是扎实的。问题出在报告和验证环节，而非研究设计或数据质量。经过彻底修订，本手稿有望对Plyometric训练文献做出有意义的贡献。')

doc.add_page_break()

# ===== FATAL FLAWS =====
h('一、致命缺陷（FF1-FF4：每个都独立会导致拒稿）', 1)

# FF1
h('FF1. I²置信区间在数学上不可能', 2)
p('位置：手稿第215行、第219行', bold=True)
p('问题描述：手稿报告的I²点估计落在其自身95%置信区间之外，这在数学上不可能——点估计必须位于由同一模型计算出的CI之内。')

tbl(
    ['分析池', '手稿报告I²', '手稿报告CI', '问题'],
    [
        ['严格手叉腰池 (line 215)', '76.2%', '[54.6%, 70.2%]', '76.2% 在 [54.6%, 70.2%] 之外！'],
        ['宽版全池 (line 219)', '66.9%', '[45.7%, 61.9%]', '66.9% 在 [45.7%, 61.9%] 之外！'],
    ]
)

p('根因分析：', bold=True)
p('• I²点估计（76.2%和66.9%）来自Analysis_Report.txt，该文件使用了REML via R/metafor对analysis_ready_effects.csv中的效应量进行计算', indent=0.5)
p('• 但CI [54.6%, 70.2%]和[45.7%, 61.9%]是从sensitivity_r_results.json复制的——该文件由独立的sensitivity_r_corr.py运行生成，重新计算了效应量并产生了不同的I²点估计（69.7%和63.4%）', indent=0.5)
p('• 作者从一次分析运行中取I²点估计，又从另一次不兼容的运行中取CI，拼接了两者', indent=0.5)

p('影响：致命。任何统计审稿人或期刊编辑会立即识别此错误，并将其作为结果未经验证的证据。', color=RED, bold=True)
p('修正方案：使用R/metafor的confint() Q-profile方法，在产生主分析结果的同一REML模型上重新计算所有I² 95% CI。修正后的CI约为Strict [63.3%, 94.5%]，Wide [59.8%, 90.7%]——但必须在最终16研究严格池上重新计算。点估计和CI都必须来自同一模型拟合。', color=RED)

# FF2
h('FF2. 总体样本量错误', 2)
p('位置：手稿第162行（以及所有引用总样本量的地方）', bold=True)

tbl(
    ['来源', '干预组 N', '对照组 N', '总计'],
    [
        ['手稿第162行', '390', '342', '732'],
        ['analysis_ready_effects.csv实际值', '367', '351', '718'],
        ['差异', '+23', '-9', '+14'],
    ]
)

p('影响：这个错误传播到摘要、结果和GRADE证据概要中关于总样本量的每一个陈述。它从基础层面削弱了可信度。', color=RED, bold=True)
p('修正方案：(a) 全文统一使用分析N（推荐），或(b) 记录哪些研究存在随机化-vs-分析N的差异，并明确说明报告的是哪种N。', color=RED)

# FF3
h('FF3. GRADE表样本量严重错误', 2)
p('位置：手稿第357行（GRADE证据概要）', bold=True)

tbl(
    ['分析池', '手稿报告 N', 'CSV实际 N', '误差'],
    [
        ['严格手叉腰池', 'N=520', 'N=341', '+179（高估52%！）'],
        ['宽版全池', 'N=718', 'N=702', '+16'],
    ]
)

p('严格池的误差尤其严重：手稿将分析的样本量高估了超过一半。16项严格手臂研究的实际参与者为341名（IG=178, CG=163）。N=520无法从数据的任何合理解释中推导出来。宽版池的误差（+16）对应R24（Rubley 2011，VJ带臂研究，n=16）被错误地计入了CMJ宽版池。', indent=0.5)

p('影响：致命。GRADE评估是手稿自身对证据确定性的总结——当基础样本计数偏差52%时，读者无法信任证据评级。', color=RED, bold=True)
p('修正方案：直接从analysis_ready_effects.csv重新计算所有GRADE的N值。通过求和该池中各研究的IG和CG列来验证每个分析池的参与者数量。在补充材料中显式展示计数。', color=RED)

# FF4
h('FF4. 参考文献书目数据存在复制粘贴错误', 2)
p('位置：参考文献列表条目第32项和第33项', bold=True)

tbl(
    ['条目', '手稿内容', '实际内容（PubMed验证）'],
    [
        ['Ref 32 (Ramirez-Campillo 2015, R02)', 'JSCR 29(5):1294-1303', 'JSCR 29(5):1317-1328'],
        ['Ref 33 (Ramirez-Campillo 2015, R03/R05)', 'JSCR 29(5):1294-1303, 年份2015', 'JSCR 28(1):97-104, 年份2014, DOI不同'],
    ]
)

p('这两条参考文献在手稿中共享相同的期刊、卷、期和页码——这是一个明显的复制粘贴错误。它们是不同的论文，有不同的DOI。此外，15项已检查的纳入研究参考文献中有9项缺少DOI。', indent=0.5)

p('影响：致命。参考文献列表不可信任。', color=RED, bold=True)
p('修正方案：根据PubMed/CrossRef记录验证每一条参考文献。确保所有条目都有DOI。纠正Ref 32和Ref 33的克隆书目数据。', color=RED)

doc.add_page_break()

# ===== MAJOR ISSUES =====
h('二、重大问题（M1-M12：投稿前必须修复）', 1)

major_issues = [
    ('M1', 'r敏感性表使用过时且不兼容的数据', '第251行',
     'r敏感性表来自sensitivity_r_results.json，该文件由独立的Python运行生成。即使使用相同的r=0.7，也产生了不同的合并估计值：主分析g=+1.128 vs 敏感性表g=+1.177, I²=76.2% vs 69.7%。r=0.7行应完全复现主分析结果。这使得敏感性分析在内部逻辑上不自洽。',
     '在同一R会话中重新运行r敏感性分析，或在r=0.7行中显式复现主分析值。'),
    ('M2', 'PRISMA流程图算法内部矛盾', '第154行',
     '算法：135（识别）- 44（标题/摘要排除）- 60（全文排除）= 31，而非29。有两项研究在流程图中未被解释。',
     '调整全文排除数为62，或调整初始识别数，或在流程图中添加一个可见的类别（例如"筛选过程中移除的2篇重复"）。每个数字必须协调。'),
    ('M3', '手稿缺少结构化摘要', '整篇手稿',
     '审核的手稿文件仅涵盖方法、结果和讨论部分。投稿需要遵循PRISMA 2020摘要检查表的结构化摘要，但未提供。',
     '添加遵循PRISMA 2020摘要检查表的结构化摘要（标题、目的、纳入标准、信息来源、偏倚风险、综合方法、纳入研究、结果、解读、资金/注册）。'),
    ('M4', '4个数据库中有3个的检索策略缺失', '第27-33行',
     '仅提供了PubMed检索策略。Scopus、Web of Science和Google Scholar的完整可复现策略缺失。PRISMA 2020第5项和PRISMA-S扩展要求所有数据库的完整检索策略。',
     '在补充材料中提供Scopus、Web of Science和Google Scholar的完整布尔检索策略。解释鉴于其可复现性限制，Google Scholar结果是如何筛选的。'),
    ('M5', '双重筛选过程未描述', '方法部分（第59-66行）',
     '手稿描述了独立的双重数据提取（附kappa值），但研究筛选过程（标题/摘要筛选、全文筛选）从未被描述。PRISMA 2020第6项要求：(a) 每条记录由多少名审核者筛选，(b) 筛选是否独立，(c) 冲突解决程序。',
     '新增"研究筛选"小节，描述筛选方法（筛选者人数、独立性、冲突解决程序）。'),
    ('M6', '缺少资金和利益冲突声明', '手稿末尾',
     '没有资金声明（PRISMA 2020第21项）和利益竞争声明（PRISMA 2020第22项）。ICMJE标准和PRISMA 2020均要求这两项。',
     '添加两项声明。即使没有资金，也应声明"本研究未获得任何资助机构的专项资助"。即使没有利益冲突，也应声明"作者声明无利益竞争"。'),
    ('M7', 'Meta回归p值内部矛盾', '表格第304行 vs 讨论第311行、第398行',
     'Meta回归表正确报告Duration的p=0.023，但讨论第311行和第398行均写成p=0.026。此外，z值舍入存在小错误：0.131/0.058=2.2586，应舍入为2.26，而非2.27。',
     '将讨论中两处p值更正为0.023。将z舍入为2.26（或更精确地重新计算）。'),
    ('M8', '参考文献格式不标准且不可验证', '第508-610行',
     '参考文献列表使用多个独立编号块（纳入研究16-46、既往Meta分析1-4、理论基础5-11、方法学12-14、临床意义15-16），产生重复的参考编号。正文引用使用[作者, 年份]格式，但与编号列表无对应关系，使验证变得不可能。',
     '使用单一连续编号系统（例如[1]-[N]）。确保每个正文引用明确映射到参考文献列表中的一条。遵循目标期刊的参考文献格式。'),
    ('M9', 'R11 Sedano Campo个体效应量使用了错误的度量指标', '第441行',
     '讨论引用R11的个体g为+3.42（仅后测SMD）。手稿中所有其他效应量都使用变化分数SMD（Hedges g），按此标准R11的g=+5.20。这使得引用的个体值无法与合并估计值或任何其他引用的研究效应量进行比较。',
     '一致地引用变化分数SMD（g=+5.20），或明确标注使用的是哪种度量指标及原因。'),
    ('M10', '年龄亚组分析遗漏了一项研究', '第281-287行',
     '年龄亚组分析总计2+2+13+10=27项研究。宽版池有28项CMJ研究。R06（Van Roie 2020，老年成年人，n=27，69-70岁）出现在表1中但被排除在年龄亚组分析之外，且无解释。',
     '将R06纳入年龄亚组分析，归入"老年"组k=1，或明确记录排除原因。'),
    ('M11', '训练类型k计数在手稿和分析报告间不匹配', '第294行 vs Analysis_Report第43-44行',
     '手稿：纯Plyo=25项研究，Plyo+力量=2项，总计=27。分析报告：纯Plyo=26项，混合=2项，总计=28。',
     '协调有争议研究的分类（可能是R19 Michailidis 2018, Plyo+COD）。确保所有研究计数正确求和为28。'),
    ('M12', 'PEDro百分比错误', '第186行',
     '"超过半数（55%）的纳入研究处于良好及以上质量区间"——实际：4项优秀（8分）+ 12项良好（6-7分）= 28项CMJ研究中的16项 = 57%，而非55%。若包含VJ研究R24：29项中的17项 = 59%。无论哪种情况，55%都是错误的。',
     '重新计算并报告正确的百分比（57%或59%）。'),
]

for mid, title, location, problem, fix in major_issues:
    h(f'{mid}. {title}', 2)
    p(f'位置：{location}', bold=True)
    p(f'问题：{problem}')
    p(f'修正：{fix}', color=ORANGE)
    p('')

doc.add_page_break()

# ===== MINOR ISSUES =====
h('三、次要问题（m1-m10：可自行决定修复）', 1)

minor_issues = [
    ('m1', 'GRADE评级内部不一致',
     '第356-362行',
     '严格池（k=16, I²=76%）评级为"低"，而青春期前亚组（k=2, I²=0%）评级为"中等"。k=2时，CI极宽[0.77, 1.98]，I²=0%在统计上无意义。青春期前亚组或许应为"极低"（精确性严重不足降2级）。'),
    ('m2', 'PROSPERO注册号无法公开验证',
     '第15行',
     '注册号CRD420261422906无法验证。如果记录正在等待批准，请明确说明。如果已获批准，请确保公众可以访问。'),
    ('m3', '方案访问信息不足',
     '方法部分',
     'PRISMA 2020第20项要求报告方案的可访问位置。如果PROSPERO是唯一方案，请明确说明。如果有单独方案文档，请提供访问信息。'),
    ('m4', 'PEDro文件中"I2"符号存在歧义',
     'PEDro_Score_Final.csv, Quality列',
     '符号"良好/I2=0"和"良好/I2=1"中的"I2"指的是PEDro第2项（随机分配）。这与手稿中显著报告的异质性I-squared统计量完全相同，容易造成混淆。重命名为"Item2"或"Rand"。'),
    ('m5', '连续变量使用了Kappa统计量',
     '第66行',
     '手稿报告了提取一致性的kappa值（0.89-0.95），但kappa适用于分类数据。对于连续变量应使用ICC或Bland-Altman一致性限。96.7%的数值一致性统计量的分析单位不明确。'),
    ('m6', 'R09 Blazevich 2003纳入资格问题',
     '严格池组成',
     'R09的PEDro=3，筛选注释注明"非随机设计(quasi-RCT)"。如果确实是非随机的，则违反了规定的仅RCT的纳入标准。可从原文核实随机化情况，或排除该研究，或保留但承认准随机设计作为局限性。'),
    ('m7', '持续时间亚组k计数分类错误',
     'Analysis_Report第27-29行',
     '报告显示短期=14，中期=10，长期=4。实际编码：短期=15，中期=9，长期=4。一项研究（可能是R09 Blazevich 2003, 5周）被错误分类。'),
    ('m8', 'R04手臂分类字段标注错误',
     'analysis_ready_effects.csv, R04行',
     'R04的cmj_arm字段显示"未明示"，但data_note确认了严格的无臂CMJ，且所有分析输出均正确将该研究置于严格池。将CSV字段更正为"手叉腰无臂"。'),
    ('m9', 'Egger检验显著性阈值不一致',
     '方法部分 vs 结果部分',
     '方法部分声明Egger阈值p<0.10，但结果部分报告p<0.001，且未参考预设阈值进行解读。需明确说明结果是否满足预设阈值。'),
    ('m10', 'R24 Rubley 2011出现在宽版池RevMan文件中',
     'RevMan_ready_wide_pool.csv',
     'R24（VJ带臂，14周，g=1.53）出现在宽版池RevMan文件中（29行）。分析报告声明宽版池=28项RCT。如果Meta分析是从此RevMan文件运行的，则R24的VJ数据被错误地纳入了CMJ宽版分析。'),
]

for mid, title, location, desc in minor_issues:
    h(f'{mid}. {title}', 2)
    p(f'位置：{location}', bold=True)
    p(f'说明：{desc}')
    p('')

doc.add_page_break()

# ===== SUGGESTIONS =====
h('四、改进建议（S1-S6：提升论文竞争力）', 1)

suggestions = [
    ('S1', '重新表述"方法学精度"主张',
     '第15行、第382-383行、第392行',
     '排除带臂CMJ提供的是"更精确"估计的主张混淆了方法学特异性与统计精度。严格池的CI更宽（统计精度更低）、生态效度更低。重新表述为："严格池提供了对下肢SSC训练效果更具特异性的估计，但在推广性和统计精度方面需要权衡。"'),
    ('S2', '限定"协同发育窗口"主张',
     '第288-289行、第406-407行',
     '该主张基于每个青少年亚组仅k=2项研究，且I² CI跨越[0%, 80%]。应表述为产生假说："青少年亚组的效应量数值较大，但每个亚组仅2项研究，这些估计是初步的，需要复现验证。"'),
    ('S3', '承认Harman 1990的局限性',
     '第13行、第392行、第417行',
     '10-15%手臂对CMJ高度的贡献引用自一项36年前的单实验室研究（n=17）。承认：(a) 这是单一研究估计，(b) 手臂摆动会定性地改变运动协调性，(c) 10-15%可能不能推广至Meta分析中的所有人群。'),
    ('S4', '解释为什么异质性高于预期',
     '讨论（第390-396行附近）',
     '本研究的I²=76%大约是Markovic 2007年I²=11%的7倍，尽管纳入了所谓更均质的标准。这一悖论值得解释。可能的因素：不同的结局度量指标（变化分数vs仅后测SMD）、不同的人群、不同的分析时代标准。'),
    ('S5', '加强对发表偏倚的讨论',
     '第320-346行',
     '手稿目前将Trim-and-Fill解释（右侧填充意味着真实异质性）前置，而将Egger检验（p<0.001）和SE-g相关性（r=0.88）后置。应将两者呈现为互补："发表偏倚信号强烈，但潜在偏倚方向不确定。"引用Sterne et al. 2011（已在参考文献中），关于k<25时统计检验优先于视觉漏斗图检查的建议。'),
    ('S6', '从经验上证明MCID的SD范围',
     '第429行',
     '用于MCID转换的4-6 cm合并SD范围是近似的。要么从短期亚组数据计算实际的加权合并SD，要么明确将4-6 cm标注为示意性近似值。'),
]

for sid, title, location, desc in suggestions:
    h(f'{sid}. {title}', 2)
    p(f'位置：{location}', bold=True)
    p(desc)
    p('')

doc.add_page_break()

# ===== CROSS-REFERENCE MATRIX =====
h('五、问题交叉引用矩阵', 1)

tbl(
    ['问题', '摘要', '方法', '结果：合并', '结果：亚组', '发表偏倚', 'GRADE', '讨论', '表格/参考文献', '数据文件'],
    [
        ['FF1 (I² CI不可能)', '—', '—', '✓ L215,219', '—', '—', '—', '—', '—', '✓'],
        ['FF2 (GRADE N错误)', '—', '—', '—', '—', '—', '✓ L357', '—', '—', '✓'],
        ['FF3 (总N错误)', '✓', '—', '✓ L162', '—', '—', '✓ L357', '—', '—', '✓'],
        ['FF4 (参考文献错误)', '—', '—', '—', '—', '—', '—', '—', '✓', '—'],
        ['M1 (r敏感性过时)', '—', '—', '✓ L251', '—', '—', '—', '—', '✓', '✓'],
        ['M2 (PRISMA流程)', '—', '—', '✓ L154', '—', '—', '—', '—', '—', '✓'],
        ['M3 (无摘要)', '✓', '—', '—', '—', '—', '—', '—', '—', '—'],
        ['M4 (缺失检索策略)', '—', '✓ L27-33', '—', '—', '—', '—', '—', '—', '—'],
        ['M5 (无筛选描述)', '—', '✓ L59-66', '—', '—', '—', '—', '—', '—', '—'],
        ['M6 (无资金/COI)', '—', '—', '—', '—', '—', '—', '—', '✓', '—'],
        ['M7 (p值笔误)', '—', '—', '—', '—', '—', '—', '✓ L311,398', '✓', '—'],
        ['M8 (参考文献格式)', '—', '—', '—', '—', '—', '—', '—', '✓', '—'],
        ['M9 (R11 SMD类型)', '—', '—', '—', '—', '—', '—', '✓ L441', '—', '✓'],
        ['M10 (R06缺失)', '—', '—', '—', '✓ L281-287', '—', '—', '—', '—', '✓'],
        ['M11 (k计数不匹配)', '—', '—', '—', '✓ L294', '—', '—', '—', '—', '✓'],
        ['M12 (PEDro %错误)', '—', '✓ L186', '—', '—', '—', '—', '—', '—', '✓'],
    ]
)

p('✓ = 受影响的章节，L = 行号，— = 不受影响', size=Pt(9))

doc.add_page_break()

# ===== STRENGTHS =====
h('六、正面评价（值得保留的核心优势）', 1)

strengths = [
    'CMJ手臂位置分层是真正的创新——首次将其作为纳入标准和分层分析变量，为运动科学Meta分析树立了新的方法学标准',
    '预测区间(Prediction Interval, PI)的报告和讨论在运动科学Meta分析中不常见，且对PI跨零的解读诚实且具有实践指导意义',
    'GRADE评估框架完整且理由充分（虽然个别评级和样本计数需要修正）',
    'r敏感性分析范围全面（0.5-0.9），诚实报告了效应量对r假设的中等程度敏感性（Δg≈+0.95），证实了结论在所有r假设下的方向稳健性',
    'Trim-and-Fill方向分析显示了对发表偏倚的深入思考——正确区分了经典发表偏倚（左侧填充）和小研究真实效应差异（右侧填充）',
    '训练安全性讨论（4.6节）的存在本身值得肯定——许多运动科学Meta分析完全忽略这一维度',
    '代码公开（GitHub）+ PROSPERO预注册体现了开放科学态度，为结果复现提供了完整路径',
    '多臂研究取组逻辑透明且有原则（优先取标准方案/保守选择），避免了选择性报告',
    '效应量计算使用了更精确的Pre-post change SMD方法（含小样本Hedges g校正），比大多数运动科学Meta分析仅使用Post-only SMD更合理',
    'Inter-rater reliability报告（kappa=0.89/0.92/0.95，一致性率96.7%）是良好的方法学实践，增强了数据提取的可信度',
]

for i, s in enumerate(strengths, 1):
    p(f'{i}. {s}')

doc.add_page_break()

# ===== FINAL VERDICT =====
h('七、最终裁决', 1)

tbl(
    ['维度', '评估'],
    [
        ['总体建议', '重大修改 (MAJOR REVISION)'],
        ['致命缺陷数量', '4（每个独立足以拒稿）'],
        ['重大问题数量', '12（系统性的PRISMA合规差距）'],
        ['研究问题重要性', '★★★★★ — 高度临床相关，填补关键证据空白'],
        ['方法学质量', '★★★★☆ — 分析方法扎实，但数字报告存在系统性验证问题'],
        ['科学贡献潜力', '★★★★☆ — 手臂位置分层是真正的创新'],
        ['当前可发表性', '★★☆☆☆ — 致命缺陷阻碍发表'],
        ['修正后可发表性', '★★★★☆ — 修订后有望在JSCR/Sports Med级别期刊发表'],
    ]
)

p('')
p('裁决理由：', bold=True, size=Pt(12))
p('本手稿解决了一个定义明确的研究问题，并做出了真正创新的方法论贡献。分析方法——变化分数SMD、REML合并、Q-profile I² CI、预测区间、全面的亚组和Meta回归分析以及GRADE评估——比大多数运动科学Meta分析更为严格。')
p('然而，当前版本的手稿不可发表。四个致命缺陷中的每一个都独立足以拒稿：')
p('1. FF1（I² CI在数学上不可能）表明数值结果在未经验证的情况下在不兼容的分析运行之间被复制。这动摇了每个报告数字的可信度。', indent=0.5)
p('2. FF2-FF3（样本量错误）影响摘要、结果和GRADE证据概要——任何Meta分析中被阅读最多也最为关键的章节。', indent=0.5)
p('3. FF4（参考文献错误）表明即使是最基本的书目数据也未被验证。', indent=0.5)
p('重大问题和次要问题反映了投稿前必须解决的系统性PRISMA 2020合规差距。其中若干项（缺少摘要、缺少检索策略、未记录的筛选过程、缺失资金/COI声明）是明确的PRISMA 2020违规项，无论分析质量如何，都会阻碍发表。')
p('')
p('投稿前必须采取的五项行动：', bold=True)
p('1. 在单一版本化分析流程中，从definitive的analysis_ready_effects.csv重新计算每一个数值结果。绝不在不同分析运行之间拼接结果。', indent=0.5)
p('2. 对照重新计算的结果，重新验证手稿中的每一个数字，特别关注样本量、I²值及其CI、所有GRADE的N值以及所有亚组的k计数。', indent=0.5)
p('3. 完成PRISMA 2020检查清单：添加结构化摘要，提供所有数据库检索策略，描述双重筛选过程，添加资金和COI声明。', indent=0.5)
p('4. 修正参考文献列表：对照PubMed验证所有条目，添加缺失的DOI，修复克隆的书目数据（Ref 32和Ref 33）。', indent=0.5)
p('5. 协调PRISMA流程图的算法（135-44-62=29，或等效调整）。', indent=0.5)
p('')
p('基础研究是可靠的，分析方法学是扎实的。问题出在报告和验证环节，而非研究设计或数据质量。经过彻底修订，本手稿有望对Plyometric训练文献做出有意义的贡献。', bold=True, size=Pt(11))

doc.add_page_break()

# ===== APPENDIX: FILES EXAMINED =====
h('附录：审核中检查的文件', 1)

files = [
    'D:\\桌面\\科研训练\\analysis_ready_effects.csv',
    'D:\\桌面\\科研训练\\output\\Analysis_Report.txt',
    'D:\\桌面\\科研训练\\output\\Table1_Study_Characteristics.csv',
    'D:\\桌面\\科研训练\\output\\PEDro_Score_Final.csv',
    'D:\\桌面\\科研训练\\output\\PRISMA_flow.json',
    'D:\\桌面\\科研训练\\output\\i2_ci_corrected.json',
    'D:\\桌面\\科研训练\\output\\sensitivity_r_results.json',
    'D:\\桌面\\科研训练\\output\\trimfill_results.json',
    'D:\\桌面\\科研训练\\output\\RevMan_ready_strict_pool.csv',
    'D:\\桌面\\科研训练\\output\\RevMan_ready_wide_pool.csv',
    'D:\\桌面\\科研训练\\output\\Manuscript_Draft_Methods_Results.md',
    'D:\\桌面\\科研训练\\screening_merged_30studies.csv',
    'D:\\桌面\\科研训练\\clean_and_compute_effects.py',
    'D:\\桌面\\科研训练\\sensitivity_analysis.py',
    'D:\\桌面\\科研训练\\sensitivity_r_corr.py',
    'D:\\桌面\\科研训练\\meta_regression.py',
    'D:\\桌面\\科研训练\\publication_bias.py',
    'D:\\桌面\\科研训练\\subgroup_analysis.py',
    'D:\\桌面\\科研训练\\meta_toolkit\\pooling.py',
    'D:\\桌面\\科研训练\\meta_toolkit\\effects.py',
    'D:\\桌面\\科研训练\\meta_toolkit\\r_bridge.py',
    'D:\\桌面\\科研训练\\meta_toolkit\\run_meta.R',
]

for f in files:
    p(f'• {f}', size=Pt(9))

# ===== SAVE =====
output_path = 'D:/桌面/Plyo_CMJ_Meta分析_第三轮综合审查报告.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
