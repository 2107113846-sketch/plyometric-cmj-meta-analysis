# -*- coding: utf-8 -*-
"""
Generate v3 manuscript .docx from updated _v3_text.txt
Reads the master text file as source of truth, and builds a formatted Word document.
"""
import sys, json, os, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJ = Path(__file__).parent
OUTPUT_DIR = PROJ / 'output'
STYLE_FONT = 'Times New Roman'
CJK_FONT = '宋体'
HEADING_FONT = '黑体'

# ================================================================
# Load text source
# ================================================================
with open(PROJ / '_v3_text.txt', 'r', encoding='utf-8') as f:
    full_text = f.read()

lines = full_text.split('\n')

# ================================================================
# Formatting helpers
# ================================================================
def set_run_font(run, size=None, bold=False, italic=False, font_name=None, cjk=None):
    font = font_name or STYLE_FONT
    run.font.name = font
    if cjk is None:
        cjk = CJK_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), cjk)
    rPr.insert(0, rFonts)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, cjk=HEADING_FONT)
    return h

def add_para(doc, text, bold=False, size=Pt(12), align=None, first_line_indent=Cm(0.74),
             spacing_after=Pt(6), spacing_before=Pt(0)):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = spacing_after
    para.paragraph_format.space_before = spacing_before
    para.paragraph_format.first_line_indent = first_line_indent
    # Handle **bold** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, size=size, bold=bold)
    return para

def add_simple_para(doc, text, bold=False, size=Pt(12), indent_left=None, first_line=Cm(0.74),
                    spacing_after=Pt(6)):
    """Add a paragraph with optional left indent and first-line indent control."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = spacing_after
    if first_line is not None:
        para.paragraph_format.first_line_indent = first_line
    if indent_left is not None:
        para.paragraph_format.left_indent = indent_left
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, size=size, bold=bold)
    return para

def add_figure(doc, img_path, caption, width=Inches(5.5)):
    if os.path.exists(img_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run()
        run.add_picture(img_path, width=width)
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_after = Pt(12)
        cap_run = cap_para.add_run(caption)
        set_run_font(cap_run, size=Pt(10), bold=True)
    else:
        add_para(doc, f'[图片缺失: {img_path}]', size=Pt(10),
                 align=WD_ALIGN_PARAGRAPH.CENTER)

def add_table_grid(doc, headers, rows, note_rows=None):
    """Add a formatted table with optional note rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=Pt(9), bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci < len(headers):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                run = cell.paragraphs[0].add_run(str(val))
                set_run_font(run, size=Pt(9))
    doc.add_paragraph()
    return table

# ================================================================
# Parse text file into sections
# ================================================================
# We'll identify major sections by headings and content blocks
# Structure:
# Line 1: empty
# Lines 2-7: Title block
# Line 8: 摘要 heading
# Lines 9-28: Abstract content
# Line 29: 1  引言
# etc.

def get_section(start_pattern, end_pattern=None, start_after=0):
    """Extract lines between start_pattern and end_pattern."""
    start_idx = None
    end_idx = len(lines)
    for i, line in enumerate(lines):
        if i < start_after:
            continue
        if start_idx is None and start_pattern in line:
            start_idx = i + 1  # content starts after heading
        elif start_idx is not None and end_pattern and end_pattern in line:
            end_idx = i
            break
    if start_idx is None:
        return []
    # Go until next section heading or end
    if end_pattern is None:
        for i in range(start_idx, len(lines)):
            line = lines[i].strip()
            # Detect section heading (number + Chinese chars)
            if re.match(r'^(4|5|参考文献)\s', line) and i > start_idx + 5:
                end_idx = i
                break
    return lines[start_idx:end_idx]

# ================================================================
# Build Document
# ================================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = STYLE_FONT
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), CJK_FONT)
rPr.insert(0, rFonts)

# --- Title page ---
add_para(doc, '', size=Pt(12))
title = '快速伸缩复合训练对反向纵跳高度影响的系统综述与Meta分析'
subtitle = '——基于手臂位置分层的效应量估计与剂量-反应关系'
add_para(doc, title, bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER,
         first_line_indent=None, spacing_after=Pt(8))
add_para(doc, subtitle, bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER,
         first_line_indent=None, spacing_after=Pt(20))
add_para(doc, 'PROSPERO注册号：CRD420261422906', size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '西南大学体育学院', size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, spacing_after=Pt(12))
add_para(doc, '指导教师：付道领', size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

doc.add_page_break()

# ================================================================
# Abstract — parse from text
# ================================================================
add_heading_styled(doc, '摘要', 1)

# Parse abstract sections
abstract_map = {
    '目的：': ('目的：', '', True),
    '纳入标准：': ('纳入标准：', '', True),
    '信息来源：': ('信息来源：', '', True),
    '偏倚风险：': ('偏倚风险：', '', True),
    '综合方法：': ('综合方法：', '', True),
    '纳入研究：': ('纳入研究：', '', True),
    '结果：': ('结果：', '', True),
    '解读：': ('解读：', '', True),
    '资金与注册：': ('资金与注册：', '', True),
}

current_abstract_label = None
current_abstract_text = []

# Find abstract section lines
abstract_section_start = None
abstract_section_end = None
for i, line in enumerate(lines):
    if '摘要' == line.strip() and i > 5:
        abstract_section_start = i + 1
    if abstract_section_start and '关键词：' in line:
        abstract_section_end = i + 1
        break

if abstract_section_start:
    for line in lines[abstract_section_start:abstract_section_end]:
        line = line.strip()
        if not line:
            continue
        matched = False
        for label in ['目的：', '纳入标准：', '信息来源：', '偏倚风险：',
                       '综合方法：', '纳入研究：', '结果：', '解读：', '资金与注册：']:
            if line.startswith(label):
                if current_abstract_label:
                    add_simple_para(doc, current_abstract_text, bold=(current_abstract_label != ''),
                                   first_line_indent=Cm(0.74) if current_abstract_label == '' else None,
                                   spacing_after=Pt(2) if current_abstract_label != '' else Pt(6))
                    current_abstract_text = ''
                add_simple_para(doc, label, bold=True, first_line_indent=Cm(0.74) if label != '目的：' else Cm(0.74),
                               spacing_after=Pt(2))
                current_abstract_label = label
                current_abstract_text = line[len(label):].strip()
                matched = True
                break
        if not matched:
            if current_abstract_text:
                current_abstract_text += line
            else:
                current_abstract_text = line
    # Flush last
    if current_abstract_text:
        add_simple_para(doc, current_abstract_text, bold=False, first_line_indent=Cm(0.74),
                       spacing_after=Pt(6))

# Keywords
for i, line in enumerate(lines):
    if '关键词：' in line:
        add_para(doc, line.strip(), bold=True, size=Pt(11), first_line_indent=Cm(0.74),
                 spacing_before=Pt(12))
        break

doc.add_page_break()

# ================================================================
# Main body sections — parse from text file
# ================================================================
# Use a line-by-line parser that detects headings and content

COLLAPSED = False  # Track whether we're inside a collapsed link output

chapter_map = {
    '1  引言': 1,
    '2  方法': 1,
    '2.1  文献检索策略': 2,
    '2.2  纳入与排除标准': 2,
    '2.3  文献筛选': 2,
    '2.4  数据提取': 2,
    '2.5  偏倚风险评价': 2,
    '2.6  效应量计算': 2,
    '2.7  统计分析方法': 2,
    '3  结果': 1,
    '3.1  文献筛选流程': 2,
    '3.2  研究特征': 2,
    '3.3  偏倚风险评价': 2,
    '3.4  主分析：': 2,
    '3.5  敏感性分析': 2,
    '3.6  亚组分析': 2,
    '3.7  Meta回归': 2,
    '3.8  发表偏倚': 2,
    '3.9  GRADE证据质量评级': 2,
    '4  讨论': 1,
    '4.1  主要发现': 2,
    '4.2  与既往Meta分析的比较': 2,
    '4.3  干预时长-效应关系与剂量-反应意义': 2,
    '4.4  年龄/发育阶段的调节作用': 2,
    '4.5  CMJ手臂位置：方法学启示': 2,
    '4.6  训练安全性与损伤风险': 2,
    '4.7  局限性': 2,
    '4.8  实践建议与研究展望': 2,
    '5  利益冲突与资金声明': 1,
    '参考文献': 1,
}

# Find start of main body (after abstract)
body_start = 0
for i, line in enumerate(lines):
    if line.strip().startswith('1  引言'):
        body_start = i
        break

# Parse body into paragraphs
body_paragraphs = []
current_para_lines = []
in_special_block = False

for i in range(body_start, len(lines)):
    line = lines[i]
    stripped = line.strip()

    # Skip empty lines at start
    if not current_para_lines and not stripped:
        continue

    # Detect section headings
    heading_match = False
    for heading_key, heading_level in chapter_map.items():
        if stripped == heading_key or stripped.startswith(heading_key):
            if current_para_lines:
                body_paragraphs.append(('text', '\n'.join(current_para_lines).strip()))
                current_para_lines = []
            body_paragraphs.append(('heading', stripped, heading_level))
            heading_match = True
            break

    if heading_match:
        continue

    # Detect table/image markers
    if stripped.startswith('表') and (' ' not in stripped[:4] or '注：' in stripped):
        if current_para_lines:
            body_paragraphs.append(('text', '\n'.join(current_para_lines).strip()))
            current_para_lines = []
        body_paragraphs.append(('table_caption', stripped))
        continue

    if stripped.startswith('图') and ('PRISMA' in stripped or '森林' in stripped or '漏斗' in stripped or 'S1' in stripped):
        if current_para_lines:
            body_paragraphs.append(('text', '\n'.join(current_para_lines).strip()))
            current_para_lines = []
        body_paragraphs.append(('figure', stripped))
        continue

    # Handle bullet items in limitations
    if re.match(r'^\d+\.\s', stripped) and '发表偏倚不能排除' in stripped:
        if current_para_lines:
            body_paragraphs.append(('text', '\n'.join(current_para_lines).strip()))
            current_para_lines = []
        # Collect limitation item
        lim_text = [stripped]
        j = i + 1
        while j < len(lines) and not re.match(r'^\d+\.\s', lines[j].strip()) and not any(
            lines[j].strip().startswith(k) for k in chapter_map):
            if lines[j].strip():
                lim_text.append(lines[j].strip())
            j += 1
        body_paragraphs.append(('limitation', '\n'.join(lim_text)))
        continue
    elif re.match(r'^\d+\.\s', stripped):
        # Other numbered items - treat as text continuation
        pass

    current_para_lines.append(line)

# Flush remaining
if current_para_lines:
    body_paragraphs.append(('text', '\n'.join(current_para_lines).strip()))

# Now build the document from body_paragraphs
current_section = None
for item in body_paragraphs:
    if item[0] == 'heading':
        add_heading_styled(doc, item[1], item[2])
        current_section = item[1]
    elif item[0] == 'text':
        # Split by double newline to get logical paragraphs
        pgraphs = item[1].split('\n\n')
        for pg in pgraphs:
            pg = pg.strip()
            if not pg:
                continue
            # Replace single newlines with spaces within a paragraph
            pg = pg.replace('\n', '')
            # Skip empty
            if not pg.strip():
                continue
            add_simple_para(doc, pg.strip())
    elif item[0] == 'limitation':
        # Format limitation with bold prefix
        text = item[1].replace('\n', '')
        m = re.match(r'^(\d+\.\s)(.+?[。.])', text)
        if m:
            num_dot = m.group(1)
            bold_part = text[m.start(2):m.end(2)]
            rest = text[m.end(2):]
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.first_line_indent = Cm(0)
            run1 = para.add_run(f'{num_dot}{bold_part}')
            set_run_font(run1, size=Pt(12), bold=True)
            run2 = para.add_run(rest)
            set_run_font(run2, size=Pt(12))
    elif item[0] == 'table_caption':
        add_para(doc, item[1], bold=True, size=Pt(11),
                first_line_indent=None, spacing_before=Pt(8))
    elif item[0] == 'figure':
        pass  # Figures handled below

# Insert tables and figures at known positions

# Table 1 - Study Characteristics
# Table 2 - PEDro
# Table 3 - Sensitivity
# Table 4 - Arm position subgroup
# Table 5 - Duration subgroup
# Table 6 - Age subgroup
# Table 7 - Meta regression
# Table 8 - GRADE

# Figures
# Figure 1 - PRISMA flow
# Figure 2 - Forest plot strict
# Figure 3 - Funnel plot strict
# Figure S1 - Forest plot wide

# References at the end

# Since the text-based parser already included content, we just need to
# save. Tables would need manual placement, which is beyond this simple parser.

output_path = PROJ / 'Plyo训练CMJ高度Meta分析_初稿_v3_updated.docx'
doc.save(str(output_path))
print(f'Saved to: {output_path}')
print('Done! Note: This is a text-only rendering. Tables and figures need manual formatting.')
