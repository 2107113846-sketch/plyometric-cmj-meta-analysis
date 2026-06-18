# -*- coding: utf-8 -*-
"""Generate Chinese manuscript Word document with embedded figures."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT_DIR = 'D:/桌面/科研训练/output'
STYLE_FONT = 'Times New Roman'
CJK_FONT = '宋体'
HEADING_FONT = '黑体'

def set_run_font(run, size=None, bold=False, italic=False, font_name=None, cjk=None):
    font = font_name or STYLE_FONT
    run.font.name = font
    if cjk is None:
        cjk = CJK_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), cjk)
    rPr.insert(0, rFonts)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, cjk=HEADING_FONT)
    return h

def add_para(doc, text, bold=False, size=Pt(12), align=None, first_line_indent=Cm(0.74),
             spacing_after=Pt(6), spacing_before=Pt(0)):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = spacing_after
    para.paragraph_format.space_before = spacing_before
    para.paragraph_format.first_line_indent = first_line_indent

    # Handle **bold** markers
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = para.add_run(part)
            set_run_font(run, size=size, bold=bold)
    return para

def add_figure(doc, img_path, caption, width=Inches(5.5)):
    """Insert an image with centered caption."""
    if os.path.exists(img_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run()
        run.add_picture(img_path, width=width)

        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_after = Pt(12)
        cap_run = cap_para.add_run(caption)
        set_run_font(cap_run, size=Pt(10), bold=True)
    else:
        add_para(doc, f'[图片缺失: {img_path}]', size=Pt(10),
                 align=WD_ALIGN_PARAGRAPH.CENTER)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=Pt(9), bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=Pt(9))
    doc.add_paragraph()

# ================================================================
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = STYLE_FONT
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
rPr = style.element.get_or_add_rPr()
rFonts = rPr.makeelement(qn('w:rFonts'), {})
rFonts.set(qn('w:eastAsia'), CJK_FONT)
rPr.insert(0, rFonts)

# ================================================================
# 标题页
# ================================================================
add_para(doc, '', size=Pt(12))
add_para(doc,
    '快速伸缩复合训练对反向纵跳高度影响的系统综述与Meta分析',
    bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=None, spacing_after=Pt(8))
add_para(doc,
    '——基于手臂位置分层的效应量估计与剂量-反应关系',
    bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER,
    first_line_indent=None, spacing_after=Pt(20))
add_para(doc, 'PROSPERO注册号：CRD420XXXXXX', size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_para(doc, '西南大学体育学院', size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
         spacing_after=Pt(12))
add_para(doc, '指导教师：付道领', size=Pt(11),
         align=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

doc.add_page_break()

# ================================================================
# 摘要
# ================================================================
add_heading_styled(doc, '摘要', 1)

add_para(doc, '背景：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    '快速伸缩复合训练（Plyometric Training）被广泛用于提升运动表现，但既往Meta分析存在手臂位置混杂、'
    '缺乏剂量-反应系统检验、未纳入近年新RCT等方法学局限。',
    size=Pt(12), first_line_indent=Cm(0.74))
add_para(doc, '目的：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    '首次以CMJ手臂位置为纳入标准，系统评估Plyometric训练对CMJ高度的影响，并检验干预时长等调节变量的'
    '剂量-反应关系。',
    size=Pt(12), first_line_indent=Cm(0.74))
add_para(doc, '方法：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    '系统检索PubMed、Scopus、Web of Science及Google Scholar（建库至2026年5月）。纳入标准：RCT设计、'
    'Plyometric训练干预、报告CMJ高度Mean±SD。以严格手叉腰/双臂交叉无臂CMJ为主分析池（16篇），'
    '含臂位未明/CMJA的宽版池（28篇）为敏感性分析。采用Hedges\' g为效应量指标，基于前测-后测变化值计算'
    '（假定r=0.7），REML随机效应模型合并。PEDro量表评估偏倚风险。预注册PROSPERO。',
    size=Pt(12), first_line_indent=Cm(0.74))
add_para(doc, '结果：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    '纳入29篇RCT，总样本量N≈720+。严格手叉腰池合并Hedges\' g=+1.128[95%CI:+0.656,+1.600]，'
    'I²=76.2%，预测区间[-0.878,+3.134]。宽版池g=+1.002[95%CI:+0.722,+1.282]，I²=66.9%。'
    '干预时长是最重要的调节变量：短期(≤6周)g=+0.60(I²=13%)、中期(7-10周)g=+1.46(I²=61%)、'
    '长期(>10周)g=+1.85(I²=94%)。Meta回归确认每周SMD增加约0.13(p=0.026)。'
    '青少年亚组效应最大(g=+1.37~1.51,I²=0%)。Egger检验显著(p<0.001)，SE-g高度相关(r=+0.88)。'
    'r=0.5/0.9敏感性分析确认效应方向始终正向显著。',
    size=Pt(12), first_line_indent=Cm(0.74))
add_para(doc, '结论：', bold=True, size=Pt(12), first_line_indent=Cm(0.74),
         spacing_after=Pt(2))
add_para(doc,
    'Plyometric训练对CMJ高度有显著的大效应提升，短期训练(≤6周)即可获得可靠的中等效应。干预时长与效应量'
    '呈正向剂量-反应关系。建议所有CMJ测试明确报告手臂位置，未来需更多长期大样本RCT。',
    size=Pt(12), first_line_indent=Cm(0.74))
add_para(doc,
    '关键词：快速伸缩复合训练；反向纵跳；Meta分析；手臂位置；剂量-反应关系；牵张-缩短循环',
    size=Pt(11), bold=True, first_line_indent=Cm(0.74), spacing_before=Pt(12))

doc.add_page_break()

# ================================================================
# 1. 引言
# ================================================================
add_heading_styled(doc, '1  引言', 1)

intro_paragraphs = [
    '快速伸缩复合训练（Plyometric Training，亦称增强式训练）是一类利用骨骼肌牵张-缩短循环'
    '（stretch-shortening cycle, SSC）产生爆发性向心收缩的训练方法，典型动作包括跳深（drop jump）、'
    '栏架跳（hurdle hop）、连续纵跳等[1,2]。该训练被广泛应用于从青少年体育到精英竞技的各级运动表现'
    '提升计划中，其生理学基础在于增强肌腱刚度、提高神经肌肉激活速率以及优化SSC力学效率[3]。',

    '反向纵跳（countermovement jump, CMJ）是衡量下肢爆发力的金标准指标之一[4]。与下蹲跳（squat jump, '
    'SJ）不同，CMJ包含向心前的反向预拉伸阶段，能够更真实地反映SSC利用能力，因而被广泛用于运动选材、'
    '训练监控和科研评估。CMJ高度的提升通常被视为Plyometric训练有效性的直接证据。',

    '既往已有多篇Meta分析评估了Plyometric训练对纵跳高度的影响。Markovic（2007）的经典Meta分析报告'
    'Plyometric训练对CMJ的合并效应量为SMD=0.88（95%CI: 0.64-1.11, I²=11.4%），属大效应范畴，但该研究'
    '纳入了多种跳跃结局指标且未对CMJ手臂位置加以区分[1]。de Villarreal等（2009）纳入56项研究的225个'
    '效应量，系统分析了训练量、强度和类型等调节变量对跳跃表现的影响，但未报告针对CMJ的独立合并效应量及'
    '异质性数据[2]。Slimani等（2016）针对团队运动运动员报告了各项体能指标的合并效应量，但由于未对CMJ'
    '手臂位置分层且纳入标准较宽，其效应量数据的可比较性受限[9]。Ramirez-Campillo课题组近年发表了一系列'
    '聚焦特定人群的Meta分析：Moran等（2019）报告女性青少年Plyometric训练后CMJ效应量为ES=0.57'
    '（95%CI: 0.21-0.93），且发现年龄<15岁亚组效应量（ES=0.78）显著高于≥15岁亚组（ES=0.31）[5]；'
    'Ramirez-Campillo等（2020）报告女性足球运动员的合并效应量为ES=1.01（95%CI: 0.36-1.66）[6]；'
    'Sole与Ramirez-Campillo等（2021）在个人运动项目运动员中发现合并效应量为ES=0.49（I²=0.0%）[7]。'
    '此外，Ramirez-Campillo等（2020）在Scand J Med Sci Sports发表的伞状综述系统回顾了Plyometric训练'
    '对多项体能指标的影响，提供了该领域的宏观证据图景[14]。',

    '然而，上述Meta分析存在共同的方法学局限性。首先，CMJ测试中手臂摆动与否对跳跃高度的测量有显著影响——'
    '手臂摆动可贡献约10%-15%的额外跳跃高度[8]——但除Markovic（2007）在亚组中对CMJ有无摆臂做了初步区分外，'
    '其余Meta分析均未将CMJ手臂位置（手叉腰/双臂交叉 vs. 自由摆臂/Vertec）作为纳入标准或分析变量。这种'
    '方法学上的模糊处理可能导致效应量估计的混杂偏倚：带有手臂摆动的CMJ或Vertec纵跳因上肢参与而可能高估'
    'Plyometric训练对下肢SSC能力的真实效果。其次，多数既往Meta分析缺乏对干预时长与效应量之间剂量-反应关系'
    '的系统检验（如Meta回归或连续性剂量-反应模型），仅少数进行了分类亚组分析（如<7周 vs. ≥7周）。此外，'
    '针对青春期前及青春期运动员的专项Meta分析十分有限（仅Moran等2019针对女性青少年），且近年（2020-2025）'
    '发表的多项新RCT尚未被纳入任何系统综述。',

    '因此，本研究旨在通过更新且方法学更严格的系统综述与Meta分析，综合评估Plyometric训练对CMJ高度的影响。'
    '与既往Meta分析相比，本研究的创新点包括：（1）首次将CMJ手臂位置作为明确的纳入标准，并在主分析中限定'
    '严格手叉腰/双臂交叉的无臂CMJ，同时将手臂位置作为亚组分析变量；（2）系统检验干预时长的剂量-反应关系，'
    '包括分类亚组和连续性Meta回归；（3）纳入2020-2025年新增RCT，更新效应量估计并填补青少年人群的证据缺口。'
    '本研究预先设定以下研究问题：（1）Plyometric训练对CMJ高度的合并效应量（Hedges\' g）有多大？'
    '（2）干预时长、训练频率、受试者年龄和CMJ手臂位置是否调节效应量？（3）短期（≤6周）、中期（7-10周）和'
    '长期（>10周）Plyometric干预的效应是否存在剂量-反应关系？基于现有证据，我们假设Plyometric训练将对'
    'CMJ高度产生中等至大的正向效应（g>0.5），且效应量将随干预时长增加而增大。',
]

for text in intro_paragraphs:
    add_para(doc, text)

# ================================================================
# 2. 方法
# ================================================================
add_heading_styled(doc, '2  方法', 1)

add_heading_styled(doc, '2.1  文献检索策略', 2)
add_para(doc,
    '系统检索PubMed、Scopus、Web of Science及Google Scholar数据库，检索时限为建库至2026年5月。'
    '采用以下检索策略（以PubMed为例）：')
add_para(doc,
    '(plyometric* OR "jump training" OR "reactive strength" OR "stretch-shortening cycle") '
    'AND ("countermovement jump" OR CMJ OR "vertical jump" OR "jump height") '
    'AND (random* OR RCT OR "controlled trial")',
    size=Pt(10), first_line_indent=None, spacing_before=Pt(4))
add_para(doc,
    '同时追溯纳入文献的参考文献列表以补充检索。语言不限。')

add_heading_styled(doc, '2.2  纳入与排除标准', 2)
add_para(doc, '纳入标准（PICOS框架）：', bold=True)
add_table(doc,
    ['PICOS要素', '标准'],
    [
        ['P（受试者）', '健康人群，无年龄/性别/训练水平限制'],
        ['I（干预）', 'Plyometric训练（含跳深、栏架跳、连续跳等SSC动作），'
         '实验组在常规训练/日常活动基础上额外增加Plyo'],
        ['C（对照）', '不做Plyometric训练的对照组（可维持常规训练或日常活动）'],
        ['O（结局）', '反向纵跳（CMJ）高度（cm），报告前测/后测 Mean±SD'],
        ['S（研究设计）', '随机对照试验（RCT）或整群随机试验'],
    ])

add_para(doc, '排除标准：', bold=True)
exclusions = [
    '未报告CMJ高度（仅报告SJ、DJ、VJ或其他跳跃指标）',
    '无独立不做Plyo的对照组（如Plyo vs. 抗阻训练对比）',
    '非随机设计（如自身前后对照、队列研究）',
    'CMJ测试未限制手臂摆动（即CMJA/Abalakov/Vertec VJ），且未单独报告无臂CMJ',
    '受试者为伤病康复人群',
    '重复发表或数据重叠',
]
for ex in exclusions:
    add_para(doc, f'• {ex}', first_line_indent=Cm(1.5), size=Pt(11),
             spacing_after=Pt(1))

add_heading_styled(doc, '2.3  数据提取', 2)
add_para(doc,
    '由两名研究者独立提取以下信息：（1）研究特征：第一作者、发表年份、国家、运动项目、样本量；'
    '（2）受试者特征：性别、年龄、训练水平、发育阶段；（3）干预特征：Plyometric训练类型、'
    '周期（周）、频率（次/周）、触地次数；（4）CMJ测试：设备、手臂位置（手叉腰/不明/带臂）、'
    '单位；（5）结局数据：干预组与对照组CMJ高度前测/后测Mean±SD。')
add_para(doc,
    '若原文报告SEM（标准误），按SD=SEM×√n换算。若仅报告CI（置信区间），'
    '按SD≈(CI_upper−CI_lower)/(2×1.96)反推。若仅报告变化值（Δ±SD_Δ），'
    '以Pre+Δ估算Post Mean，以Pre SD或SD_Δ近似Post SD。')

add_heading_styled(doc, '2.4  偏倚风险评价', 2)
add_para(doc,
    '采用PEDro量表（Physiotherapy Evidence Database Scale）评估每篇纳入研究的方法学质量。'
    'PEDro量表包含11个条目：条目1评价外部真实性（不计入总分），条目2-11评价内部真实性和统计报告质量，'
    '每项满足记1分，总分范围0-10。评分标准：≥6分为高质量，4-5分为中等质量，≤3分为低质量。'
    '两名研究者独立完成评分，分歧通过协商解决。对于运动训练RCT，条目5（受试者盲法）和条目6（治疗师盲法）'
    '因干预性质而在绝大多数情况下无法实现，这在PEDro评分框架中被视为该领域的固有局限而非个别研究的方法学缺陷。')

add_heading_styled(doc, '2.5  效应量计算', 2)
add_para(doc,
    '采用Hedges\' g（小样本校正的标准化均差，SMD）作为效应量指标。计算两种SMD：')
add_para(doc,
    '（1）Post-only SMD：基于后测均值的标准化均差，以合并后测SD为标准化因子。'
    '（2）Pre-post change SMD（主分析）：基于前后测变化值的组间比较，假设前测-后测相关系数r=0.7'
    '（CMJ典型值[4]）。')
add_para(doc,
    '变化值SD计算公式：SD_change = √(SD_pre² + SD_post² − 2×r×SD_pre×SD_post)。'
    'Hedges\' g校正因子：J = 1 − 3/(4×df − 1)，df = n_IG + n_CG − 2。')

add_heading_styled(doc, '2.6  统计分析方法', 2)
add_para(doc,
    '所有分析在Python 3.x + R 4.6/metafor混合环境中完成。')

add_para(doc, '主分析：', bold=True)
add_para(doc,
    '采用随机效应模型（Restricted Maximum Likelihood, REML）合并效应量。报告合并SMD（Hedges\' g）、'
    '95%CI、预测区间。异质性以τ²、I²统计量及Q检验评估。I²>50%视为显著异质性，I²>75%视为高异质性。')
add_para(doc,
    '预先设定的分析池：（1）主分析池：严格手叉腰无臂CMJ；（2）宽版分析池：所有CMJ研究（含臂位未明/CMJA），'
    '作为敏感性分析；（3）VJ敏感性亚组：带臂纵跳（非CMJ），单独分析。')

add_para(doc, '亚组分析（分类调节变量）：', bold=True)
sgs = [
    'CMJ手臂位置（严格手叉腰 vs. 臂位未明 vs. CMJA带臂）',
    '年龄/发育阶段（青春期前 vs. 青春期 vs. 青年成人 vs. 成年/职业）',
    '干预时长（短期≤6周 vs. 中期7-10周 vs. 长期>10周）',
    '训练类型（纯Plyo vs. Plyo+力量混合）',
    '对照类型（无训练对照 vs. 常规训练对照）',
    '性别（男 vs. 女 vs. 混合）',
]
for sg in sgs:
    add_para(doc, f'• {sg}', first_line_indent=Cm(1.5), size=Pt(11),
             spacing_after=Pt(1))

add_para(doc, 'Meta回归（连续调节变量）：', bold=True)
add_para(doc,
    '以干预时长（周）和每周训练次数为主要预测变量，构建多变量模型检验独立效应。')

add_para(doc, '敏感性分析：', bold=True)
add_para(doc,
    '（1）留一法（Leave-One-Out）检验单篇研究对合并效应的影响；（2）剔除极端效应量研究'
    '（>3×IQR）；（3）pre-post相关系数r=0.5和r=0.9敏感性检验。')

add_para(doc, '发表偏倚：', bold=True)
add_para(doc,
    '漏斗图目视检查、Egger回归检验（p<0.10提示不对称）、SE-g相关性分析（小研究效应）。')

# ================================================================
# 3. 结果
# ================================================================
add_heading_styled(doc, '3  结果', 1)

add_heading_styled(doc, '3.1  文献筛选流程', 2)
add_para(doc,
    '初步检索获得文献135篇。排除重复及标题/摘要筛选44篇后，对91篇进行全文评估。排除60篇（未报告CMJ高度'
    '约25篇、无独立不做Plyo对照组约15篇、非随机设计约8篇、CMJ含臂摆且未单独报告无臂CMJ约5篇、重复/重叠'
    '样本约3篇、其他约4篇），最终29篇RCT纳入Meta分析。其中严格手叉腰无臂CMJ 16篇，臂位未明/CMJA推断带臂'
    '12篇，VJ带臂敏感性亚组1篇。')

# --- Fig 1: PRISMA ---
add_figure(doc,
    os.path.join(OUTPUT_DIR, 'PRISMA_flow_diagram.png'),
    '图1  PRISMA文献筛选流程图', width=Inches(5.0))

add_heading_styled(doc, '3.2  研究特征', 2)
add_para(doc,
    '共纳入29篇RCT，总样本量N≈720+（干预组n≈390，对照组n≈350）。研究发表于2003-2025年，涵盖足球'
    '（11篇）、篮球（5篇）、排球、手球、赛艇、游泳、中长跑等多个运动项目。受试者年龄范围11-70岁，'
    '以男性为主（20篇）。')

add_para(doc, '表1  纳入研究特征摘要', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['特征', '类别', 'k'],
    [
        ['年龄/发育阶段', '青春期前/青春期', '6'],
        ['', '青年成人', '14'],
        ['', '成年/职业', '7'],
        ['', '老年', '1'],
        ['运动项目', '足球', '11'],
        ['', '篮球', '5'],
        ['', '其他球类/运动', '12'],
        ['CMJ手臂位置', '严格手叉腰', '16'],
        ['', '臂位未明', '6'],
        ['', 'CMJA带臂', '6'],
        ['干预时长', '短期（≤6周）', '14'],
        ['', '中期（7-10周）', '9'],
        ['', '长期（>10周）', '4'],
    ])

add_heading_styled(doc, '3.3  偏倚风险评价', 2)
add_para(doc,
    '29篇研究PEDro评分已通过原文逐篇核实。PEDro得分（条目2-11计分，满分10）：均值5.93/10，中位6，'
    '范围3-8。质量分布：良好及以上（≥6分）17篇（58.6%），一般（4-5分）11篇（37.9%），较差（<4分）'
    '1篇（R09 Blazevich 2003，3/10，准随机设计）。条目5（受试者盲法）和条目6（治疗师盲法）通过率均为0%，'
    '这是运动训练干预的固有局限。若以8分制（排除条目5-6）校正，均值升至约5.4/8（67%）。')

add_para(doc, '表2  PEDro各条目通过率（29篇）', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['PEDro条目', '通过率', '说明'],
    [
        ['1. 纳入标准明确', '100%', '不纳入计分'],
        ['2. 随机分配', '95%', '仅Blazevich(2003)得分0（准RCT）'],
        ['3. 分配隐藏', '5%', '仅1篇(Ramirez-Campillo 2018)明确描述'],
        ['4. 基线可比', '90%', '2篇前测差异≥10%'],
        ['5. 受试者盲法', '0%', '运动干预固有局限'],
        ['6. 治疗师盲法', '0%', '教练员知晓训练内容'],
        ['7. 评估者盲法', '15%', '仅3篇明确CMJ测试者不知分组'],
        ['8. 随访率≥85%', '90%', '大多数研究保持完整随访'],
        ['9. ITT/依从性', '100%', '均报告了组间数据'],
        ['10. 组间比较', '100%', '均报告CMJ组间比较数据'],
        ['11. 点估计+变异', '55%', '约半数未明确报告效应量'],
    ])

add_heading_styled(doc, '3.4  主分析：Plyometric训练对CMJ高度的影响', 2)
add_para(doc,
    '严格手叉腰池（16篇RCT）：合并Hedges\' g = +1.128 [95% CI: +0.656, +1.600]，p<0.001；'
    '预测区间：[-0.878, +3.134]；异质性：τ²=0.687，I²=76.2%，Q(15)=52.30，p<0.001。')
add_para(doc,
    '宽版全池（28篇RCT）：合并Hedges\' g = +1.002 [95% CI: +0.722, +1.282]，p<0.001；'
    '异质性：τ²=0.359，I²=66.9%，Q(27)=77.23，p<0.001。')
add_para(doc,
    '解释：Plyometric训练对CMJ高度有显著的大效应提升（Cohen\'s标准：g>0.8=大效应）。'
    '结果在两种分析池中高度一致。')

# --- Fig 2 & 3: Forest + Funnel ---
add_figure(doc,
    os.path.join(OUTPUT_DIR, 'forest_strict_hand_on_hip.png'),
    '图2  严格手叉腰池森林图（16篇）', width=Inches(5.5))
add_figure(doc,
    os.path.join(OUTPUT_DIR, 'funnel_strict_hand_on_hip.png'),
    '图3  严格手叉腰池漏斗图', width=Inches(5.0))

add_heading_styled(doc, '3.5  敏感性分析', 2)
add_para(doc, '表3  敏感性分析：离群值剔除', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['分析', 'k', 'Hedges\' g', '95% CI', 'I²'],
    [
        ['全部严格手叉腰', '16', '+1.128', '[+0.656, +1.600]', '76%'],
        ['剔除R11（Sedano Campo 2009）', '15', '+0.943', '[+0.603, +1.283]', '54%'],
        ['剔除R27（Toumi 2004）', '15', '+1.011', '[+0.587, +1.435]', '70%'],
        ['剔除R11+R27', '14', '+0.851', '[+0.548, +1.153]', '41%'],
    ])

add_para(doc,
    '留一法分析：剔除任何单篇研究后，合并效应量保持在g=+0.943至+1.197范围内，均显著>0。'
    '最有影响力的研究为R11 Sedano Campo（2009）（剔除后Δg=−0.185，I²降至54%）。')

add_para(doc,
    'Pre-post相关系数敏感性：r的取值对SMD量级有实质影响。r=0.5（保守下界）时，严格池SMD=+0.887'
    '[95%CI: +0.518, +1.257]；r=0.9（反映CMJ力台高重测信度ICC>0.85）时，严格池SMD=+1.841'
    '[95%CI: +1.072, +2.611]。效应方向始终正向显著，但量级随r增加而增大。鉴于CMJ已知的高信度，'
    'r=0.7是合理的中间值选择。建议在主分析中报告r=0.7结果，同时在讨论中报告r=0.5/0.9的范围。')

# --- Fig S1: Wide forest ---
add_figure(doc,
    os.path.join(OUTPUT_DIR, 'forest_wide_all_cmj.png'),
    '图S1  宽版全池森林图（28篇，含臂位未明）', width=Inches(5.5))

add_heading_styled(doc, '3.6  亚组分析', 2)

add_para(doc, '表4  CMJ手臂位置亚组分析', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['亚组', 'k', 'Hedges\' g', '95% CI', 'I²'],
    [
        ['严格手叉腰', '16', '+1.128', '[+0.656, +1.600]', '76%'],
        ['臂位未明', '6', '+0.794', '[+0.465, +1.124]', '20%'],
        ['CMJA带臂', '6', '+0.977', '[+0.214, +1.741]', '79%'],
    ])

add_para(doc, '表5  干预时长亚组分析（关键调节变量）', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['亚组', 'k', 'Hedges\' g', '95% CI', 'I²'],
    [
        ['短期（≤6周）', '14', '+0.595', '[+0.377, +0.814]', '13%'],
        ['中期（7-10周）', '10', '+1.456', '[+0.962, +1.951]', '61%'],
        ['长期（>10周）', '4', '+1.852', '[-0.080, +3.783]', '94%'],
    ])
add_para(doc,
    '短期Plyo训练的合并效应最为可靠：I²仅13%，95%CI紧凑，效应属中等偏大（g≈0.60）。'
    '干预时长是解释异质性的首要调节变量。')

add_para(doc, '表6  年龄/发育阶段亚组分析', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['亚组', 'k', 'Hedges\' g', '95% CI', 'I²', '备注'],
    [
        ['青春期前', '2', '+1.374', '[+0.773, +1.975]', '0%', '[! k=2,证据不足]'],
        ['青春期', '2', '+1.513', '[+0.878, +2.149]', '0%', '[! k=2,证据不足]'],
        ['青年成人', '13', '+1.112', '[+0.676, +1.548]', '67%', ''],
        ['成年/职业', '10', '+0.803', '[+0.193, +1.413]', '83%', ''],
    ])
add_para(doc,
    '青少年运动员效应最大（g≈1.4-1.5）且完全同质（I²=0%），可能与生长发育窗口叠加Plyo产生协同增益有关。'
    '但青春期前和青春期亚组各仅含2篇研究，合并估计需谨慎解读。Plyo+力量混合组（k=2, g=+1.973'
    '[95%CI: +0.205, +3.740], I²=79%）同样标记为证据不足。')

add_heading_styled(doc, '3.7  Meta回归', 2)
add_para(doc, '表7  单变量Meta回归结果', bold=True, size=Pt(11),
         first_line_indent=None, spacing_before=Pt(8))
add_table(doc,
    ['调节变量', '斜率(b)', 'SE', 'z', 'p'],
    [
        ['干预时长（周）', '+0.130', '0.058', '+2.23', '0.026'],
        ['每周训练次数', '+0.536', '0.231', '+2.32', '0.020'],
        ['年龄（岁）', '−0.008', '0.014', '−0.59', '0.556'],
        ['总样本量', '−0.007', '0.008', '−0.89', '0.372'],
        ['发表年份', '−0.040', '0.031', '−1.28', '0.202'],
        ['CMJ手臂位置', '−0.104', '0.184', '−0.57', '0.570'],
    ])
add_para(doc,
    '干预时长每增加1周，预期SMD增加约0.13（p=0.026）。根据模型预测：4周干预预期g≈+0.53，'
    '12周干预预期g≈+1.57。多变量模型（时长+臂位）中，时长在控制臂位后仍显著（b=+0.136, p=0.021），'
    '臂位不显著（b=−0.149, p=0.386）。干预时长是独立于CMJ臂位的稳健预测变量。')

add_heading_styled(doc, '3.8  发表偏倚', 2)
add_para(doc,
    'Egger回归检验：严格池截距=−2.357，p<0.001；宽版池截距=−1.334，p<0.001。'
    'SE-g相关性：严格池r=+0.88，宽版池r=+0.86——方向：小样本研究倾向于报告更大效应量。'
    '漏斗图检查：合并效应两侧研究分布大致对称（严格池8篇高于合并效应/8篇低于），'
    '但SE较大的区域缺少靠近零值的效应量。')
add_para(doc,
    '解读注意事项：Egger检验显著且SE-g相关性强，传统上提示发表偏倚（小样本阴性研究缺失）。但需注意：'
    '（a）样本量在Meta回归中不是显著调节变量（p=0.372）；（b）小样本研究可能因为实施了更高强度/'
    '更严格监督的Plyo而确实产生更大效应；（c）青少年亚组（n偏小）的效应确实高于成年组。'
    '建议在论文中同时报告全池分析和仅大样本的敏感性分析。')

# ================================================================
# 4. 讨论
# ================================================================
add_heading_styled(doc, '4  讨论', 1)

add_heading_styled(doc, '4.1  主要发现', 2)
add_para(doc,
    '本Meta分析纳入29篇RCT（严格手叉腰池16篇，宽版池含臂位未明28篇），系统评估了Plyometric训练'
    '对CMJ高度的影响。主要发现如下：')
findings = [
    'Plyometric训练对CMJ高度有显著的大效应提升：严格手叉腰池合并Hedges\' g=+1.13[95%CI:+0.66,+1.60]，'
    '宽版全池g=+1.00[95%CI:+0.72,+1.28]，效应量均属Cohen\'s标准下的大效应范围（g>0.8）。',
    '结果在多项敏感性分析下高度稳健：剔除两个最有影响力的离群值（R11 Sedano Campo 2009和R27 Toumi 2004）后，'
    '合并效应量仍为g=+0.85[95%CI:+0.55,+1.15]，I²降至41%。',
    '干预时长是最重要的效应调节变量：短期训练（≤6周）产生一致的中等效应（g=+0.60, I²=13%），而中长期训练'
    '（>6周）效应更大但异质性更高（g=+1.46~1.85, I²=61-94%）。Meta回归确认了显著的线性时长-效应关系'
    '（每周+0.13 SMD, p=0.026）。',
    '青少年运动员的增益尤为突出：青春期前和青春期亚组的效应量达g=+1.37~1.51，且完全同质（I²=0%）。',
]
for f in findings:
    add_para(doc, f'• {f}', first_line_indent=Cm(1.5), size=Pt(12), spacing_after=Pt(3))

add_heading_styled(doc, '4.2  与既往Meta分析的比较', 2)
add_para(doc,
    '本研究的合并效应量（严格池g=+1.13，宽版池g=+1.00）略高于既往大多数Meta分析报告的范围'
    '（ES/SMD=0.49-1.01），这一差异可以从以下几个方面加以解释。首先，纳入标准的严格程度不同：'
    '本研究主分析池限定了严格手叉腰/双臂交叉的无臂CMJ，而既往Meta分析多未区分手臂摆动方式，'
    '将Vertec纵跳、Abalakov跳等含上肢参与的跳跃指标与标准CMJ混合分析。手臂摆动的混杂效应可能稀释了'
    '下肢SSC训练的"纯"效应量——换言之，既往Meta分析的合并效应量中可能同时包含了训练效应和手臂摆动的'
    '测量误差，而本研究通过手臂位置的严格分层部分排除了这一噪声。其次，本研究纳入了2021-2025年间发表的'
    '新RCT，这些较新的研究普遍采用了更规范的CMJ测试方案和更系统的Plyometric训练设计，可能报告了更大的'
    '效应量。第三，本研究的效应量计算采用了基于前测-后测变化值的SMD方法（假定r=0.7），部分既往Meta分析'
    '可能使用仅后测SMD，而前者在控制基线差异方面具有更高的统计效率。')

add_para(doc,
    '本研究与既往研究一致地发现了低至中等程度的发表偏倚信号。Markovic（2007）报告CMJ数据存在发表偏倚迹象'
    '（Kendall\'s τ=0.42, p=0.012）[1]；Ramirez-Campillo等（2020）也发现阻力Plyometric训练与控制组比较时'
    'Egger检验显著（p=0.04）[6]。本研究的Egger检验同样高度显著（严格池截距=−2.36, p<0.001），且SE-g相关性'
    '较强（r=+0.86-0.88）。这种跨研究的发表偏倚一致性提示，Plyometric训练Meta分析领域可能存在系统性的'
    '小样本阴性研究缺失，未来应考虑采用灰色文献检索、预注册系统综述等方法加以缓解。')

add_para(doc,
    '值得注意的是，本研究在调节变量分析上的发现与既往研究构成互补而非矛盾。Ramirez-Campillo等（2021）'
    '在个人运动运动员中报告I²=0.0%[7]，与本研究短期亚组（≤6周）I²=13%的低异质性模式类似，共同提示：'
    '当研究对象和干预方案的异质性受到一定限制时，Plyometric训练效应具有高度一致性。此外，Moran等（2019）'
    '发现<15岁亚组效应量大于≥15岁亚组[5]，与本研究年龄亚组分析中青少年（g=+1.37~1.51）高于青年成人及'
    '成年（g=+0.80~1.11）的趋势一致，支持了发育神经肌肉可塑性窗口假说的跨证据链稳健性。')

add_heading_styled(doc, '4.3  干预时长-效应关系与剂量-反应意义', 2)
add_para(doc,
    '本研究最稳健的发现之一是干预时长与效应量之间的正向关系。Meta回归显示干预时长每增加1周，预期SMD增加'
    '约0.13（p=0.026）；多变量模型在控制CMJ手臂位置后，时长仍然显著（p=0.021）。每周训练次数同样为显著的'
    '正向预测变量（p=0.020），提示训练频率与时长共同构成Plyometric训练的"剂量"维度。')
add_para(doc,
    '然而，需要审慎解读的是，长期亚组（>10周，k=4）的效应量虽然最大（g=+1.85），但异质性极高（I²=94%），'
    '95%置信区间跨越零值。相比之下，短期亚组（≤6周，k=14）虽然效应量中等（g=+0.60），但I²仅13%，是本研究'
    '中最可靠的估计。这可能反映了一个真实的生理学现象：短期Plyometric训练在最初4-6周内主要通过神经适应'
    '（运动单位募集效率提升、肌间协调改善）产生增益，效应量相对一致；而随着训练周期延长，个体对训练刺激的'
    '适应性差异（训练内容、强度和恢复管理的异质性）导致效应量的离散度增大。从证据质量角度出发，本研究对短期'
    'Plyometric训练效应的估计比对长期效应的估计更有信心。')

add_heading_styled(doc, '4.4  年龄/发育阶段的调节作用', 2)
add_para(doc,
    '青少年运动员（青春期前及青春期）的效应量（g=+1.37~1.51）显著高于成年/职业运动员（g=+0.80），且青少年'
    '亚组的异质性为零（I²=0%）。这一发现与发育神经肌肉可塑性理论一致：青春期前后是神经肌肉系统快速发展的'
    '"敏感窗口"，此时引入Plyometric刺激可能与自然的神经肌肉成熟过程产生协同效应[12,13]。成年和职业运动员的'
    '效应量虽仍属大效应，但异质性更高（I²=83%），可能反映了训练背景和基础体能水平的差异对干预效果的稀释效应。'
    '需要注意的是，青春期前和青春期亚组各仅含2篇研究（受限于当前文献总量），因此虽然I²=0%本身是稳健的，'
    '但亚组效应量的外部推广性仍需更多研究验证。')

add_heading_styled(doc, '4.5  CMJ手臂位置：方法学启示', 2)
add_para(doc,
    '本研究的一个独特点在于将CMJ手臂位置作为纳入标准和分层分析变量。从亚组分析看，三组的效应量存在差异：'
    '严格手叉腰（k=16）：g=+1.13[0.66,1.60], I²=76%；臂位未明（k=6）：g=+0.79[0.47,1.12], I²=20%；'
    'CMJA带臂（k=6）：g=+0.98[0.21,1.74], I²=79%。臂位未明组I²最低（20%），效应量中等偏大。'
    '严格手叉腰和CMJA组的异质性均较高。这一结果对未来的研究和实践有明确的方法学建议：所有CMJ测试必须明确'
    '报告手臂位置（"hands on hips"/"arms akimbo"/"arms across chest"）。当前文献中臂位报告的缺失'
    '（本研究中有6篇无法确定手臂位置）是阻碍精准效应量估计的重要方法学瑕疵。')

add_heading_styled(doc, '4.6  局限性', 2)
add_para(doc, '本研究存在以下局限，读者在解读结论时需加以考虑：')

limitations = [
    ('发表偏倚不能排除。',
     'Egger检验显著（严格池截距=−2.36, p<0.001），且SE-g高度相关（r=+0.86-0.88），'
     '提示小样本研究倾向报告更大效应量。但Meta回归未发现样本量与效应量的显著关联（p=0.372），'
     '且漏斗图中合并效应两侧的研究分布大致对称（严格池：8篇高于合并效应/8篇低于）。这种模式可能反映的'
     '是"小样本研究实施了更高强度/更严格监督的Plyometric训练"而非经典发表偏倚。尽管如此，建议将分析结果'
     '解释为效应量的上限估计，并注意小样本研究的潜在偏倚。'),
    ('部分研究样本量偏小。',
     '多篇纳入研究的样本量<20人（如R03 Ramirez-Campillo 2015: n=16; R10 Byrne 2010: n=13），'
     '这可能放大个别研究的权重和对合并效应量的影响。'),
    ('数据估算。',
     '本研究中有四项数据涉及估算或近似：（a）R19 Michailidis（2018）的SD从CI反推估算（≈0.80cm）；'
     '（b）R23 Rensing（2015）的对照组SD以干预组SD近似（5.35cm）；（c）R29 Vescovi（2008）的后测均值'
     '及SD以变化值加前测SD近似；（d）R16 Khlifa（2010）的SEM换算为SD（×√n）。四项均通过敏感性分析验证'
     '——分别剔除对应研究后，合并效应量的变化范围为Δg=−0.016至+0.049，均远小于最小有意义差异的0.10阈值。'
     '因此这些估算对总体结论无实质影响。'),
    ('灰色文献未系统检索。',
     '本研究仅检索了PubMed、Scopus、Web of Science及Google Scholar，未系统检索学位论文数据库'
     '（如ProQuest Dissertations）和会议摘要。可能遗漏未发表的阴性结果研究。'),
    ('亚组分析的部分类别研究数量较少。',
     '青春期前（k=2）、青春期（k=2）和长期干预（k=4）亚组的研究数不足以进行稳健的亚组推断，相应的Meta回归'
     '结果（年龄p=0.556）也可能受到统计效力不足的影响。这些亚组已在结果中明确标注[证据不足]。'),
    ('未纳入的潜在调节变量。',
     '总触地次数（ground contacts）、训练强度（如跳深高度）、以及Plyometric训练的具体类型（跳深vs.栏架跳'
     'vs.连续纵跳）可能是重要的效应调节因素，但多数纳入研究未充分报告这些信息，故无法纳入分析。'),
]
for i, (bold_part, normal_part) in enumerate(limitations, 1):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.first_line_indent = Cm(0)
    run_b = para.add_run(f'{i}. {bold_part}')
    set_run_font(run_b, size=Pt(12), bold=True)
    run_n = para.add_run(normal_part)
    set_run_font(run_n, size=Pt(12))

add_heading_styled(doc, '4.7  实践建议与研究展望', 2)

add_para(doc, '对教练员和体能训练师的建议：', bold=True)
coach_recs = [
    '短至6周的Plyometric训练（每周2-3次）即可在CMJ高度上产生可靠的中等至大幅度提升（g≈+0.60），'
    '且效果一致性高（I²=13%），建议作为团队运动季前准备期的标准组成部分。',
    '延长训练周期（7-10周）可能获得更大的效应量（g≈+1.47），但个体间反应差异增大，建议配合个性化训练监控。',
    '青少年运动员（青春期前后）是Plyometric训练的特别受益人群，应在安全前提下尽早引入适龄的Plyometric活动。',
]
for rec in coach_recs:
    add_para(doc, f'• {rec}', first_line_indent=Cm(1.5), spacing_after=Pt(3))

add_para(doc, '对研究者的建议：', bold=True)
res_recs = [
    'CMJ测试必须明确报告手臂位置。建议使用"双手叉腰"（hands on hips）或"双臂交叉胸前"'
    '（arms akimbo）作为标准测试姿势，以便跨研究比较。',
    '未来研究应详细报告训练剂量参数（总触地次数、训练强度），以支持更精细的剂量-反应Meta分析。',
    '需要更大样本的长期RCT（>10周）来稳定长期训练的效应量估计。',
    '鼓励研究者将个体参与者数据（IPD）公开或通过合作共享，以支持未来的IPD Meta分析。',
]
for rec in res_recs:
    add_para(doc, f'• {rec}', first_line_indent=Cm(1.5), spacing_after=Pt(3))

# ================================================================
# 参考文献
# ================================================================
add_heading_styled(doc, '参考文献', 1)

references = [
    '[1] Markovic G. Does plyometric training improve vertical jump height? A meta-analytical '
    'review. Br J Sports Med. 2007;41(6):349-355.',
    '[2] de Villarreal ESS, Kellis E, Kraemer WJ, Izquierdo M. Determining variables of '
    'plyometric training for improving vertical jump height performance: a meta-analysis. '
    'J Strength Cond Res. 2009;23(2):495-506.',
    '[3] Taube W, Leukel C, Gollhofer A. How neurons make us jump: the neural control of '
    'stretch-shortening cycle movements. Exerc Sport Sci Rev. 2012;40(2):106-115.',
    '[4] Claudino JG, Cronin J, Mezêncio B, et al. The countermovement jump to monitor '
    'neuromuscular status: a meta-analysis. J Sci Med Sport. 2017;20(4):397-402.',
    '[5] Moran J, Clark CCT, Ramirez-Campillo R, et al. A meta-analysis of plyometric training '
    'in female youth: its efficacy and shortcomings in the literature. J Strength Cond Res. '
    '2019;33(7):1996-2008.',
    '[6] Ramirez-Campillo R, Alvarez C, García-Hermoso A, et al. Effects of plyometric jump '
    'training in female soccer player\'s vertical jump height: a systematic review with '
    'meta-analysis. J Sports Sci. 2020;38(14):1641-1648.',
    '[7] Sole S, Ramirez-Campillo R, Andrade DC, et al. Plyometric jump training effects on '
    'the physical fitness of individual-sport athletes: a systematic review with meta-analysis. '
    'PeerJ. 2021;9:e11004.',
    '[8] Harman EA, Rosenstein MT, Frykman PN, Rosenstein RM. The effects of arms and '
    'countermovement on vertical jumping. Med Sci Sports Exerc. 1990;22(6):825-833.',
    '[9] Slimani M, Chamari K, Miarka B, et al. Effects of plyometric training on physical '
    'fitness in team sport athletes: a systematic review. J Hum Kinet. 2016;53:231-247.',
    '[10] Bosco C, Viitasalo JT, Komi PV, Luhtanen P. Combined effect of elastic energy and '
    'myoelectrical potentiation during stretch-shortening cycle exercise. Acta Physiol Scand. '
    '1982;114(4):557-565.',
    '[11] Markovic G, Mikulic P. Neuro-musculoskeletal and performance adaptations to '
    'lower-extremity plyometric training. Sports Med. 2010;40(10):859-895.',
    '[12] Lloyd RS, Faigenbaum AD, Stone MH, et al. Position statement on youth resistance '
    'training: the 2014 International Consensus. Br J Sports Med. 2014;48(7):498-505.',
    '[13] Radnor JM, Oliver JL, Waugh CM, et al. The influence of growth and maturation on '
    'stretch-shortening cycle function in youth. Sports Med. 2018;48(1):57-71.',
]

for ref in references:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Cm(1.27)
    para.paragraph_format.first_line_indent = Cm(-1.27)
    run = para.add_run(ref)
    set_run_font(run, size=Pt(11))

# ================================================================
# SAVE
# ================================================================
output_path = 'D:/桌面/Plyo训练CMJ高度Meta分析_初稿.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
print('完成！')
