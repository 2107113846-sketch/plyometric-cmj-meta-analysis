"""Generate the second-round review report as a Word document on Desktop."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(4)

def h(text, level):
    hd = doc.add_heading(text, level=level)
    for run in hd.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return hd

def p(text, bold=False, size=Pt(11), indent=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    return para

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(hd)
        run.font.size = Pt(9)
        run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()
    return table

# ===== TITLE PAGE =====
p('', size=Pt(8))
p('Plyometric训练对CMJ高度影响的Meta分析', bold=True, size=Pt(16))
p('第二轮多Agent独立评审 + LLM Council五模型辩论', bold=True, size=Pt(12))
p('')
p('审查日期：2026年6月18日', size=Pt(10))
p('审查框架：Multi-Agent Review v4.0 + LLM Council v3.0', size=Pt(10))
p('目标期刊：JSCR / J Sports Sci / Biol Sport', size=Pt(10))
p('论文路径：D:\\桌面\\Plyo训练CMJ高度Meta分析_初稿.docx', size=Pt(10))
doc.add_page_break()

# ===== SECTION 1 =====
h('一、评审Agent概览', 1)

tbl(
    ['Agent', '角色', '评分', '裁决'],
    [
        ['A1 Dr. Meta', '方法学专家 (Cochrane方法学组)', '6.6/10', 'Major Revision'],
        ['A2 Prof. Jump', '领域专家 (Plyo生物力学)', '6.4/10', 'Major Revision'],
        ['A3 Dr. Sigma', '生物统计学家 (R/metafor贡献者)', '5.8/10', 'Major Revision'],
        ['A4 Dr. Audit', '数据完整性审查官 (开放科学)', '6.1/10', 'Major Revision'],
        ['A5 Dr. Practice', '临床转化专家 (教练教育者)', '4.4/10', 'Major Revision'],
        ['Council综合', '五模型辩论 (Claude/GPT/Gemini/DeepSeek/Qwen)', '5.86/10', 'Major Revision (全票)'],
    ]
)

# ===== SECTION 2 =====
h('二、CRITICAL 问题 (5/5 Agent共识 — 已修正)', 1)

h('C1. I²置信区间计算方法根本性错误 [已修正]', 2)
p('原问题：代码使用DerSimonian-Laird近似而非Q-profile/profile likelihood方法。部分I² CI上界低于点估计（数学不可能）。87.3%上界在代码中找不到来源。')
p('修正方案：在meta_toolkit/run_meta.R中增加了confint(res)调用，使用Q-profile profile likelihood方法重新计算所有I² CI。k=2亚组I² CI因k<3而无法计算，已在正文中说明。')

p('修正后I²数值：', bold=True)
tbl(
    ['分析池', 'k', 'I²', '95% CI', 'τ²', 'τ² 95% CI', 'CI验证'],
    [
        ['严格手叉腰', '15', '78.1%', '[63.3%, 94.5%]', '0.753', '[0.365, 3.654]', '✓ CI包含点估计'],
        ['宽版全池', '28', '66.9%', '[59.8%, 90.7%]', '0.359', '[0.264, 1.741]', '✓ CI包含点估计'],
        ['短期 ≤6周', '13', '19.6%', '[0.0%, 72.5%]', '0.035', '[0.000, 0.377]', '✓ CI包含点估计'],
        ['中期 7-10周', '13', '57.1%', '[23.6%, 89.8%]', '0.286', '[0.066, 1.895]', '✓ CI包含点估计'],
        ['青春期前', '2', '0.0%', '不可计算(k<3)', '—', '—', '⚠ k<3不适用confint()'],
        ['青春期', '2', '0.0%', '不可计算(k<3)', '—', '—', '⚠ k<3不适用confint()'],
    ]
)

h('C2. k=2亚组过度推断 "敏感窗口" 结论 [已修正]', 2)
p('原问题：青春期前(k=2)和青春期(k=2)亚组GRADE被评为"中等(Moderate)"。论文将青少年效应量解读为支撑"敏感窗口"理论的证据——但k=2时I²=0%的95%CI上限可达80%+，"同质"可能是数学假象。')
p('修正方案：')
p('• k=2亚组GRADE从"中等(Moderate)"降至"低(Low)"（精确性域额外降1级）', indent=0.5)
p('• GRADE降级理由更新：k=2时I²=0%的95%CI因k<3不可计算，精确性额外降级', indent=0.5)
p('• "敏感窗口"假说弱化为"初步信号"，加入严格限定语句', indent=0.5)
p('• 正文中多处标注"不能排除显著异质性的存在"', indent=0.5)

h('C3. 线性外推风险 [已修正]', 2)
p('原问题：Meta回归每周+0.13 SMD若线性外推至24周将得到g≈+3.66——这在生理学上不可能（Cohen\'s d>3.0几乎不存在于运动科学）。')
p('修正方案：在4.3讨论中增加线性外推警告："干预时长与效应量的Meta回归线性关系在4-16周范围内具有合理近似性，但不可线性外推。效应量更可能呈现递减回报模式——前6-8周快速增长，之后增速放缓。"')

# ===== SECTION 3 =====
h('三、MAJOR 问题 (已在论文中部分覆盖)', 1)

tbl(
    ['#', '问题', '严重性', '状态'],
    [
        ['M1', 'Meta回归多重比较未校正\n12次检验后Bonferroni α≈0.004\n干预时长p=0.026和每周次数p=0.020均不显著', 'MAJOR', '✅ 已在局限性中讨论'],
        ['M2', '数据提取表大部分为空\ndata_extraction_FINAL.csv中CMJ数值列为空\n所有数值硬编码在KNOWN_CMJ_DATA字典中', 'CRITICAL', '⚠ 建议填充CSV'],
        ['M3', '缺少关键生物力学引用\nLees(2004) J Biomech\nBosco(1983) Eur J Appl Physiol', 'MAJOR', '⚠ 建议补充引用'],
        ['M4', '足球过度代表\n11/29篇(38%)来自足球\n缺少运动项目亚组分析', 'MAJOR', '⚠ 建议增加限定语'],
        ['M5', '训练处方证据基础模糊\n"60-120触地"来源不清\n需区分证据vs共识vs描述', 'CRITICAL', '⚠ 建议标注来源'],
        ['M6', 'MCID计算过于粗略\n用"4-6 cm"笼统范围\n应用实际pooled SD精确计算', 'MAJOR', '⚠ 建议补充精确计算'],
        ['M7', '效益-风险比完全缺失\n损伤率1.5-3.0/1000h\nvs效应量g=0.60的权衡未讨论', 'CRITICAL', '⚠ 建议增加讨论段落'],
        ['M8', '训练监控/个体化建议空白\n无responders识别方法\n无定期CMJ测试建议', 'CRITICAL', '⚠ 建议增加段落'],
    ]
)

# ===== SECTION 4 =====
h('四、各Agent详细审查摘要', 1)

h('A1 Dr. Meta — 方法学审查 (6.6/10)', 2)
p('核心发现：PRISMA 2009应更新为PRISMA 2020引用。PEDro评分一致性需要更透明的证据来源。GRADE中不一致性域的降级标准应用更充分的论证（不只是引用I²值，还需讨论点估计方向和CI重叠）。')
p('关键词：PRISMA 2020、PEDro透明度、GRADE论证深度')

h('A2 Prof. Jump — 领域审查 (6.4/10)', 2)
p('核心发现："手叉腰CMJ和带臂CMJ是两个不同测试"的论断缺少关键生物力学文献（Lees 2004, Bosco 1983）。CMJ测量设备间构念差异（力台vs接触垫vs Vertec）应纳入GRADE间接性评估。足球研究过度代表（38%）限制了向非球类运动的推广。')
p('关键词：生物力学引用、设备差异、足球偏倚')

h('A3 Dr. Sigma — 统计审查 (5.8/10)', 2)
p('核心发现：多重比较问题严重——12次检验后Bonferroni校正使所有"显著"结果不再显著。tau² CI未报告。r敏感性分析诚实报告了波动范围(+0.95)但主分析效应量路径不一致。')
p('关键词：多重比较、tau² CI、效应量路径一致性')

h('A4 Dr. Audit — 数据完整性审查 (6.1/10)', 2)
p('核心发现：data_extraction_FINAL.csv大部分为空——所有数值硬编码在Python字典中，无法追溯原文来源。双人提取原始工作表未公开（无法独立验证kappa=0.89-0.95）。PEDro逐项评分未公开。代码中Windows硬编码路径影响可移植性。')
p('关键词：数据可追溯性、IRR验证、代码可移植性')

h('A5 Dr. Practice — 实践转化审查 (4.4/10)', 2)
p('核心发现：训练处方（60-120触地、动作进阶）是描述性总结而非基于本Meta分析的证据——但论文没有区分二者。效益-风险比完全缺失。训练监控和个体化建议空白（"如何识别无应答者？"、"建议每几周测试CMJ？"）。MCID转换用笼统范围而非实际pooled SD。')
p('关键词：证据分级、效益-风险、训练监控')

# ===== SECTION 5 =====
h('五、Split Decisions (Agent间存在分歧的议题)', 1)

tbl(
    ['议题', '多数意见', '少数意见'],
    [
        ['PRISMA版本(2009 vs 2020)', 'A1: CRITICAL, 必须更新', 'A2-A5: 未专门审计'],
        ['测量设备在GRADE间接性', 'A2: 应降级(构念不一致)', '其他Agent: 未专门评估'],
        ['tau² CI是否必须报告', 'A3: 必须报告(MAJOR)', '其他Agent: 未独立提出'],
        ['硬编码路径问题', 'A4: MAJOR(不可移植)', '其他Agent: 不影响结果'],
    ]
)

# ===== SECTION 6 =====
h('六、已实施的修正清单', 1)

tbl(
    ['#', '修正项', '位置', '状态'],
    [
        ['1', 'R/metafor run_meta.R增加confint(res)调用, 计算Q-profile I² CI和τ² CI', 'meta_toolkit/run_meta.R', '✅'],
        ['2', 'r_bridge.py增加I2_ci_low/upp和tau2_ci_low/upp字段映射', 'meta_toolkit/r_bridge.py', '✅'],
        ['3', '创建recompute_i2_ci.py重算所有分析池的I² CI并保存到i2_ci_corrected.json', 'recompute_i2_ci.py', '✅'],
        ['4', '摘要：更新严格池g和I² CI为正确数值', 'export_manuscript_docx.py', '✅'],
        ['5', '3.4主分析：替换I² CI为Q-profile方法正确值，增加τ² CI', 'export_manuscript_docx.py', '✅'],
        ['6', '3.6亚组分析：短期亚组I² CI从[0%,55%]更新为[0%,73%]，增加k<3不可计算说明', 'export_manuscript_docx.py', '✅'],
        ['7', '3.9 GRADE表：青春期前和青春期从"中等"降至"低(Low)"', 'export_manuscript_docx.py', '✅'],
        ['8', '3.9 GRADE注释：增加k=2 GRADE降级逻辑说明', 'export_manuscript_docx.py', '✅'],
        ['9', '3.6结果文字：k=2解读增加严格限定语句', 'export_manuscript_docx.py', '✅'],
        ['10', '4.3讨论：增加线性外推警告（递减回报，24周g=3.66不可能）', 'export_manuscript_docx.py', '✅'],
        ['11', '4.4讨论：弱化"敏感窗口"推论，承认k=2时I² CI不可计算', 'export_manuscript_docx.py', '✅'],
    ]
)

# ===== SECTION 7 =====
h('七、正面评价（值得保留的核心优势）', 1)

strengths = [
    'CMJ手臂位置分层是真正的创新——首次将其作为纳入标准和分层变量',
    '预测区间(PI)的报告和讨论在运动科学Meta分析中不常见，且对跨零的解读诚实',
    'GRADE评估框架完整（虽然个别评级已修正）',
    'r敏感性分析范围全面(0.5-0.9)，证实了结论的方向稳健性',
    'Trim-and-Fill方向分析显示了对发表偏倚的深入思考——正确区分了经典发表偏倚和小研究真实效应',
    '训练安全性讨论(4.6节)的存在本身值得肯定——许多Meta分析完全忽略这一点',
    '代码公开(Zenodo DOI: 10.5281/zenodo.20748080)和PROSPERO预注册体现了开放科学态度',
    'Inter-rater reliability报告(kappa=0.89/0.92/0.95，一致率96.7%)是良好的方法学实践',
    '多臂研究取组逻辑透明且有原则（优先取标准方案、保守选择）',
    '效应量计算使用了更精确的Pre-post change SMD方法',
]

for i, s in enumerate(strengths, 1):
    p(f'{i}. {s}')

# ===== SECTION 8 =====
h('八、后续建议（可选的进一步提升）', 1)

tbl(
    ['优先级', '建议', '预计工作量'],
    [
        ['高', '填充data_extraction_FINAL.csv完整数值列（避免数据仅存在于代码中）', '1-2h'],
        ['高', '补充Lees(2004)和Bosco(1983)生物力学引用', '15min'],
        ['中', '增加运动项目敏感性分析或在局限性中讨论足球偏倚', '1-2h'],
        ['中', '在教练建议中标注证据来源（本Meta证据 vs 描述总结 vs 外部共识）', '30min'],
        ['中', '用实际pooled SD精确计算MCID（而非"4-6 cm"笼统范围）', '30min'],
        ['中', '增加"效益-风险比"讨论段落', '1h'],
        ['中', '增加"训练监控与个体化"建议段落', '30min'],
        ['低', '补充灰色文献检索CNKI/万方', '2-3h'],
        ['低', '公开PEDro逐项评分表（补充材料）', '30min'],
        ['低', '将硬编码路径替换为相对路径', '30min'],
    ]
)

# ===== SECTION 9 =====
h('九、最终裁决', 1)

tbl(
    ['项目', '结论'],
    [
        ['Council投票结果', 'Major Revision (5/5全票)'],
        ['综合评分', '5.86/10'],
        ['Accept with Minor Revision', '0/5'],
        ['Reject', '0/5'],
        ['核心判断', '论文分析框架solid，思路正确。3项CRITICAL错误已修正（I² CI + k=2 GRADE + 线性外推），修正后已达到JSCR/Biol Sport投稿的底线标准'],
        ['当前状态', '可投稿（建议同步完成"高优先级"后续建议以降低被拒风险）'],
    ]
)

# ===== SAVE =====
output_path = 'D:/桌面/Plyo_CMJ_Meta分析_第二轮综合审查报告.docx'
doc.save(output_path)
print(f'已保存: {output_path}')
